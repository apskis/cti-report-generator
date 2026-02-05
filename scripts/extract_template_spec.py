"""
Extract a Word template (.docx) into a structured JSON specification.

The output is designed for AI/Claude to read and apply the same formatting
when generating reports (e.g. with python-docx). Run from repo root:

  python scripts/extract_template_spec.py

Output: CTI_Weekly_Report_Template_Spec.json

Use the spec to:
- Apply exact fonts, sizes, colors (color_hex), alignment, spacing (space_before_pt, space_after_pt).
- Recreate section order: header block, Executive Summary, This Week at a Glance (metric tables),
  Vulnerability Exposure table, Sector Threat Activity, Exploitation Indicators, Recommended Actions, footer.
- Match table styling: header row shading, cell shading_hex, and run-level formatting in cells.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

# Allow importing from project root
REPO_ROOT = Path(__file__).resolve().parent.parent
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from docx import Document
from docx.oxml.ns import qn


def _rgb_to_hex(rgb) -> str | None:
    """Convert python-docx RGBColor or (r,g,b) to hex string."""
    if rgb is None:
        return None
    if hasattr(rgb, "rgb"):
        rgb = (rgb.rgb >> 16, (rgb.rgb >> 8) & 0xFF, rgb.rgb & 0xFF)
    if isinstance(rgb, (list, tuple)) and len(rgb) >= 3:
        return f"{int(rgb[0]):02X}{int(rgb[1]):02X}{int(rgb[2]):02X}"
    return None


def _get_cell_shading_hex(cell) -> str | None:
    """Read cell fill color from OOXML (w:tcPr/w:shd @w:fill)."""
    tc_pr = cell._element.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if fill:
        return fill.strip().upper().lstrip("#")
    return None


def _serialize_paragraph(para) -> dict:
    """Serialize a paragraph for JSON: style, alignment, spacing, runs."""
    pf = para.paragraph_format
    # Alignment
    alignment = None
    if para.alignment is not None:
        alignment = str(para.alignment).replace("WD_ALIGN_PARAGRAPH.", "")

    # Spacing (in points when available)
    space_before = getattr(pf, "space_before", None)
    space_after = getattr(pf, "space_after", None)
    line_spacing = getattr(pf, "line_spacing", None)
    if space_before is not None and hasattr(space_before, "pt"):
        space_before = round(space_before.pt, 1)
    if space_after is not None and hasattr(space_after, "pt"):
        space_after = round(space_after.pt, 1)

    runs = []
    for run in para.runs:
        r = {"text": run.text}
        if run.font.name:
            r["font_name"] = run.font.name
        if run.font.size is not None:
            r["font_size_pt"] = round(run.font.size.pt, 1)
        if run.font.bold:
            r["bold"] = True
        if run.font.italic:
            r["italic"] = True
        if run.font.underline is not None:
            r["underline"] = run.font.underline
        rgb = _rgb_to_hex(run.font.color.rgb) if run.font.color.rgb else None
        if rgb:
            r["color_hex"] = rgb
        if r.get("text") or len(r) > 1:
            runs.append(r)

    out = {}
    if para.style and para.style.name:
        out["style"] = para.style.name
    if alignment:
        out["alignment"] = alignment
    if space_before is not None:
        out["space_before_pt"] = space_before
    if space_after is not None:
        out["space_after_pt"] = space_after
    if line_spacing is not None:
        out["line_spacing"] = line_spacing
    if runs:
        out["runs"] = runs
    else:
        # Empty paragraph (spacer)
        full_text = (para.text or "").strip()
        if full_text:
            out["text"] = full_text
        else:
            out["_comment"] = "empty paragraph (spacer)"
    return out


def _serialize_table(table) -> dict:
    """Serialize a table: dimensions, then rows of cells with text and shading."""
    rows_spec = []
    for row in table.rows:
        cells_spec = []
        for cell in row.cells:
            # Cell shading from XML
            fill = _get_cell_shading_hex(cell)
            cell_data = {}
            # All text from cell (merge paragraphs)
            parts = []
            for p in cell.paragraphs:
                p_spec = _serialize_paragraph(p)
                if p_spec.get("runs"):
                    parts.append(" ".join(r.get("text", "") for r in p_spec["runs"]))
                elif p_spec.get("text"):
                    parts.append(p_spec["text"])
            cell_data["text"] = "\n".join(parts).strip() if parts else ""
            if fill:
                cell_data["shading_hex"] = fill
            cells_spec.append(cell_data)
        rows_spec.append({"cells": cells_spec})
    return {"rows": len(table.rows), "cols": len(table.columns), "cells": rows_spec}


def extract_template_spec(docx_path: Path) -> dict:
    """Build a full template specification from a .docx file."""
    doc = Document(str(docx_path))

    # Section/page setup (first section)
    section = doc.sections[0]
    page = {
        "page_width_inches": round(section.page_width.inches, 2),
        "page_height_inches": round(section.page_height.inches, 2),
        "left_margin_inches": round(section.left_margin.inches, 2),
        "right_margin_inches": round(section.right_margin.inches, 2),
        "top_margin_inches": round(section.top_margin.inches, 2),
        "bottom_margin_inches": round(section.bottom_margin.inches, 2),
    }

    # Document body: iterate body elements (paragraphs and tables in order)
    body = []
    for element in doc.element.body:
        tag = element.tag
        if qn("w:p") in tag or tag.endswith("}p"):
            # Find corresponding paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    body.append({"type": "paragraph", "content": _serialize_paragraph(para)})
                    break
        elif qn("w:tbl") in tag or tag.endswith("}tbl"):
            for table in doc.tables:
                if table._tbl is element:
                    body.append({"type": "table", "content": _serialize_table(table)})
                    break

    # If body order is wrong (python-docx order can differ), use paragraph/table iteration
    # and try to preserve order by scanning body XML for p vs tbl
    body_ordered = []
    para_idx = 0
    tbl_idx = 0
    for element in doc.element.body:
        tag = element.tag
        if qn("w:p") in tag or tag.endswith("}p"):
            if para_idx < len(doc.paragraphs):
                body_ordered.append({
                    "type": "paragraph",
                    "content": _serialize_paragraph(doc.paragraphs[para_idx])
                })
                para_idx += 1
        elif qn("w:tbl") in tag or tag.endswith("}tbl"):
            if tbl_idx < len(doc.tables):
                body_ordered.append({
                    "type": "table",
                    "content": _serialize_table(doc.tables[tbl_idx])
                })
                tbl_idx += 1

    # Prefer body_ordered; if it undercounts, fall back to body
    if len(body_ordered) >= len(body):
        body = body_ordered
    else:
        body = body_ordered or body

    return {
        "template_name": docx_path.name,
        "description": "Structured specification of document structure and formatting for AI/code to apply when generating reports.",
        "page_setup": page,
        "body": body,
    }


def main() -> None:
    template_path = REPO_ROOT / "CTI_Weekly_Report_Template_Example.docx"
    if not template_path.exists():
        print(f"Template not found: {template_path}", file=sys.stderr)
        sys.exit(1)

    spec = extract_template_spec(template_path)
    out_path = REPO_ROOT / "CTI_Weekly_Report_Template_Spec.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(spec, f, indent=2, ensure_ascii=False)

    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
