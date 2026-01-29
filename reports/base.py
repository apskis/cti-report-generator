"""
Base report generator class.

Provides common functionality for all report types.
"""
from abc import ABC, abstractmethod
from datetime import datetime
from io import BytesIO
import logging
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

logger = logging.getLogger(__name__)


class BrandColors:
    """Brand color constants matching the template."""

    ORANGE_PRIMARY = RGBColor(0xE6, 0x51, 0x00)  # #E65100 - Main title
    ILLUMINA_BLUE = RGBColor(0x00, 0x5D, 0xAA)  # #005DAA - Illumina Blue
    GRAY_DARK = RGBColor(0x55, 0x55, 0x55)  # #555555 - Body text emphasis
    GRAY_MEDIUM = RGBColor(0x66, 0x66, 0x66)  # #666666 - Subtitles, notes
    RED_CRITICAL = RGBColor(0xFF, 0x00, 0x00)  # Red for critical severity
    ORANGE_HIGH = RGBColor(0xFF, 0xA5, 0x00)  # Orange for high severity
    YELLOW_P1 = "FFFF00"  # Yellow highlight for P1
    GREEN_LOW = RGBColor(0x00, 0x80, 0x00)  # Green for LOW risk
    ORANGE_TABLE_HEADER = "E65100"  # Orange for table headers (hex string for shading)


class FontSizes:
    """Font size constants matching the template."""

    TITLE = Pt(18)
    HEADING_1 = Pt(14)
    HEADING_2 = Pt(12)
    BODY = Pt(10.5)
    BODY_SMALL = Pt(10)
    SUBTITLE = Pt(9)
    FOOTNOTE = Pt(8)


class BaseReportGenerator(ABC):
    """
    Abstract base class for report generators.

    Subclasses must implement:
    - report_type: Property returning the report type name
    - generate(): Method to generate the report document
    """

    def __init__(self):
        self.doc: Document | None = None
        self.created_at = datetime.now()

    @property
    @abstractmethod
    def report_type(self) -> str:
        """Return the report type identifier (e.g., 'weekly', 'monthly')."""
        pass

    @property
    @abstractmethod
    def filename_prefix(self) -> str:
        """Return the filename prefix for this report type."""
        pass

    @abstractmethod
    def generate(self, analysis_result: Dict[str, Any]) -> Document:
        """
        Generate the report document.

        Args:
            analysis_result: Dictionary containing analysis results

        Returns:
            Document object (python-docx Document)
        """
        pass

    def get_filename(self) -> str:
        """Generate the filename for this report."""
        date_str = self.created_at.strftime("%Y-%m-%d")
        return f"{self.filename_prefix}_{date_str}.docx"

    def to_bytes(self) -> bytes:
        """Convert the document to bytes for upload."""
        if self.doc is None:
            raise ValueError("Document not generated. Call generate() first.")

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =========================================================================
    # Common styling utilities
    # =========================================================================

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        """Apply background shading to a table cell."""
        tc_pr = cell._element.get_or_add_tcPr()
        
        # Remove any existing shading first
        existing_shd = tc_pr.find(qn("w:shd"))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Create new shading element
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
        
        # Ensure color_hex doesn't have # prefix and is uppercase
        color_hex = color_hex.replace("#", "").upper()
        shd.set(qn("w:fill"), color_hex)
        shd.set(qn("w:val"), "clear")  # Clear any pattern
    
    def _set_cell_shading_rgb(self, cell, r: int, g: int, b: int) -> None:
        """Apply background shading to a table cell using RGB values."""
        tc_pr = cell._element.get_or_add_tcPr()
        
        # Remove any existing shading first
        existing_shd = tc_pr.find(qn("w:shd"))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Create new shading element
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
        
        # Convert RGB to hex (Word uses hex format)
        # RGB(55, 46, 0) = #372E00
        color_hex = f"{r:02X}{g:02X}{b:02X}"
        shd.set(qn("w:fill"), color_hex)
        shd.set(qn("w:val"), "solid")  # Use solid fill, not clear
        shd.set(qn("w:color"), "auto")  # Auto color for text (not needed for fill)

    def _set_cell_borders(self, cell, color_hex: str = "808080", size: str = "4") -> None:
        """
        Apply borders to a table cell.
        
        Args:
            cell: The table cell to apply borders to
            color_hex: Border color in hex format (default: "808080" for gray)
            size: Border size (default: "4" for thin border)
        """
        tc_pr = cell._element.get_or_add_tcPr()
        
        # Create border elements
        def create_border(side: str):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color_hex)
            return border
        
        # Get or create tcBorders element
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        
        # Add all borders
        tc_borders.append(create_border("top"))
        tc_borders.append(create_border("bottom"))
        tc_borders.append(create_border("left"))
        tc_borders.append(create_border("right"))

    def _apply_severity_color(self, cell, severity: str) -> None:
        """Apply color coding to severity cell text."""
        if not cell.paragraphs[0].runs:
            return

        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        severity_upper = severity.upper()

        if severity_upper == "CRITICAL":
            run.font.color.rgb = BrandColors.RED_CRITICAL
            run.font.bold = True
        elif severity_upper == "HIGH":
            run.font.color.rgb = BrandColors.ORANGE_HIGH
            run.font.bold = True

    def _apply_risk_color(self, cell, risk: str) -> None:
        """Apply color coding to risk level cell."""
        if not cell.paragraphs[0].runs:
            return

        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        risk_upper = risk.upper()

        if risk_upper in ("CRITICAL", "P1"):
            run.font.color.rgb = BrandColors.RED_CRITICAL
            run.font.bold = True
        elif risk_upper in ("HIGH", "P2"):
            run.font.color.rgb = BrandColors.ORANGE_HIGH
            run.font.bold = True

    def _set_paragraph_font(
        self,
        paragraph,
        size: Pt | None = None,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        font_name: str = "Arial"
    ) -> None:
        """Set font properties for a paragraph's runs."""
        for run in paragraph.runs:
            run.font.name = font_name
            if size:
                run.font.size = size
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

    def _add_styled_paragraph(
        self,
        text: str,
        size: Pt = FontSizes.BODY,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        style: str | None = None
    ):
        """Add a paragraph with custom styling."""
        if style:
            para = self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text)

        if alignment:
            para.alignment = alignment

        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = size
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

        return para

    def _create_metric_card_table(
        self,
        metrics: list[tuple[str, str, str]]
    ):
        """
        Create a metric card table (like "4 New This Week" cards).

        Args:
            metrics: List of tuples (number, title, subtitle)
        """
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.autofit = True

        for i, (number, title, subtitle) in enumerate(metrics):
            cell = table.rows[0].cells[i]

            # Clear default paragraph
            cell.paragraphs[0].clear()

            # Add number (large, bold)
            num_para = cell.paragraphs[0]
            num_run = num_para.add_run(number)
            num_run.font.name = "Arial"
            num_run.font.size = Pt(24)
            num_run.font.bold = True
            num_run.font.color.rgb = BrandColors.ORANGE_PRIMARY

            # Add title
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = "Arial"
            title_run.font.size = FontSizes.BODY_SMALL
            title_run.font.bold = True

            # Add subtitle
            sub_para = cell.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = "Arial"
            sub_run.font.size = FontSizes.FOOTNOTE
            sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            sub_run.font.italic = True

            # Center align all paragraphs
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return table

    def _format_date_range(self, start_date: datetime | None = None, end_date: datetime | None = None) -> str:
        """Format a date range string."""
        if end_date is None:
            end_date = self.created_at
        if start_date is None:
            # Default to 7 days before end_date for weekly
            from datetime import timedelta
            start_date = end_date - timedelta(days=6)

        return f"{start_date.strftime('%B %d')} to {end_date.strftime('%d, %Y')}"

    def _get_week_number(self) -> int:
        """Get ISO week number for the report date."""
        return self.created_at.isocalendar()[1]

    def _get_year(self) -> int:
        """Get year for the report date."""
        return self.created_at.year

    def _get_banner_path(self) -> str | None:
        """
        Get the path to the report banner image.
        
        Returns:
            Path to Bulletin_banner.jpg if found, None otherwise
        """
        # Try root directory first
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        banner_path = os.path.join(root_dir, "Bulletin_banner.jpg")
        
        if os.path.exists(banner_path):
            return banner_path
        
        # Try current working directory
        alt_path = "Bulletin_banner.jpg"
        if os.path.exists(alt_path):
            return alt_path
        
        return None

    def _add_image(self, image_path: str, width: Inches | None = None, height: Inches | None = None) -> None:
        """
        Add an image to the document.
        
        Args:
            image_path: Path to the image file
            width: Optional width in Inches (defaults to page width)
            height: Optional height in Inches (maintains aspect ratio if only width specified)
        """
        if not os.path.exists(image_path):
            logger.warning(f"Image file not found: {image_path}")
            return
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run()
        
        # Default to page width if not specified
        if width is None:
            width = Inches(6.5)  # Standard page width minus margins
        
        run.add_picture(image_path, width=width, height=height)

    def _set_default_font(self, font_name: str = "Arial") -> None:
        """
        Set the default font for the entire document.
        
        Args:
            font_name: Name of the font to use (default: "Arial")
        """
        if self.doc is None:
            return
        
        # Set font for Normal style (default paragraph style)
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = font_name
        
        # Also set for Heading styles
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            if style_name in self.doc.styles:
                heading_style = self.doc.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = font_name
                # Set Heading 1 font size to 14pt
                if style_name == 'Heading 1':
                    heading_font.size = Pt(14)

    def _configure_page_settings(self) -> None:
        """
        Configure page margins, header/footer distances, and default paragraph spacing.
        
        Based on the template requirements:
        - Margins: Top 1.25", Bottom 0.88", Left 0.75", Right 0.75"
        - Header from edge: 0.49"
        - Footer from edge: 0.49"
        - Paragraph spacing: Before 0 pt, After 2 pt, Line spacing: Single
        """
        if self.doc is None:
            return
        
        # Get the first (and typically only) section
        section = self.doc.sections[0]
        
        # Set page margins
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(0.88)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Set header and footer distances from edge
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)
        
        # Configure default paragraph spacing for Normal style
        normal_style = self.doc.styles['Normal']
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.space_before = Pt(0)
        normal_paragraph_format.space_after = Pt(2)
        normal_paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Also configure heading styles with the same spacing
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            if style_name in self.doc.styles:
                heading_style = self.doc.styles[style_name]
                heading_paragraph_format = heading_style.paragraph_format
                heading_paragraph_format.space_before = Pt(0)
                heading_paragraph_format.space_after = Pt(2)
                heading_paragraph_format.line_spacing = 1.0

    def _apply_font_to_run(self, run, font_name: str = "Arial") -> None:
        """
        Apply font name to a text run.
        
        Args:
            run: The text run to apply font to
            font_name: Name of the font to use (default: "Arial")
        """
        run.font.name = font_name

    def _add_section_divider(self) -> None:
        """
        Add a faint horizontal divider line below a major section title.
        Creates a paragraph with a bottom border to visually separate sections.
        """
        if self.doc is None:
            return
        
        divider_para = self.doc.add_paragraph()
        divider_para.paragraph_format.space_before = Pt(0)  # No space before (tight with content above)
        divider_para.paragraph_format.space_after = Pt(12)  # Space after (before next section title)
        
        # Add bottom border to create divider line
        p_pr = divider_para._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        
        # Create bottom border (gray, 20% opacity = light gray)
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")  # Thin line
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "CCCCCC")  # Light gray (approximately 20% opacity effect)
        p_bdr.append(bottom_border)

    def _add_banner_header(self) -> None:
        """
        Add the standard banner image to the document header section.
        This method should be called at the start of _add_header() in subclasses.
        """
        banner_path = self._get_banner_path()
        if banner_path:
            # Get the document section and its header
            section = self.doc.sections[0]
            header = section.header
            
            # Clear any existing content in the header
            header.paragraphs[0].clear()
            
            # Add the banner image to the header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = header_para.add_run()
            # Calculate width to fit within margins (page width minus left and right margins)
            # Standard page width is 8.5", minus left (0.75") and right (0.75") margins = 7"
            run.add_picture(banner_path, width=Inches(7.0))
        else:
            logger.warning("Bulletin_banner.jpg not found - banner will be omitted from report")
