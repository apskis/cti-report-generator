"""
Quarterly Strategic CTI Report Generator.

Generates quarterly strategic threat intelligence briefs for leadership.
"""
from datetime import datetime, timedelta
import logging
import os
import re
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reports.base import BaseReportGenerator, BrandColors, FontSizes
from reports.registry import register_report_generator

logger = logging.getLogger(__name__)


class RiskLevel:
    """Risk level constants for quarterly assessments."""
    HIGH = "HIGH"
    MEDIUM = "MEDIUM"
    LOW = "LOW"
    UNCHANGED = "Unchanged"
    INCREASED = "↑"
    DECREASED = "↓"


@register_report_generator("quarterly")
class QuarterlyReportGenerator(BaseReportGenerator):
    """
    Quarterly Strategic CTI Report Generator.

    Generates leadership-focused strategic threat briefs matching
    the CTI_Quarterly_Strategic_Report_Example.docx format.

    Template structure:
    - Report ID header (CTI-QTR-YYYY-QN)
    - Title: "Cyber Threat Intelligence" / "Quarterly Strategic Brief"
    - Quarter date range
    - Executive Summary (strategic overview)
    - Quarterly Risk Assessment (risk metric cards)
    - Industry Breach Landscape (incidents table)
    - Geopolitical Threat Landscape (by nation-state)
    - Looking Ahead (next quarter outlook)
    - Recommendations for Leadership
    - Footer with sources
    """

    @property
    def report_type(self) -> str:
        return "quarterly"

    @property
    def filename_prefix(self) -> str:
        return "CTI_Quarterly_Strategic_Brief"

    def generate(self, analysis_result: Dict[str, Any]) -> Document:
        """
        Generate the quarterly strategic report document.

        Args:
            analysis_result: Dictionary containing:
                - executive_summary: str (strategic overview)
                - risk_assessment: dict with risk levels
                - breach_landscape: dict with incident statistics
                - incidents_by_type: list of incident type breakdowns
                - geopolitical_threats: dict by country (china, russia, north_korea, etc.)
                - looking_ahead: dict with outlook, initiatives, watch_items
                - recommendations: list of strategic recommendations

        Returns:
            Document object
        """
        try:
            logger.info("Generating Quarterly Strategic CTI Report")
            self.doc = Document()
            
            # Set default font to Arial for the entire document
            self._set_default_font("Arial")
            
            # Configure page settings (margins, header/footer distances, paragraph spacing)
            self._configure_page_settings()

            # Calculate quarter info
            self._calculate_quarter_info()

            # Add sections in order
            self._add_header()
            self._add_executive_summary(analysis_result)
            self._add_risk_assessment(analysis_result)
            self._add_breach_landscape(analysis_result)
            self._add_geopolitical_landscape(analysis_result)
            self._add_looking_ahead(analysis_result)
            self._add_leadership_recommendations(analysis_result)
            self._add_footer()

            logger.info("Quarterly Strategic CTI Report generated successfully")
            return self.doc

        except Exception as e:
            logger.error(f"Error generating quarterly report: {str(e)}", exc_info=True)
            raise

    def _calculate_quarter_info(self) -> None:
        """Calculate the quarter's date range and number."""
        today = self.created_at
        # Determine quarter
        month = today.month
        year = today.year

        if month <= 3:
            self.quarter = 1
            self.quarter_start = datetime(year, 1, 1)
            self.quarter_end = datetime(year, 3, 31)
        elif month <= 6:
            self.quarter = 2
            self.quarter_start = datetime(year, 4, 1)
            self.quarter_end = datetime(year, 6, 30)
        elif month <= 9:
            self.quarter = 3
            self.quarter_start = datetime(year, 7, 1)
            self.quarter_end = datetime(year, 9, 30)
        else:
            self.quarter = 4
            self.quarter_start = datetime(year, 10, 1)
            self.quarter_end = datetime(year, 12, 31)

    def _get_previous_quarter(self) -> str:
        """Get the previous quarter string (e.g., 'Q4 2025')."""
        if self.quarter == 1:
            return f"Q4 {self._get_year() - 1}"
        return f"Q{self.quarter - 1} {self._get_year()}"

    def _add_header(self) -> None:
        """Add report header with banner image, ID, title, and date range."""
        year = self._get_year()

        # Add banner image at the top
        self._add_banner_header()

        # Report ID (e.g., CTI-QTR-2026-Q1) - positioned at top-right
        report_id = f"CTI-QTR-{year}-Q{self.quarter}"
        id_para = self.doc.add_paragraph()
        id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        id_run = id_para.add_run(report_id)
        id_run.font.name = "Arial"
        id_run.font.size = FontSizes.SUBTITLE
        id_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        # Main title - Quarterly Strategic Brief (styled like Book Title, orange, 20pt, centered)
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run("Quarterly Strategic Brief")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(20)  # Font size 20pt
        title_run.font.bold = True
        title_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned
        # Reduce spacing after title
        title_para.paragraph_format.space_after = Pt(0)

        # Subtitle - Quarter date range (using Subtitle style, gray, centered)
        start_month = self.quarter_start.strftime("%B")
        end_month = self.quarter_end.strftime("%B")
        date_range = f"Q{self.quarter} {year} ({start_month} to {end_month})"

        subtitle_para = self.doc.add_paragraph(date_range, style="Subtitle")
        for run in subtitle_para.runs:
            run.font.name = "Arial"
            run.font.color.rgb = BrandColors.GRAY_DARK
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned
        # Reduce spacing after subtitle
        subtitle_para.paragraph_format.space_after = Pt(0)

        self.doc.add_paragraph()

    def _add_executive_summary(self, analysis_result: Dict[str, Any]) -> None:
        """Add executive summary section."""
        logger.info("Adding Executive Summary section")

        # Executive Summary heading - Heading 1
        summary_heading = self.doc.add_heading("Executive Summary", level=1)
        for run in summary_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        # Add space after heading
        summary_heading.paragraph_format.space_after = Pt(6)

        # Get summary paragraphs
        summary = analysis_result.get("executive_summary", "")
        if not summary:
            summary = self._generate_default_executive_summary(analysis_result)

        # Split into paragraphs if it's a long string
        paragraphs = summary.split("\n\n") if "\n\n" in summary else [summary]

        for para_text in paragraphs:
            if para_text.strip():
                para = self.doc.add_paragraph(para_text.strip())
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = FontSizes.BODY

        self.doc.add_paragraph()

    def _generate_default_executive_summary(self, analysis_result: Dict[str, Any]) -> str:
        """Generate a default executive summary from available data."""
        stats = analysis_result.get("breach_landscape", {})
        total_incidents = stats.get("total_incidents", 0)
        apt_groups = len(analysis_result.get("geopolitical_threats", {}).get("actors", []))

        return f"""The threat landscape for the genomics, life sciences, and precision manufacturing sectors \
remained elevated throughout Q{self.quarter} {self._get_year()}, with {total_incidents} publicly disclosed \
breaches affecting peer organizations in the industry.

No direct threats to the organization were identified this quarter; however, the threat actors, techniques, \
and vulnerabilities observed are consistent with those historically used against genomics companies. \
{apt_groups} threat actor groups were observed targeting the sector with varying levels of sophistication."""

    def _add_risk_assessment(self, analysis_result: Dict[str, Any]) -> None:
        """Add quarterly risk assessment section with risk cards."""
        logger.info("Adding Quarterly Risk Assessment section")

        # Quarterly Risk Assessment heading - Heading 1
        risk_heading = self.doc.add_heading("Quarterly Risk Assessment", level=1)
        for run in risk_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        # Add space before and after heading
        risk_heading.paragraph_format.space_before = Pt(12)
        risk_heading.paragraph_format.space_after = Pt(6)

        risk_data = analysis_result.get("risk_assessment", {})
        breach_data = analysis_result.get("breach_landscape", {})
        
        # Helper function to calculate percentage change
        def calculate_percentage(current, previous):
            """Calculate percentage change between current and previous values."""
            if not previous or previous == 0:
                return None
            if not current:
                return None
            try:
                change = ((current - previous) / previous) * 100
                return f"+{int(change)}%" if change > 0 else f"{int(change)}%"
            except (TypeError, ValueError):
                return None

        # Create risk assessment cards table in horizontal layout (1 row, 4 columns)
        # Calculate percentages for trends where possible
        ransomware_pct = calculate_percentage(
            breach_data.get("ransomware_count"),
            breach_data.get("prev_ransomware")
        )
        
        risks = [
            (
                "Nation-State Espionage",  # Single line, no newline
                risk_data.get("nation_state", RiskLevel.HIGH),
                risk_data.get("nation_state_trend", RiskLevel.UNCHANGED),
                None,  # No percentage data available
            ),
            (
                "Ransomware & Extortion",  # Single line, no newline
                risk_data.get("ransomware", RiskLevel.HIGH),
                risk_data.get("ransomware_trend", RiskLevel.INCREASED),
                ransomware_pct,  # Use calculated percentage if available
            ),
            (
                "Supply Chain Compromise",  # Single line, no newline
                risk_data.get("supply_chain", RiskLevel.MEDIUM),
                risk_data.get("supply_chain_trend", RiskLevel.UNCHANGED),
                None,  # No percentage data available
            ),
            (
                "Insider Threat",
                risk_data.get("insider", RiskLevel.LOW),
                risk_data.get("insider_trend", RiskLevel.UNCHANGED),
                None,  # No percentage data available
            ),
        ]

        # Create horizontal table (1 row, 4 columns)
        table = self.doc.add_table(rows=1, cols=len(risks))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT  # Left align table like the example
        
        # Set table width to 100%
        tbl = table._element
        # Get or create tblPr element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)  # Insert at the beginning
        
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "5000")  # 5000 = 100% in Word's 50ths of a percent
        tbl_w.set(qn("w:type"), "pct")  # Percentage type

        # Fill table cells with risk data
        for i, (category, level, trend, calculated_pct) in enumerate(risks):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].clear()
            
            # Set cell vertical alignment to center for better alignment
            tc_pr = cell._element.get_or_add_tcPr()
            v_align = tc_pr.find(qn("w:vAlign"))
            if v_align is None:
                v_align = OxmlElement("w:vAlign")
                tc_pr.append(v_align)
            v_align.set(qn("w:val"), "center")  # Center vertically
            
            # Set cell margins for lean/compact spacing
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            
            # Set all margins (top, bottom, left, right) to minimal padding for lean look
            for margin_type in ["top", "bottom", "left", "right"]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), "72")  # 72 twips = 0.05 inches (reduced from 0.1)
                margin.set(qn("w:type"), "dxa")

            # Set cell background to "Light gray - background 2" (Word theme color)
            # This is typically #F2F2F2 or similar very light gray
            self._set_cell_shading(cell, "F2F2F2")  # Light gray - background 2

            # Set cell borders (light gray, thinner - 1/2 pt)
            self._set_cell_borders(cell, color_hex="C0C0C0", size="1")

            # Category name - centered, single line, lean spacing
            cat_para = cell.paragraphs[0]
            cat_para.paragraph_format.space_before = Pt(0)
            cat_para.paragraph_format.space_after = Pt(0)
            # Replace any newlines with spaces to keep on single line
            category_single_line = category.replace("\n", " ")
            cat_run = cat_para.add_run(category_single_line)
            cat_run.font.name = "Arial"
            cat_run.font.size = Pt(9)  # Font size 9 as requested
            cat_run.font.bold = True
            cat_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)  # Darker text for light background (brighter than before)
            cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centered alignment

            # Risk level - now as colored text (not colored background box)
            level_para = cell.add_paragraph()
            level_para.paragraph_format.space_before = Pt(0)
            level_para.paragraph_format.space_after = Pt(0)
            level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            level_run = level_para.add_run(level)
            level_run.font.name = "Arial"
            level_run.font.size = Pt(14)  # Font size 14 as requested
            level_run.font.bold = True
            
            # Set text color based on risk level - use colors that match the example
            risk_level_color = None
            if level == RiskLevel.HIGH:
                risk_level_color = RGBColor(0xF4, 0x7F, 0x7D)  # Pinkish-red like example
                level_run.font.color.rgb = risk_level_color
            elif level == RiskLevel.MEDIUM:
                risk_level_color = RGBColor(0xFF, 0xA5, 0x00)  # Orange
                level_run.font.color.rgb = risk_level_color
            elif level == RiskLevel.LOW:
                risk_level_color = RGBColor(0x7F, 0xCB, 0x7F)  # Green like example
                level_run.font.color.rgb = risk_level_color
            else:
                # Default styling if level doesn't match
                level_run.font.color.rgb = RGBColor(0xD0, 0xD0, 0xD0)

            # Trend vs previous quarter - format with percentage if available
            prev_quarter_short = f"Q{self.quarter - 1 if self.quarter > 1 else 4}"  # Short format like "Q4"
            trend_para = cell.add_paragraph()
            trend_para.paragraph_format.space_before = Pt(0)
            trend_para.paragraph_format.space_after = Pt(0)
            trend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse trend to extract percentage or determine display
            trend_text = str(trend).strip()
            trend_display = None
            trend_color = None
            
            # Prefer calculated percentage if available
            if calculated_pct:
                trend_display = f"vs {prev_quarter_short}: {calculated_pct}"
                if "+" in calculated_pct and risk_level_color:
                    trend_color = risk_level_color
                else:
                    trend_color = RGBColor(0x33, 0x33, 0x33)  # Brighter/darker text for light background
            # Check if trend contains percentage
            elif "%" in trend_text:
                # Already has percentage format
                trend_display = f"vs {prev_quarter_short}: {trend_text}"
                # Use risk level color if it's an increase
                if "+" in trend_text and risk_level_color:
                    trend_color = risk_level_color
                else:
                    trend_color = RGBColor(0x33, 0x33, 0x33)  # Brighter text for unchanged
            elif trend_text.upper() in ["UNCHANGED", "UNCHANGED"]:
                trend_display = f"vs {prev_quarter_short}: Unchanged"
                trend_color = RGBColor(0x33, 0x33, 0x33)  # Brighter text for "Unchanged"
            elif trend_text in ["↑", "INCREASED", "INCREASE"]:
                # Show as "Unchanged" if no percentage available
                trend_display = f"vs {prev_quarter_short}: Unchanged"
                trend_color = RGBColor(0x33, 0x33, 0x33)
            elif trend_text in ["↓", "DECREASED", "DECREASE"]:
                trend_display = f"vs {prev_quarter_short}: Unchanged"
                trend_color = RGBColor(0x33, 0x33, 0x33)  # Brighter text
            else:
                # Try to extract percentage from trend string
                percentage_match = re.search(r'([+-]?\d+(?:\.\d+)?%)', trend_text)
                if percentage_match:
                    percentage = percentage_match.group(1)
                    trend_display = f"vs {prev_quarter_short}: {percentage}"
                    if "+" in percentage and risk_level_color:
                        trend_color = risk_level_color
                    else:
                        trend_color = RGBColor(0x33, 0x33, 0x33)
                else:
                    # Default: show as "Unchanged"
                    trend_display = f"vs {prev_quarter_short}: Unchanged"
                    trend_color = RGBColor(0x33, 0x33, 0x33)  # Brighter text
            
            # Create trend run
            trend_run = trend_para.add_run(trend_display)
            trend_run.font.name = "Arial"
            trend_run.font.size = FontSizes.FOOTNOTE
            if trend_color:
                trend_run.font.color.rgb = trend_color

        self.doc.add_paragraph()

    def _add_breach_landscape(self, analysis_result: Dict[str, Any]) -> None:
        """Add industry breach landscape section."""
        logger.info("Adding Industry Breach Landscape section")

        # Industry Breach Landscape heading - Heading 1
        breach_heading = self.doc.add_heading("Industry Breach Landscape", level=1)
        for run in breach_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        # Add space before and after heading
        breach_heading.paragraph_format.space_before = Pt(12)
        breach_heading.paragraph_format.space_after = Pt(6)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        sub_run = subtitle.add_run(
            f"Publicly disclosed security incidents affecting life sciences, pharmaceutical, "
            f"biotechnology, healthcare, and advanced manufacturing organizations during "
            f"Q{self.quarter} {self._get_year()}."
        )
        sub_run.font.name = "Arial"
        sub_run.font.size = FontSizes.SUBTITLE
        sub_run.font.italic = True
        sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

        breach_data = analysis_result.get("breach_landscape", {})
        prev_quarter = self._get_previous_quarter()
        
        # Helper function to calculate percentage change
        def calculate_pct(current, previous):
            """Calculate percentage change."""
            if not previous or previous == 0 or previous == "N/A":
                return None
            if not current or current == "N/A":
                return None
            try:
                curr_val = float(str(current).replace("$", "").replace("M", "").replace(",", ""))
                prev_val = float(str(previous).replace("$", "").replace("M", "").replace(",", ""))
                if prev_val == 0:
                    return None
                change = ((curr_val - prev_val) / prev_val) * 100
                return f"+{int(change)}%" if change > 0 else f"{int(change)}%"
            except (ValueError, TypeError):
                return None

        # Metric cards with percentage calculations
        metrics = [
            (
                str(breach_data.get("total_incidents", 0)),
                "Total Incidents",
                breach_data.get('prev_total_incidents', 'N/A'),
                calculate_pct(breach_data.get("total_incidents", 0), breach_data.get('prev_total_incidents'))
            ),
            (
                f"${breach_data.get('total_impact_millions', 0)}M",
                "Est. Total Impact",
                f"${breach_data.get('prev_total_impact', 'N/A')}M",
                calculate_pct(breach_data.get('total_impact_millions', 0), breach_data.get('prev_total_impact'))
            ),
            (
                str(breach_data.get("ransomware_count", 0)),
                "Ransomware",
                breach_data.get('prev_ransomware', 'N/A'),
                calculate_pct(breach_data.get("ransomware_count", 0), breach_data.get('prev_ransomware'))
            ),
            (
                f"{breach_data.get('records_exposed_millions', 0)}M",
                "Records Exposed",
                f"{breach_data.get('prev_records', 'N/A')}M",
                calculate_pct(breach_data.get('records_exposed_millions', 0), breach_data.get('prev_records'))
            ),
        ]

        self._create_breach_metric_cards(metrics, prev_quarter)
        self.doc.add_paragraph()

        # Incidents by Type heading - Heading 2, font size 12, Arial, black color
        incidents_heading = self.doc.add_heading("Incidents by Type", level=2)
        for run in incidents_heading.runs:
            run.font.name = "Arial"  # Arial font
            run.font.size = Pt(12)  # Font size 12pt
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color (default)
        # Add space after heading
        incidents_heading.paragraph_format.space_after = Pt(6)

        incidents = analysis_result.get("incidents_by_type", [])
        if incidents:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = None  # Remove table style to avoid overrides
            
            # Set column widths: narrow for Q1 and Q4 (just fit headers), expand Notable Example
            # Column widths in inches
            # Incident Type: 1.5 inches, Q1 2026: 0.5 inches (narrow to fit header), Q4 2025: 0.5 inches (narrow to fit header), Notable Example: auto (remaining)
            columns = table.columns
            columns[0].width = Inches(1.5)  # Incident Type
            columns[1].width = Inches(0.5)    # Q1 2026 - narrow to fit header text
            columns[2].width = Inches(0.5)    # Q4 2025 - narrow to fit header text
            columns[3].width = Inches(4.5)    # Notable Example - takes remaining space

            # Column headers: "Incident Type", "Q1 2026", "Q4 2025", "Notable Example"
            headers = ["Incident Type", "Q1 2026", "Q4 2025", "Notable Example"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].paragraphs[0].clear()
                header_run = header_cells[i].paragraphs[0].add_run(header)
                header_run.font.name = "Arial"
                header_run.font.bold = True
                header_run.font.size = Pt(11.43)  # 114300 EMU = 11.43pt (matches example)
                header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
                # Center-align Q1 2026 and Q4 2025 headers, center others
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Center-align vertically for Q1 and Q4 header cells
                if i == 1 or i == 2:  # Q1 2026 or Q4 2025 columns
                    tc_pr = header_cells[i]._element.get_or_add_tcPr()
                    v_align = tc_pr.find(qn("w:vAlign"))
                    if v_align is None:
                        v_align = OxmlElement("w:vAlign")
                        tc_pr.append(v_align)
                    v_align.set(qn("w:val"), "center")
                # Orange header background (#E65100) with val=clear
                self._set_cell_shading(header_cells[i], "E65100")  # Orange background
                # Add borders to header cells
                self._set_cell_borders(header_cells[i], color_hex="C0C0C0", size="1")

            for incident in incidents:
                row = table.add_row()
                cells = row.cells
                
                # Column 1: Incident Type - Arial 11.43pt, default black, left-aligned, no background
                cells[0].paragraphs[0].clear()
                type_run = cells[0].paragraphs[0].add_run(incident.get("type", ""))
                type_run.font.name = "Arial"
                type_run.font.size = Pt(11.43)  # 114300 EMU = 11.43pt
                type_run.font.bold = False
                type_run.font.italic = False
                # No color set (default black)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Column 2: Q1 2026 - Arial 11.43pt, bold, default black, center-aligned (horizontally and vertically)
                cells[1].paragraphs[0].clear()
                current_run = cells[1].paragraphs[0].add_run(str(incident.get("current_count", 0)))
                current_run.font.name = "Arial"
                current_run.font.size = Pt(11.43)  # 114300 EMU = 11.43pt
                current_run.font.bold = True  # Bold numbers in Q1 2026 column
                current_run.font.italic = False
                # No color set (default black)
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned horizontally
                # Center-align vertically
                tc_pr = cells[1]._element.get_or_add_tcPr()
                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is None:
                    v_align = OxmlElement("w:vAlign")
                    tc_pr.append(v_align)
                v_align.set(qn("w:val"), "center")
                
                # Column 3: Q4 2025 - Arial 11.43pt, not bold, gray (#666666), center-aligned (horizontally and vertically)
                cells[2].paragraphs[0].clear()
                prev_run = cells[2].paragraphs[0].add_run(str(incident.get("prev_count", 0)))
                prev_run.font.name = "Arial"
                prev_run.font.size = Pt(11.43)  # 114300 EMU = 11.43pt
                prev_run.font.bold = False
                prev_run.font.italic = False
                prev_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)  # Gray #666666
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned horizontally
                # Center-align vertically
                tc_pr = cells[2]._element.get_or_add_tcPr()
                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is None:
                    v_align = OxmlElement("w:vAlign")
                    tc_pr.append(v_align)
                v_align.set(qn("w:val"), "center")
                
                # Column 4: Notable Example - Arial 10.16pt, italic, dark gray (#555555), left-aligned
                cells[3].paragraphs[0].clear()
                example_run = cells[3].paragraphs[0].add_run(incident.get("notable_example", ""))
                example_run.font.name = "Arial"
                example_run.font.size = Pt(10.16)  # 101600 EMU = 10.16pt
                example_run.font.italic = True  # Italic text
                example_run.font.bold = False
                example_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)  # Dark gray #555555
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Add borders to data row cells (no background shading - uses default white)
                for cell in cells:
                    # No background shading (None = default white)
                    # Add borders
                    self._set_cell_borders(cell, color_hex="C0C0C0", size="1")  # Thin light gray borders

        self.doc.add_paragraph()

        # Common factors
        common_factors = analysis_result.get("common_factors", "")
        if common_factors:
            factors_para = self.doc.add_paragraph()
            # Text before colon: blue and bold
            factors_label = factors_para.add_run("Common factors across incidents: ")
            factors_label.font.name = "Arial"
            factors_label.font.size = FontSizes.BODY_SMALL
            factors_label.font.bold = True
            factors_label.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
            # Text after colon: normal (not bold)
            factors_text = factors_para.add_run(common_factors)
            factors_text.font.name = "Arial"
            factors_text.font.size = FontSizes.BODY_SMALL
            factors_text.font.bold = False

        # Don't add paragraph here - divider will be added right after

    def _create_metric_cards(self, metrics: List[tuple]) -> None:
        """Create metric cards for breach landscape."""
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (number, title, subtitle) in enumerate(metrics):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].clear()

            # Number
            num_para = cell.paragraphs[0]
            num_run = num_para.add_run(number)
            num_run.font.name = "Arial"
            num_run.font.size = Pt(24)
            num_run.font.bold = True
            num_run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange for emphasis
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Title
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = "Arial"
            title_run.font.size = FontSizes.BODY_SMALL
            title_run.font.bold = True
            title_run.font.color.rgb = BrandColors.GRAY_DARK  # Dark text
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle (previous quarter comparison) - check for percentage increases
            sub_para = cell.add_paragraph()
            # Extract percentage from subtitle if present
            if "+" in subtitle and "%" in subtitle:
                # Has percentage increase - color it green
                sub_run = sub_para.add_run(subtitle)
                sub_run.font.name = "Arial"
                sub_run.font.size = FontSizes.FOOTNOTE
                sub_run.font.color.rgb = BrandColors.GREEN_LOW  # Green for increases
            else:
                sub_run = sub_para.add_run(subtitle)
                sub_run.font.name = "Arial"
                sub_run.font.size = FontSizes.FOOTNOTE
                sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM  # Gray for other text
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_breach_metric_cards(self, metrics: List[tuple], prev_quarter: str) -> None:
        """Create metric cards for breach landscape with dark olive background."""
        # Create horizontal table (1 row, 4 columns) - same structure as risk boxes
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Remove any table style that might override cell shading
        table.style = None
        
        # Set table width to 100%
        tbl = table._element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        
        # Remove any table style reference that might override
        tbl_style = tbl_pr.find(qn("w:tblStyle"))
        if tbl_style is not None:
            tbl_pr.remove(tbl_style)
        
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "5000")  # 5000 = 100% in Word's 50ths of a percent
        tbl_w.set(qn("w:type"), "pct")

        for i, (number, title, prev_value, percentage) in enumerate(metrics):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].clear()
            
            # Set cell vertical alignment to center
            tc_pr = cell._element.get_or_add_tcPr()
            v_align = tc_pr.find(qn("w:vAlign"))
            if v_align is None:
                v_align = OxmlElement("w:vAlign")
                tc_pr.append(v_align)
            v_align.set(qn("w:val"), "center")
            
            # Set cell margins for lean/compact spacing
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            
            for margin_type in ["top", "bottom", "left", "right"]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), "72")  # 72 twips = 0.05 inches
                margin.set(qn("w:type"), "dxa")

            # Set cell background to dark olive (#372E00) with val=clear
            self._set_cell_shading(cell, "372E00")  # Dark olive background

            # Set cell borders (light gray, thinner - 1/2 pt)
            self._set_cell_borders(cell, color_hex="C0C0C0", size="1")

            # Main value (font size 18, bold, white/light gray)
            num_para = cell.paragraphs[0]
            num_para.paragraph_format.space_before = Pt(0)
            num_para.paragraph_format.space_after = Pt(0)
            num_run = num_para.add_run(number)
            num_run.font.name = "Arial"
            num_run.font.size = Pt(18)  # Font size 18 as requested
            num_run.font.bold = True
            num_run.font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)  # Light gray/white for dark background
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Title/label (font size 8.5)
            title_para = cell.add_paragraph()
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run(title)
            title_run.font.name = "Arial"
            title_run.font.size = Pt(8.5)  # Font size 8.5 as requested
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)  # Light gray/white for dark background
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Comparison line (previous quarter with percentage)
            comp_para = cell.add_paragraph()
            comp_para.paragraph_format.space_before = Pt(0)
            comp_para.paragraph_format.space_after = Pt(0)
            comp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format: "Q4 2025: 36 +31%" where percentage is orange
            if percentage:
                # Add base text
                base_text = f"{prev_quarter}: {prev_value} "
                comp_run = comp_para.add_run(base_text)
                comp_run.font.name = "Arial"
                comp_run.font.size = Pt(7)  # Font size 7 for base text
                comp_run.font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)  # Light gray/white
                
                # Add percentage in orange (font size 7)
                pct_run = comp_para.add_run(percentage)
                pct_run.font.name = "Arial"
                pct_run.font.size = Pt(7)  # Font size 7 as requested
                pct_run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange for percentage
            else:
                comp_text = f"{prev_quarter}: {prev_value}"
                comp_run = comp_para.add_run(comp_text)
                comp_run.font.name = "Arial"
                comp_run.font.size = Pt(7)  # Font size 7
                comp_run.font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)  # Light gray/white

    def _add_geopolitical_landscape(self, analysis_result: Dict[str, Any]) -> None:
        """Add geopolitical threat landscape section."""
        logger.info("Adding Geopolitical Threat Landscape section")

        geo_heading = self.doc.add_heading("Geopolitical Threat Landscape", level=1)
        for run in geo_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange color
        # Add space before and after heading
        geo_heading.paragraph_format.space_before = Pt(12)
        geo_heading.paragraph_format.space_after = Pt(6)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        sub_run = subtitle.add_run(
            f"Nation-state cyber activity with implications for the life sciences and "
            f"genomics sector during Q{self.quarter} {self._get_year()}."
        )
        sub_run.font.name = "Arial"
        sub_run.font.size = FontSizes.SUBTITLE
        sub_run.font.italic = True
        sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

        geopolitical = analysis_result.get("geopolitical_threats", {})

        # Default countries to cover if not provided
        countries = ["china", "russia", "north_korea", "iran"]

        for country in countries:
            country_data = geopolitical.get(country, {})
            if country_data or country in ["china", "russia", "north_korea"]:
                self._add_country_section(country, country_data)

    def _add_country_section(self, country: str, data: Dict[str, Any]) -> None:
        """Add a country-specific threat section."""
        # Country heading - styled in blue/teal color
        country_names = {
            "china": "China",
            "russia": "Russia",
            "north_korea": "North Korea",
            "iran": "Iran"
        }
        country_heading = self.doc.add_heading(country_names.get(country, country.title()), level=2)
        # Style country headings - Heading 2, font size 12, black, bold
        for run in country_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)  # Font size 12pt
            run.font.bold = True  # Bold
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        # Add space after heading
        country_heading.paragraph_format.space_after = Pt(6)

        # Strategic Context
        strategic_context = data.get("strategic_context", self._get_default_strategic_context(country))
        context_para = self.doc.add_paragraph()
        context_label = context_para.add_run("Strategic Context: ")
        context_label.font.name = "Arial"
        context_label.font.size = Pt(10.5)  # 10.5pt font
        context_label.font.bold = True
        context_label.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        context_text = context_para.add_run(strategic_context)
        context_text.font.name = "Arial"
        context_text.font.size = FontSizes.BODY
        context_text.font.bold = False

        # Very small space between subsections
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(3)  # Very small spacing

        # Quarter Activity
        activity = data.get("activity", self._get_default_activity(country))
        activity_para = self.doc.add_paragraph()
        activity_label = activity_para.add_run(f"Q{self.quarter} Activity: ")
        activity_label.font.name = "Arial"
        activity_label.font.size = Pt(10.5)  # 10.5pt font
        activity_label.font.bold = True
        activity_label.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        activity_text = activity_para.add_run(activity)
        activity_text.font.name = "Arial"
        activity_text.font.size = FontSizes.BODY
        activity_text.font.bold = False

        # Very small space between subsections
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(3)  # Very small spacing

        # Business Implications
        implications = data.get("implications", self._get_default_implications(country))
        impl_para = self.doc.add_paragraph()
        impl_label = impl_para.add_run("Business Implications: ")
        impl_label.font.name = "Arial"
        impl_label.font.size = Pt(10.5)  # 10.5pt font
        impl_label.font.bold = True
        impl_label.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        impl_text = impl_para.add_run(implications)
        impl_text.font.name = "Arial"
        impl_text.font.size = FontSizes.BODY
        impl_text.font.bold = False

        # Add divider after each country section
        self._add_section_divider()

    def _get_default_strategic_context(self, country: str) -> str:
        """Get default strategic context for a country."""
        defaults = {
            "china": "China's national plans designate biotechnology as a strategic priority, with emphasis on genomics, precision medicine, and biomanufacturing.",
            "russia": "Russian state cyber interests in life sciences remain opportunistic, focusing on healthcare disruption capabilities. Russian-speaking criminal groups pose significant ransomware risk.",
            "north_korea": "North Korean cyber operations serve dual purposes: revenue generation to circumvent sanctions and acquisition of medical/pharmaceutical research for domestic programs.",
            "iran": "Iranian cyber operations target healthcare and pharmaceutical sectors primarily for intelligence gathering and potential disruptive capabilities."
        }
        return defaults.get(country, "Strategic context requires analysis.")

    def _get_default_activity(self, country: str) -> str:
        """Get default activity description for a country."""
        return f"Activity analysis for {country.title()} threat actors pending data collection from Intel471 and CrowdStrike sources."

    def _get_default_implications(self, country: str) -> str:
        """Get default business implications for a country."""
        defaults = {
            "china": "Theft of proprietary research, sequencing technology designs, or manufacturing processes could erode competitive advantage.",
            "russia": "Ransomware incidents in life sciences and manufacturing can result in significant recovery costs and operational disruption.",
            "north_korea": "Credential compromise of research or executive personnel could provide access to sensitive environments and IP repositories.",
            "iran": "Potential for both espionage and disruptive operations targeting pharmaceutical supply chains."
        }
        return defaults.get(country, "Business implications require assessment.")

    def _add_looking_ahead(self, analysis_result: Dict[str, Any]) -> None:
        """Add looking ahead section for next quarter."""
        logger.info("Adding Looking Ahead section")

        next_quarter = self.quarter + 1 if self.quarter < 4 else 1
        next_year = self._get_year() if self.quarter < 4 else self._get_year() + 1

        looking_ahead_heading = self.doc.add_heading(f"Looking Ahead: Q{next_quarter} {next_year}", level=1)
        # Style "Looking Ahead" heading - Heading 1, font size 14, orange
        for run in looking_ahead_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange color
        # Add space before and after heading
        looking_ahead_heading.paragraph_format.space_before = Pt(12)
        looking_ahead_heading.paragraph_format.space_after = Pt(6)

        looking_ahead = analysis_result.get("looking_ahead", {})

        # Threat Outlook
        outlook = looking_ahead.get("threat_outlook",
            "We anticipate continued pressure from state-sponsored espionage campaigns as genomics "
            "research and precision manufacturing technology becomes increasingly valuable.")
        outlook_para = self.doc.add_paragraph()
        outlook_bold = outlook_para.add_run("Threat Outlook: ")
        outlook_bold.font.name = "Arial"
        outlook_bold.font.size = Pt(10.5)  # 10.5pt font
        outlook_bold.font.bold = True
        outlook_bold.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        outlook_text = outlook_para.add_run(outlook)
        outlook_text.font.name = "Arial"
        outlook_text.font.size = FontSizes.BODY
        outlook_text.font.bold = False

        # Very small space between subsections
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(3)  # Very small spacing

        # Planned Initiatives
        initiatives = looking_ahead.get("planned_initiatives",
            "Continued monitoring of threat landscape and enhancement of detection capabilities.")
        init_para = self.doc.add_paragraph()
        init_bold = init_para.add_run("Planned Initiatives: ")
        init_bold.font.name = "Arial"
        init_bold.font.size = Pt(10.5)  # 10.5pt font
        init_bold.font.bold = True
        init_bold.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        init_text = init_para.add_run(initiatives)
        init_text.font.name = "Arial"
        init_text.font.size = FontSizes.BODY
        init_text.font.bold = False

        # Very small space between subsections
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(0)
        spacing_para.paragraph_format.space_after = Pt(3)  # Very small spacing

        # Watch Items
        watch_items = looking_ahead.get("watch_items",
            "Potential escalation in state-sponsored activity around major industry events and product announcements.")
        watch_para = self.doc.add_paragraph()
        watch_bold = watch_para.add_run("Watch Items: ")
        watch_bold.font.name = "Arial"
        watch_bold.font.size = Pt(10.5)  # 10.5pt font
        watch_bold.font.bold = True
        watch_bold.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
        watch_text = watch_para.add_run(watch_items)
        watch_text.font.name = "Arial"
        watch_text.font.size = FontSizes.BODY
        watch_text.font.bold = False

        # Don't add paragraph here - divider will be added right after

    def _add_leadership_recommendations(self, analysis_result: Dict[str, Any]) -> None:
        """Add recommendations for leadership section."""
        logger.info("Adding Recommendations for Leadership section")

        recommendations_heading = self.doc.add_heading("Recommendations for Leadership", level=1)
        # Style "Recommendations" heading - Heading 1, font size 14, orange
        for run in recommendations_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange color
        # Add space before and after heading
        recommendations_heading.paragraph_format.space_before = Pt(12)
        recommendations_heading.paragraph_format.space_after = Pt(6)

        recommendations = analysis_result.get("recommendations", [])

        if not recommendations:
            recommendations = [
                ("Executive Awareness", "Consider targeted security awareness for executives and key research personnel given sustained social engineering campaigns via professional networks."),
                ("Vendor Risk Review", "Evaluate security posture of critical software and laboratory equipment vendors given supply chain compromise activity."),
                ("Manufacturing Environment Security", "Review network segmentation between IT and OT/manufacturing systems. Ensure incident response plans address manufacturing disruption scenarios."),
                ("Incident Response Readiness", "Confirm ransomware response plans address regulatory disclosure timelines, notification requirements, and manufacturing continuity scenarios."),
                ("Board Reporting", "Peer incidents and regulatory enforcement may prompt board inquiries. CTI team available to support sector threat context preparation."),
            ]

        for rec in recommendations:
            if isinstance(rec, tuple):
                title, description = rec
                para = self.doc.add_paragraph(style="List Bullet")
                title_run = para.add_run(f"{title}: ")
                title_run.font.name = "Arial"
                title_run.font.bold = True
                title_run.font.size = Pt(10.5)  # 10.5pt font
                title_run.font.color.rgb = RGBColor(0x00, 0x8A, 0xC9)  # #008AC9
                desc_run = para.add_run(description)
                desc_run.font.name = "Arial"
                desc_run.font.size = FontSizes.BODY
                # Add spacing after each bullet point
                para.paragraph_format.space_after = Pt(6)
            else:
                para = self.doc.add_paragraph(rec, style="List Bullet")
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = FontSizes.BODY
                    run.font.bold = True
                # Add spacing after each bullet point
                para.paragraph_format.space_after = Pt(6)

        self.doc.add_paragraph()

    def _add_footer(self) -> None:
        """Add footer with contact info, sources, TLP classification, and page number."""
        logger.info("Adding footer")

        # Contact info
        contact = self.doc.add_paragraph()
        contact_run = contact.add_run("Prepared by: Cyber Threat Intelligence | cti@illumina.com")
        contact_run.font.name = "Arial"
        contact_run.font.size = FontSizes.BODY_SMALL
        contact_run.font.bold = True
        contact_run.font.color.rgb = BrandColors.GRAY_DARK  # Dark text
        contact_run.font.underline = False  # No underline
        contact.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left align

        self.doc.add_paragraph()

        # Data sources
        sources = self.doc.add_paragraph()
        sources_run = sources.add_run(
            "Sources: CrowdStrike Falcon Intelligence, Intel471 Titan, HHS Breach Portal, "
            "FBI IC3, SEC filings, state attorney general notifications, FDA guidance publications, "
            "and open source intelligence. Breach counts based on public disclosures and may not "
            "reflect total incidents."
        )
        sources_run.font.name = "Arial"
        sources_run.font.size = FontSizes.FOOTNOTE
        sources_run.font.bold = True
        sources_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White text
        sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Underline "open source intelligence" if present
        for run in sources.runs:
            if "open source intelligence" in run.text.lower():
                run.font.underline = True

        # Add TLP classification and page number to document footer
        section = self.doc.sections[0]
        footer = section.footer
        
        # Clear any existing footer content
        if footer.paragraphs:
            footer.paragraphs[0].clear()
        else:
            footer.add_paragraph()
        
        # Create footer paragraph
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "TLP:" in light gray, italic
        tlp_label = footer_para.add_run("TLP: ")
        tlp_label.font.name = "Arial"
        tlp_label.font.size = FontSizes.FOOTNOTE
        tlp_label.font.italic = True
        tlp_label.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray
        
        # "AMBER+" in orange, italic
        amber = footer_para.add_run("AMBER+")
        amber.font.name = "Arial"
        amber.font.size = FontSizes.FOOTNOTE
        amber.font.italic = True
        amber.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange
        
        # "STRICT" in orange, italic, no underline
        strict = footer_para.add_run("STRICT")
        strict.font.name = "Arial"
        strict.font.size = FontSizes.FOOTNOTE
        strict.font.italic = True
        strict.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange
        strict.font.underline = False  # No underline
        
        # Pipe separator in light gray, italic
        pipe = footer_para.add_run(" | ")
        pipe.font.name = "Arial"
        pipe.font.size = FontSizes.FOOTNOTE
        pipe.font.italic = True
        pipe.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray
        
        # Page number in light gray, italic
        # Use simple text for now - user can manually add page number field in Word if needed
        # This avoids XML corruption issues
        page_text = footer_para.add_run("Page 1")
        page_text.font.name = "Arial"
        page_text.font.size = FontSizes.FOOTNOTE
        page_text.font.italic = True
        page_text.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray
