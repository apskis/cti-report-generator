"""
Weekly CTI Report Generator.

Generates weekly threat intelligence reports matching the branded template.
"""
from datetime import datetime, timedelta
import logging
from typing import Dict, Any, List

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reports.base import BaseReportGenerator, BrandColors, FontSizes
from reports.registry import register_report_generator

logger = logging.getLogger(__name__)


@register_report_generator("weekly")
class WeeklyReportGenerator(BaseReportGenerator):
    """
    Weekly CTI Report Generator.

    Generates reports matching the CTI_Weekly_Report_Template_Example.docx format.

    Template structure:
    - Report ID header (CTI-WK-YYYY-WW)
    - Title: "Cyber Threat Intelligence Weekly Report" (orange, 18pt bold)
    - Date range subtitle
    - Executive Summary
    - This Week at a Glance (metric cards)
    - Vulnerability Exposure (CVE table with weeks tracked)
    - Sector Threat Activity (threat actor table)
    - Exploitation Indicators
    - Recommended Actions
    - Footer with contact info and data sources
    """

    @property
    def report_type(self) -> str:
        return "weekly"

    @property
    def filename_prefix(self) -> str:
        return "CTI_Weekly_Report"

    def generate(self, analysis_result: Dict[str, Any]) -> Document:
        """
        Generate the weekly report document.

        Args:
            analysis_result: Dictionary containing:
                - executive_summary: str
                - statistics: dict with counts
                - cve_analysis: list of CVE dicts
                - apt_activity: list of APT dicts
                - recommendations: list of strings
                - exploitation_indicators: list of strings (optional)

        Returns:
            Document object
        """
        try:
            logger.info("Generating Weekly CTI Report")
            self.doc = Document()

            # Calculate date range (Monday to Sunday of current week)
            self._calculate_date_range()

            # Add sections in order
            self._add_header()
            self._add_executive_summary(analysis_result)
            self._add_week_at_glance(analysis_result)
            self._add_vulnerability_exposure(analysis_result)
            self._add_sector_threat_activity(analysis_result)
            self._add_exploitation_indicators(analysis_result)
            self._add_recommended_actions(analysis_result)
            self._add_footer()

            logger.info("Weekly CTI Report generated successfully")
            return self.doc

        except Exception as e:
            logger.error(f"Error generating weekly report: {str(e)}", exc_info=True)
            raise

    def _calculate_date_range(self) -> None:
        """Calculate the week's date range (Monday to Sunday)."""
        today = self.created_at
        # Find Monday of current week
        self.week_start = today - timedelta(days=today.weekday())
        # Find Sunday of current week
        self.week_end = self.week_start + timedelta(days=6)

    def _add_header(self) -> None:
        """Add report header with ID, title, and date range."""
        # Report ID (e.g., CTI-WK-2026-04)
        week_num = self._get_week_number()
        year = self._get_year()
        report_id = f"CTI-WK-{year}-{week_num:02d}"

        id_para = self.doc.add_paragraph()
        id_run = id_para.add_run(report_id)
        id_run.font.size = FontSizes.SUBTITLE
        id_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        id_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Main title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run("Cyber Threat Intelligence Weekly Report")
        title_run.font.size = FontSizes.TITLE
        title_run.font.bold = True
        title_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Date range subtitle
        date_range = f"Week {week_num} | {self.week_start.strftime('%B %d')} to {self.week_end.strftime('%d, %Y')}"
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(date_range)
        date_run.font.size = FontSizes.BODY_SMALL
        date_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add spacing
        self.doc.add_paragraph()

    def _add_executive_summary(self, analysis_result: Dict[str, Any]) -> None:
        """Add executive summary section."""
        logger.info("Adding Executive Summary section")

        heading = self.doc.add_heading("Executive Summary", level=1)

        summary = analysis_result.get("executive_summary", "No executive summary available.")
        para = self.doc.add_paragraph(summary)
        for run in para.runs:
            run.font.size = FontSizes.BODY

        self.doc.add_paragraph()

    def _add_week_at_glance(self, analysis_result: Dict[str, Any]) -> None:
        """Add 'This Week at a Glance' section with metric cards."""
        logger.info("Adding This Week at a Glance section")

        self.doc.add_heading("This Week at a Glance", level=1)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        sub_run = subtitle.add_run(
            f"Vulnerability and threat actor metrics from automated collection "
            f"({self.week_start.strftime('%B %d')} to {self.week_end.strftime('%d, %Y')})."
        )
        sub_run.font.size = FontSizes.SUBTITLE
        sub_run.font.italic = True
        sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

        stats = analysis_result.get("statistics", {})

        # First row of metric cards
        metrics_row1 = [
            (str(stats.get("new_this_week", stats.get("total_cves", 0))),
             "New This Week",
             "First appearance in Rapid7 scans"),
            (str(stats.get("persistent_count", stats.get("critical_count", 0))),
             "Persistent (3+ Wks)",
             "Unresolved from prior reports"),
            (str(stats.get("resolved_count", 0)),
             "Resolved",
             "Remediated since last week"),
        ]
        self._create_metric_cards(metrics_row1)

        self.doc.add_paragraph()

        # Second row of metric cards
        metrics_row2 = [
            (str(stats.get("total_exposed", stats.get("total_cves", 0))),
             "Total Exposed",
             "CVEs on assets"),
            (str(stats.get("exploited_count", stats.get("actively_exploited", 0))),
             "Actively Exploited",
             "Confirmed threat actor use"),
            (str(stats.get("apt_groups", 0)),
             "Actor Groups",
             "Targeting biotech sector"),
        ]
        self._create_metric_cards(metrics_row2)

        self.doc.add_paragraph()

    def _create_metric_cards(self, metrics: List[tuple]) -> None:
        """Create a row of metric cards."""
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (number, title, subtitle) in enumerate(metrics):
            cell = table.rows[0].cells[i]

            # Clear and build content
            cell.paragraphs[0].clear()

            # Number (large, bold, orange)
            num_para = cell.paragraphs[0]
            num_run = num_para.add_run(number)
            num_run.font.size = Pt(28)
            num_run.font.bold = True
            num_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Title (bold)
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = FontSizes.BODY_SMALL
            title_run.font.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle (small, gray, italic)
            sub_para = cell.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = FontSizes.FOOTNOTE
            sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            sub_run.font.italic = True
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_vulnerability_exposure(self, analysis_result: Dict[str, Any]) -> None:
        """Add vulnerability exposure section with CVE table."""
        logger.info("Adding Vulnerability Exposure section")

        self.doc.add_heading("Vulnerability Exposure", level=1)

        # Source note
        note = self.doc.add_paragraph()
        note_run = note.add_run(
            "Wks = consecutive weeks detected. Items at 3+ weeks highlighted. "
            "Source: Rapid7 InsightVM scans."
        )
        note_run.font.size = FontSizes.FOOTNOTE
        note_run.font.italic = True
        note_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

        cve_analysis = analysis_result.get("cve_analysis", [])

        if cve_analysis:
            # Create table with headers
            table = self.doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"

            # Header row
            headers = ["CVE ID", "Affected Product", "Exposure", "Exploited By", "Risk", "Wks"]
            header_cells = table.rows[0].cells

            for i, header in enumerate(headers):
                header_cells[i].text = header
                para = header_cells[i].paragraphs[0]
                para.runs[0].font.bold = True
                para.runs[0].font.size = FontSizes.SUBTITLE
                self._set_cell_shading(header_cells[i], "E0E0E0")  # Light gray header

            # Add CVE rows
            for cve in cve_analysis:
                row = table.add_row()
                cells = row.cells

                cells[0].text = cve.get("cve_id", "N/A")
                cells[1].text = cve.get("affected_product", cve.get("product", "N/A"))
                cells[2].text = cve.get("exposure", cve.get("description", "N/A")[:50])
                cells[3].text = cve.get("exploited_by", "None known")
                cells[4].text = cve.get("risk", cve.get("severity", "N/A"))
                cells[5].text = str(cve.get("weeks_detected", cve.get("wks", 1)))

                # Apply risk coloring
                self._apply_risk_color(cells[4], cells[4].text)

                # Highlight rows with 3+ weeks
                weeks = cve.get("weeks_detected", cve.get("wks", 1))
                if isinstance(weeks, str):
                    try:
                        weeks = int(weeks)
                    except ValueError:
                        weeks = 1

                if weeks >= 3:
                    for cell in cells:
                        self._set_cell_shading(cell, "FFF3CD")  # Light yellow

                # Set font size for all cells
                for cell in cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = FontSizes.SUBTITLE
        else:
            self.doc.add_paragraph("No vulnerability data available.")

        self.doc.add_paragraph()

    def _add_sector_threat_activity(self, analysis_result: Dict[str, Any]) -> None:
        """Add sector threat activity section with threat actor table."""
        logger.info("Adding Sector Threat Activity section")

        self.doc.add_heading("Sector Threat Activity", level=1)

        # Intro text
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run(
            "The following threat actors have been observed targeting organizations "
            "in the biotech, pharmaceutical, healthcare, manufacturing, and related sectors this week."
        )
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = BrandColors.GRAY_DARK

        self.doc.add_paragraph()

        apt_activity = analysis_result.get("apt_activity", [])

        if apt_activity:
            # Create threat actor table
            table = self.doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            # Header row
            headers = ["Origin / Motivation", "Activity Observed", "What to Monitor"]
            header_cells = table.rows[0].cells

            for i, header in enumerate(headers):
                header_cells[i].text = header
                para = header_cells[i].paragraphs[0]
                para.runs[0].font.bold = True
                para.runs[0].font.size = FontSizes.SUBTITLE
                self._set_cell_shading(header_cells[i], "E0E0E0")

            # Add APT rows
            for apt in apt_activity:
                row = table.add_row()
                cells = row.cells

                # Origin / Motivation
                actor = apt.get("actor", apt.get("name", "Unknown"))
                country = apt.get("country", apt.get("origin", "Unknown"))
                motivation = apt.get("motivation", "Unknown")
                cells[0].text = f"{actor}\n({country})\n{motivation}"

                # Activity Observed
                activity = apt.get("activity", apt.get("description", ""))
                ttps = apt.get("ttps", [])
                if ttps and isinstance(ttps, list):
                    activity += f"\nTTPs: {', '.join(ttps[:3])}"
                cells[1].text = activity

                # What to Monitor
                monitoring = apt.get("what_to_monitor", apt.get("indicators", ""))
                if isinstance(monitoring, list):
                    monitoring = "; ".join(monitoring[:3])
                cells[2].text = monitoring

                # Set font size
                for cell in cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = FontSizes.SUBTITLE
        else:
            self.doc.add_paragraph("No threat actor activity data available.")

        self.doc.add_paragraph()

    def _add_exploitation_indicators(self, analysis_result: Dict[str, Any]) -> None:
        """Add exploitation indicators section."""
        logger.info("Adding Exploitation Indicators section")

        self.doc.add_heading("Exploitation Indicators for This Week's CVEs", level=2)

        indicators = analysis_result.get("exploitation_indicators", [])

        # If no explicit indicators, generate from CVE data
        if not indicators:
            cve_analysis = analysis_result.get("cve_analysis", [])
            for cve in cve_analysis[:5]:  # Limit to top 5
                cve_id = cve.get("cve_id", "")
                product = cve.get("affected_product", cve.get("product", "Unknown"))
                description = cve.get("exploitation_indicator", cve.get("description", ""))
                if description:
                    indicators.append(f"{cve_id} ({product}): {description[:150]}")

        if indicators:
            for indicator in indicators:
                para = self.doc.add_paragraph(indicator, style="List Bullet")
                for run in para.runs:
                    run.font.size = FontSizes.BODY_SMALL
                    run.font.bold = True
        else:
            self.doc.add_paragraph("No exploitation indicators available for this week's CVEs.")

        self.doc.add_paragraph()

    def _add_recommended_actions(self, analysis_result: Dict[str, Any]) -> None:
        """Add recommended actions section."""
        logger.info("Adding Recommended Actions section")

        self.doc.add_heading("Recommended Actions", level=1)

        recommendations = analysis_result.get("recommendations", [])

        if recommendations:
            for i, rec in enumerate(recommendations):
                para = self.doc.add_paragraph(rec, style="List Bullet")
                for run in para.runs:
                    run.font.size = FontSizes.BODY_SMALL

                    # Bold important/urgent recommendations
                    lower_rec = rec.lower()
                    if any(word in lower_rec for word in ["urgent", "critical", "immediate", "persistent"]):
                        run.font.bold = True
        else:
            default_recs = [
                "Review scan results for exposed vulnerabilities; validate asset ownership and remediation timelines",
                "Brief development teams on current threat campaigns",
                "Verify endpoint protection has latest behavioral IOAs enabled",
                "Confirm SIEM is receiving logs from affected systems"
            ]
            for rec in default_recs:
                para = self.doc.add_paragraph(rec, style="List Bullet")
                for run in para.runs:
                    run.font.size = FontSizes.BODY_SMALL

        self.doc.add_paragraph()

    def _add_footer(self) -> None:
        """Add footer with contact info and data sources."""
        logger.info("Adding footer")

        # Contact info
        contact = self.doc.add_paragraph()
        contact_run = contact.add_run(
            "Questions or suspicious activity: secops@company.com | ServiceNow | cti@company.com"
        )
        contact_run.font.size = FontSizes.BODY_SMALL
        contact_run.font.bold = True
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        # Data sources
        sources = self.doc.add_paragraph()
        sources_run = sources.add_run(
            "Data Sources: Compiled via automated collection from NVD, CrowdStrike Falcon Intelligence, "
            "Intel471, Rapid7 InsightVM, and ThreatQ. Analysis performed by AI-assisted threat analysis."
        )
        sources_run.font.size = FontSizes.FOOTNOTE
        sources_run.font.bold = True
        sources_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        sources.alignment = WD_ALIGN_PARAGRAPH.CENTER
