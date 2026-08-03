"""Inject carryover-aware recommendations into the weekly report.

Replaces the default bullet-list Recommended Actions with a table that
distinguishes new recommendations from persistent (recurring) ones,
showing age and target inline.
"""

from __future__ import annotations

import logging

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from carryover.store import Action
from carryover.week import Week

logger = logging.getLogger(__name__)

_FONT = "Arial"
_BODY_SIZE = Pt(9)
_BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
_HEADER_FILL = "E65100"
_HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
_NEW_BADGE_COLOR = RGBColor(0x2E, 0x7D, 0x32)  # Green for "New"
_PERSISTENT_BADGE_COLOR = RGBColor(0xE6, 0x51, 0x00)  # Orange for persistent
_OVERDUE_COLOR = RGBColor(0xC6, 0x28, 0x28)  # Red for overdue
_BORDER_COLOR = "CCCCCC"
_BORDER_SIZE = "4"
_CAPTION_SIZE = Pt(8)
_CAPTION_COLOR = RGBColor(0x66, 0x66, 0x66)
_STATUS_COL_FILL = "F5F5F5"

_COLUMNS = ["Recommendation", "Status", "Age", "Target"]


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), _BORDER_SIZE)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER_COLOR)
        borders.append(el)
    tc_pr.append(borders)


def inject_recommendations_table(
    doc: Document,
    actions: list[Action],
    report_week: Week,
    new_action_ids: set[str] | None = None,
) -> Document:
    """Append a recommendations table into the document.

    Renders persistent (recurring) actions first, then new ones,
    with age and overdue status visible.

    Args:
        doc: The existing python-docx Document
        actions: All actions to include (from actions_for_table)
        report_week: The report week for age calculation
        new_action_ids: IDs of actions created this week (shown as "New")

    Returns:
        The modified Document
    """
    if not actions:
        return doc

    new_ids = new_action_ids or set()

    # Split into persistent and new
    persistent = [a for a in actions if a.id not in new_ids and a.status != "complete"]
    new = [a for a in actions if a.id in new_ids]
    completed = [a for a in actions if a.status == "complete"]

    # Build combined list: persistent first (sorted by age desc), then new
    ordered = persistent + new + completed

    # Add table
    table = doc.add_table(rows=1, cols=len(_COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, col_name in enumerate(_COLUMNS):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        _set_cell_borders(cell)
        _set_cell_shading(cell, _HEADER_FILL)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = _FONT
                run.font.size = _BODY_SIZE
                run.font.color.rgb = _HEADER_TEXT_COLOR
                run.font.bold = True

    # Data rows
    for action in ordered:
        row = table.add_row()
        is_new = action.id in new_ids
        is_overdue = action.is_overdue(report_week)
        age = action.age_weeks(report_week)

        # Column 0: Recommendation (title)
        cell0 = row.cells[0]
        cell0.text = ""
        para0 = cell0.paragraphs[0]
        run0 = para0.add_run(action.title)
        run0.font.name = _FONT
        run0.font.size = _BODY_SIZE
        run0.font.color.rgb = _BODY_COLOR
        _set_cell_borders(cell0)

        # Column 1: Status badge
        cell1 = row.cells[1]
        cell1.text = ""
        para1 = cell1.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if action.status == "complete":
            badge_text = "Complete"
            badge_color = _NEW_BADGE_COLOR
        elif is_new:
            badge_text = "New"
            badge_color = _NEW_BADGE_COLOR
        elif is_overdue:
            badge_text = "Overdue"
            badge_color = _OVERDUE_COLOR
        else:
            badge_text = "Persistent"
            badge_color = _PERSISTENT_BADGE_COLOR

        run1 = para1.add_run(badge_text)
        run1.font.name = _FONT
        run1.font.size = _BODY_SIZE
        run1.font.bold = True
        run1.font.color.rgb = badge_color
        _set_cell_borders(cell1)
        _set_cell_shading(cell1, _STATUS_COL_FILL)

        # Column 2: Age
        cell2 = row.cells[2]
        cell2.text = ""
        para2 = cell2.paragraphs[0]
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        age_text = f"{age} wk" if age != 1 else "1 wk"
        if is_new:
            age_text = "—"
        run2 = para2.add_run(age_text)
        run2.font.name = _FONT
        run2.font.size = _BODY_SIZE
        run2.font.color.rgb = _OVERDUE_COLOR if is_overdue else _BODY_COLOR
        run2.font.bold = is_overdue
        _set_cell_borders(cell2)

        # Column 3: Target
        cell3 = row.cells[3]
        cell3.text = ""
        para3 = cell3.paragraphs[0]
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = para3.add_run(action.target_week)
        run3.font.name = _FONT
        run3.font.size = _BODY_SIZE
        run3.font.color.rgb = _OVERDUE_COLOR if is_overdue else _BODY_COLOR
        run3.font.bold = is_overdue
        _set_cell_borders(cell3)

    # Column widths
    table.columns[0].width = Inches(4.2)
    table.columns[1].width = Inches(1.0)
    table.columns[2].width = Inches(0.7)
    table.columns[3].width = Inches(0.9)

    # Caption
    persistent_count = len(persistent)
    new_count = len(new)
    overdue_count = sum(1 for a in actions if a.is_overdue(report_week))
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    parts = []
    if persistent_count:
        parts.append(f"{persistent_count} persistent")
    if new_count:
        parts.append(f"{new_count} new")
    if overdue_count:
        parts.append(f"{overdue_count} overdue")
    caption_text = f"Table: {', '.join(parts)}."
    run = caption.add_run(caption_text)
    run.font.name = _FONT
    run.font.size = _CAPTION_SIZE
    run.font.color.rgb = _CAPTION_COLOR
    run.font.italic = True

    doc.add_paragraph()
    return doc
