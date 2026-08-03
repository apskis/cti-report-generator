"""Carryover table rendering: Word (.docx) and Markdown (standalone CLI output)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from carryover.store import Action
from carryover.week import Week

_COLUMNS = ["Action ID", "Recommendation", "Status", "Age", "Target"]

_FONT = "Arial"
_BODY_SIZE = Pt(9)
_BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
_HEADER_FILL = "E65100"
_HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
_ID_COL_FILL = "F5F5F5"
_STATUS_COL_FILL = "E3F2FD"
_BORDER_COLOR = "CCCCCC"
_BORDER_SIZE = "4"
_CAPTION_SIZE = Pt(8)
_CAPTION_COLOR = RGBColor(0x66, 0x66, 0x66)


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tc_pr.append(shd)


def _set_cell_borders(cell) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), _BORDER_SIZE)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER_COLOR)
        borders.append(el)
    tc_pr.append(borders)


def _style_cell(cell, bold: bool = False, fill: str | None = None) -> None:
    _set_cell_borders(cell)
    if fill:
        _set_cell_shading(cell, fill)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = _FONT
            run.font.size = _BODY_SIZE
            run.font.color.rgb = _BODY_COLOR
            run.font.bold = bold


def _row_data(action: Action, report_week: Week) -> list[str]:
    age = action.age_weeks(report_week)
    status_display = action.status.replace("_", " ").title()
    is_overdue = action.is_overdue(report_week)
    if is_overdue:
        status_display = "Overdue"
    return [
        action.id,
        action.title,
        status_display,
        f"{age} weeks",
        action.target_week,
    ]


def render_docx(actions: list[Action], report_week: Week, output_path: Path) -> Path:
    """Render the carryover table as a standalone .docx file."""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    table = doc.add_table(rows=1, cols=len(_COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col_name in enumerate(_COLUMNS):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        _set_cell_borders(cell)
        _set_cell_shading(cell, _HEADER_FILL)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = _FONT
                run.font.size = _BODY_SIZE
                run.font.color.rgb = _HEADER_TEXT_COLOR
                run.font.bold = True

    overdue_count = 0
    for action in actions:
        row = table.add_row()
        values = _row_data(action, report_week)
        is_overdue = action.is_overdue(report_week)
        if is_overdue:
            overdue_count += 1

        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = val

            fill = None
            if i == 0:
                fill = _ID_COL_FILL
            elif i == 2:  # Status column
                fill = _STATUS_COL_FILL

            bold = is_overdue and i == 4  # Bold target if overdue
            _style_cell(cell, bold=bold, fill=fill)

    table.columns[0].width = Inches(1.1)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(0.9)
    table.columns[3].width = Inches(0.7)
    table.columns[4].width = Inches(0.9)

    open_count = sum(1 for a in actions if a.status != "complete")
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    run = caption.add_run(f"Table: {open_count} open actions carried forward, {overdue_count} overdue.")
    run.font.name = _FONT
    run.font.size = _CAPTION_SIZE
    run.font.color.rgb = _CAPTION_COLOR
    run.font.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def render_markdown(actions: list[Action], report_week: Week) -> str:
    """Render the carryover table as a Markdown string."""
    lines: list[str] = []

    lines.append("| " + " | ".join(_COLUMNS) + " |")
    lines.append("| " + " | ".join(["---"] * len(_COLUMNS)) + " |")

    overdue_count = 0
    for action in actions:
        values = _row_data(action, report_week)
        is_overdue = action.is_overdue(report_week)
        if is_overdue:
            overdue_count += 1
            values[4] = f"**{values[4]}**"
        lines.append("| " + " | ".join(values) + " |")

    open_count = sum(1 for a in actions if a.status != "complete")
    lines.append("")
    lines.append(f"*Table: {open_count} open actions carried forward, {overdue_count} overdue.*")

    return "\n".join(lines)
