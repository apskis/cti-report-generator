"""
Base report generator class.

Provides common functionality for all report types.
"""
from abc import ABC, abstractmethod
from datetime import datetime
from io import BytesIO
import logging
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

logger = logging.getLogger(__name__)


class BrandColors:
    """Brand color constants matching the template."""

    ORANGE_PRIMARY = RGBColor(0xE6, 0x51, 0x00)  # #E65100 - Main title
    ORANGE_DESIGN = RGBColor(0xFF, 0x5C, 0x41)  # #ff5c41 - Header orange
    ILLUMINA_BLUE = RGBColor(0x00, 0x5D, 0xAA)  # #005DAA - Illumina Blue
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)  # White text on dark background
    GRAY_DARK = RGBColor(0x55, 0x55, 0x55)  # #555555 - Body text emphasis
    GRAY_MEDIUM = RGBColor(0x66, 0x66, 0x66)  # #666666 - Subtitles, notes
    GRAY_LIGHT = RGBColor(0xB0, 0xB0, 0xB0)  # Light gray for metric subtext/borders
    RED_CRITICAL = RGBColor(0xFF, 0x00, 0x00)  # Red for critical severity
    ORANGE_HIGH = RGBColor(0xFF, 0xA5, 0x00)  # Orange for high severity
    YELLOW_P1 = "FFFF00"  # Yellow highlight for P1
    GREEN_LOW = RGBColor(0x00, 0x80, 0x00)  # Green for LOW risk
    GREEN_RESOLVED = RGBColor(0x2E, 0xCC, 0x71)  # Green #2ECC71 for Resolved / positive metrics
    RED_HIGH_RISK = RGBColor(0xFF, 0x3B, 0x3B)  # Red #FF3B3B for high-risk, Actively Exploited, Actor Groups, Total Exposed
    ORANGE_TABLE_HEADER = "E65100"  # Orange for table headers (hex string for shading)
    DARK_BG_HEX = "1E1E1E"  # Dark page background (design reference)
    METRIC_BOX_DARK = "2A2A2A"  # Slightly lighter than page for metric boxes (#2A2A2A)
    BORDER_LIGHT_GRAY = "B0B0B0"  # Light gray box borders (dark theme)
    RED_LABEL = RGBColor(0xE0, 0x30, 0x30)  # Red "Generated Report" label in header
    # Print-style theme: white paper, off-white boxes — document independent of editor/app theme (e.g. dark mode)
    PAGE_WHITE = "FFFFFF"  # Page background; use so doc always renders as white paper
    BOX_OFF_WHITE = "F1F1F1"  # Metric box background (light)
    BORDER_BOX_LIGHT = "CCCCCC"  # Light gray box border
    TABLE_CELL_DARK = "2A2A2A"  # Default fill for table data cells (no bright white in dark mode)
    TABLE_CELL_DARK_RGB = (0x2A, 0x2A, 0x2A)  # 42, 42, 42 for _set_cell_shading_rgb
    TEXT_DARK = RGBColor(0x33, 0x33, 0x33)  # Dark gray body text on light
    TEXT_LIGHT = RGBColor(0xE0, 0xE0, 0xE0)  # Light gray text on dark backgrounds
    BULLET_DARK = RGBColor(0x66, 0x66, 0x66)  # Centered bullet on white (visible)
    # Vulnerability Exposure table — user-specified shading (dark mode)
    TABLE_HEADER_ORANGE = "993D22"  # Muted orange for header
    TABLE_HEADER_ORANGE_RGB = (0x99, 0x3D, 0x22)  # 153, 61, 34 for _set_cell_shading_rgb
    TABLE_BORDER_GRAY = "606060"  # Dark gray borders
    # Exploited By: threat actor / Risk High — #483135 (dark reddish-brown)
    EXPLOITED_BG_RED = "483135"
    EXPLOITED_BG_RED_RGB = (0x48, 0x31, 0x35)  # 72, 49, 53
    EXPLOITED_TEXT_RED = RGBColor(0xE8, 0x88, 0x88)
    # Exploited By: None observed / Risk Low — #273929 (dark green)
    EXPLOITED_BG_GREEN = "273929"
    EXPLOITED_BG_GREEN_RGB = (0x27, 0x39, 0x29)  # 39, 57, 41
    EXPLOITED_TEXT_GREEN = RGBColor(0x5C, 0xCC, 0x7A)
    # Exploited By: PoC available / Risk Medium / Wks 3+ — #372E00 (dark olive-brown)
    EXPLOITED_BG_AMBER = "372E00"
    EXPLOITED_BG_AMBER_RGB = (0x37, 0x2E, 0x00)  # 55, 46, 0
    EXPLOITED_TEXT_AMBER = RGBColor(0xCC, 0xA0, 0x44)
    RISK_HIGH_BG = "483135"
    RISK_HIGH_BG_RGB = (0x48, 0x31, 0x35)
    RISK_HIGH_TEXT = RGBColor(0xE8, 0x88, 0x88)
    RISK_MED_BG = "372E00"
    RISK_MED_BG_RGB = (0x37, 0x2E, 0x00)
    RISK_MED_TEXT = RGBColor(0xCC, 0xA0, 0x44)
    RISK_LOW_BG = "273929"
    RISK_LOW_BG_RGB = (0x27, 0x39, 0x29)
    RISK_LOW_TEXT = RGBColor(0x5C, 0xCC, 0x7A)
    WKS_HIGHLIGHT_BG = "372E00"
    WKS_HIGHLIGHT_BG_RGB = (0x37, 0x2E, 0x00)
    WKS_HIGHLIGHT_TEXT = RGBColor(0xCC, 0xA0, 0x44)

    # Light-mode table colors (from CTI_Weekly_Report_Template_Spec.json)
    # Table header: orange background
    TABLE_HEADER_BG = "E65100"  # Orange for header row
    # Exploited By column backgrounds
    EXPLOITED_ACTOR_BG = "FFEBEE"  # Light pink for threat actors
    EXPLOITED_NONE_BG = "E8F5E9"  # Light green for "None observed"
    EXPLOITED_POC_BG = "FFF8E1"  # Light yellow for "PoC available"
    # Risk column backgrounds
    RISK_HIGH_BG_LIGHT = "FFEBEE"  # Light pink for High/Critical
    RISK_MED_BG_LIGHT = "FFF3E0"  # Light orange for Medium
    RISK_LOW_BG_LIGHT = "E8F5E9"  # Light green for Low
    # Wks column backgrounds
    WKS_3PLUS_BG = "FFF8E1"  # Light yellow for 3+ weeks
    WKS_OVERDUE_BG = "FFEBEE"  # Light red for long overdue (4+)
    # Metric card background
    METRIC_CARD_BG = "F5F5F5"  # Light gray for metric cards
    # Sector Threat Activity table
    SECTOR_ORIGIN_BG = "F5F5F5"  # Light gray for Origin/Motivation
    SECTOR_MONITOR_BG = "E3F2FD"  # Light blue for What to Monitor


class FontSizes:
    """Font size constants matching the template."""

    TITLE = Pt(18)
    HEADING_1 = Pt(14)
    HEADING_2 = Pt(12)
    BODY = Pt(10.5)
    BODY_SMALL = Pt(10)
    SUBTITLE = Pt(9)
    FOOTNOTE = Pt(8)


class BaseReportGenerator(ABC):
    """
    Abstract base class for report generators.

    Subclasses must implement:
    - report_type: Property returning the report type name
    - generate(): Method to generate the report document
    """

    def __init__(self):
        self.doc: Document | None = None
        self.created_at = datetime.now()

    @property
    @abstractmethod
    def report_type(self) -> str:
        """Return the report type identifier (e.g., 'weekly', 'monthly')."""
        pass

    @property
    @abstractmethod
    def filename_prefix(self) -> str:
        """Return the filename prefix for this report type."""
        pass

    @abstractmethod
    def generate(self, analysis_result: Dict[str, Any]) -> Document:
        """
        Generate the report document.

        Args:
            analysis_result: Dictionary containing analysis results

        Returns:
            Document object (python-docx Document)
        """
        pass

    def get_filename(self) -> str:
        """Generate the filename for this report."""
        date_str = self.created_at.strftime("%Y-%m-%d")
        return f"{self.filename_prefix}_{date_str}.docx"

    def to_bytes(self) -> bytes:
        """Convert the document to bytes for upload."""
        if self.doc is None:
            raise ValueError("Document not generated. Call generate() first.")

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =========================================================================
    # Common styling utilities
    # =========================================================================

    def _set_cell_shading(self, cell, color_hex: str) -> None:
        """Apply solid background shading to a table cell. Uses w:val='solid' so fill is applied."""
        tc_pr = cell._element.get_or_add_tcPr()
        existing_shd = tc_pr.find(qn("w:shd"))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
        color_hex = (color_hex or "").replace("#", "").strip().upper()
        if len(color_hex) != 6:
            return
        shd.set(qn("w:fill"), color_hex)
        shd.set(qn("w:val"), "solid")
        shd.set(qn("w:color"), "auto")

    def _clear_cell_shading(self, cell) -> None:
        """Remove cell shading so the cell has no fill (inherits page/table background)."""
        tc_pr = cell._element.get_or_add_tcPr()
        existing_shd = tc_pr.find(qn("w:shd"))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)

    def _clear_table_style(self, table) -> None:
        """Remove table style from XML so Word does not override our cell shading (dark mode)."""
        try:
            tbl = table._tbl
            tbl_pr = tbl.find(qn("w:tblPr"))
            if tbl_pr is not None:
                tbl_style = tbl_pr.find(qn("w:tblStyle"))
                if tbl_style is not None:
                    tbl_pr.remove(tbl_style)
        except Exception:
            pass

    def _set_cell_shading_rgb(self, cell, r: int, g: int, b: int) -> None:
        """Apply background shading to a table cell using RGB values."""
        tc_pr = cell._element.get_or_add_tcPr()
        
        # Remove any existing shading first
        existing_shd = tc_pr.find(qn("w:shd"))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Create new shading element
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
        
        # Convert RGB to hex (Word uses hex format)
        # RGB(55, 46, 0) = #372E00
        color_hex = f"{r:02X}{g:02X}{b:02X}"
        shd.set(qn("w:fill"), color_hex)
        shd.set(qn("w:val"), "solid")  # Use solid fill, not clear
        shd.set(qn("w:color"), "auto")  # Auto color for text (not needed for fill)

    def _set_cell_borders(self, cell, color_hex: str = "808080", size: str = "4") -> None:
        """
        Apply borders to a table cell.
        
        Args:
            cell: The table cell to apply borders to
            color_hex: Border color in hex format (default: "808080" for gray)
            size: Border size (default: "4" for thin border)
        """
        tc_pr = cell._element.get_or_add_tcPr()
        
        # Create border elements
        def create_border(side: str):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color_hex)
            return border
        
        # Get or create tcBorders element
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        
        # Add all borders
        tc_borders.append(create_border("top"))
        tc_borders.append(create_border("bottom"))
        tc_borders.append(create_border("left"))
        tc_borders.append(create_border("right"))

    def _set_cell_margin_bottom(self, cell, points: float) -> None:
        """
        Set the bottom margin of a table cell to create row gap (e.g. between two rows of boxes).
        points: margin in points (e.g. 16 for ~16–24px visual gap). Stored as dxa (1/20 pt).
        """
        dxa = int(points * 20)  # Word tcMar uses twentieths of a point
        tc_pr = cell._element.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        bottom = tc_mar.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            tc_mar.append(bottom)
        bottom.set(qn("w:w"), str(dxa))
        bottom.set(qn("w:type"), "dxa")

    def _set_cell_margins_tight(self, cell, points: float = 4) -> None:
        """Set all cell margins to a small value for tight padding. points in pt, stored as dxa."""
        dxa = int(points * 20)
        tc_pr = cell._element.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for tag in ("top", "start", "bottom", "end"):
            el = tc_mar.find(qn(f"w:{tag}"))
            if el is None:
                el = OxmlElement(f"w:{tag}")
                tc_mar.append(el)
            el.set(qn("w:w"), str(dxa))
            el.set(qn("w:type"), "dxa")

    def _apply_severity_color(self, cell, severity: str) -> None:
        """Apply color coding to severity cell text."""
        if not cell.paragraphs[0].runs:
            return

        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        severity_upper = severity.upper()

        if severity_upper == "CRITICAL":
            run.font.color.rgb = BrandColors.RED_CRITICAL
            run.font.bold = True
        elif severity_upper == "HIGH":
            run.font.color.rgb = BrandColors.ORANGE_HIGH
            run.font.bold = True

    def _apply_risk_color(self, cell, risk: str) -> None:
        """Apply color coding to risk level cell."""
        if not cell.paragraphs[0].runs:
            return

        run = cell.paragraphs[0].runs[0]
        run.font.name = "Arial"
        risk_upper = risk.upper()

        if risk_upper in ("CRITICAL", "P1"):
            run.font.color.rgb = BrandColors.RED_CRITICAL
            run.font.bold = True
        elif risk_upper in ("HIGH", "P2"):
            run.font.color.rgb = BrandColors.ORANGE_HIGH
            run.font.bold = True

    def _set_paragraph_font(
        self,
        paragraph,
        size: Pt | None = None,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        font_name: str = "Arial"
    ) -> None:
        """Set font properties for a paragraph's runs."""
        for run in paragraph.runs:
            run.font.name = font_name
            if size:
                run.font.size = size
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

    def _add_styled_paragraph(
        self,
        text: str,
        size: Pt = FontSizes.BODY,
        bold: bool = False,
        italic: bool = False,
        color: RGBColor | None = None,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        style: str | None = None
    ):
        """Add a paragraph with custom styling."""
        if style:
            para = self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text)

        if alignment:
            para.alignment = alignment

        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = size
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

        return para

    def _create_metric_card_table(
        self,
        metrics: list[tuple[str, str, str]]
    ):
        """
        Create a metric card table (like "4 New This Week" cards).

        Args:
            metrics: List of tuples (number, title, subtitle)
        """
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.autofit = True

        for i, (number, title, subtitle) in enumerate(metrics):
            cell = table.rows[0].cells[i]

            # Clear default paragraph
            cell.paragraphs[0].clear()

            # Add number (large, bold)
            num_para = cell.paragraphs[0]
            num_run = num_para.add_run(number)
            num_run.font.name = "Arial"
            num_run.font.size = Pt(24)
            num_run.font.bold = True
            num_run.font.color.rgb = BrandColors.ORANGE_PRIMARY

            # Add title
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.name = "Arial"
            title_run.font.size = FontSizes.BODY_SMALL
            title_run.font.bold = True

            # Add subtitle
            sub_para = cell.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = "Arial"
            sub_run.font.size = FontSizes.FOOTNOTE
            sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            sub_run.font.italic = True

            # Center align all paragraphs
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return table

    def _format_date_range(self, start_date: datetime | None = None, end_date: datetime | None = None) -> str:
        """Format a date range string."""
        if end_date is None:
            end_date = self.created_at
        if start_date is None:
            # Default to 7 days before end_date for weekly
            from datetime import timedelta
            start_date = end_date - timedelta(days=6)

        return f"{start_date.strftime('%B %d')} to {end_date.strftime('%d, %Y')}"

    def _get_week_number(self) -> int:
        """Get ISO week number for the report date."""
        return self.created_at.isocalendar()[1]

    def _get_year(self) -> int:
        """Get year for the report date."""
        return self.created_at.year

    def _get_banner_path(self) -> str | None:
        """
        Get the path to the report banner image.
        
        Returns:
            Path to Bulletin_banner.jpg if found, None otherwise
        """
        # Try root directory first
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        banner_path = os.path.join(root_dir, "Bulletin_banner.jpg")
        
        if os.path.exists(banner_path):
            return banner_path
        
        # Try current working directory
        alt_path = "Bulletin_banner.jpg"
        if os.path.exists(alt_path):
            return alt_path
        
        return None

    def _add_image(self, image_path: str, width: Inches | None = None, height: Inches | None = None) -> None:
        """
        Add an image to the document.
        
        Args:
            image_path: Path to the image file
            width: Optional width in Inches (defaults to page width)
            height: Optional height in Inches (maintains aspect ratio if only width specified)
        """
        if not os.path.exists(image_path):
            logger.warning(f"Image file not found: {image_path}")
            return
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run()
        
        # Default to page width if not specified
        if width is None:
            width = Inches(6.5)  # Standard page width minus margins
        
        run.add_picture(image_path, width=width, height=height)

    def _set_default_font(self, font_name: str = "Arial") -> None:
        """
        Set the default font for the entire document.
        
        Args:
            font_name: Name of the font to use (default: "Arial")
        """
        if self.doc is None:
            return
        
        # Set font for Normal style (default paragraph style)
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = font_name
        
        # Also set for Heading styles
        for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
            if style_name in self.doc.styles:
                heading_style = self.doc.styles[style_name]
                heading_font = heading_style.font
                heading_font.name = font_name
                # Set Heading 1 font size to 14pt
                if style_name == 'Heading 1':
                    heading_font.size = Pt(14)

    def _configure_page_settings(self) -> None:
        """
        Configure page setup to match CTI template (Page Setup dialog).
        Margins: Top 1.25", Bottom 0.88", Left 0.75", Right 0.75", Gutter 0"
        Paper: Letter 8.5" x 11", Portrait
        Layout: Header from edge 0.49", Footer from edge 0.49", Vertical alignment Top
        Paragraph: Before 0 pt, After 3 pt, Line spacing Single
        """
        if self.doc is None:
            return

        section = self.doc.sections[0]

        # Paper (Letter 8.5" x 11", Portrait)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.orientation = WD_ORIENT.PORTRAIT

        # Margins
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(0.88)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.gutter = Inches(0)

        # Layout: header/footer from edge
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)

        # Vertical alignment: Top (via sectPr if not default)
        self._set_section_vertical_alignment_top(section)

        # Paragraph spacing (Before 0 pt, After 3 pt)
        normal_style = self.doc.styles["Normal"]
        pf = normal_style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            if style_name in self.doc.styles:
                pf = self.doc.styles[style_name].paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(3)
                pf.line_spacing = 1.0

    def _set_section_vertical_alignment_top(self, section) -> None:
        """Set section vertical alignment to Top (Layout tab)."""
        sect_pr = section._sectPr
        pg_pr = sect_pr.find(qn("w:pgPr"))
        if pg_pr is None:
            pg_pr = OxmlElement("w:pgPr")
            sect_pr.insert(0, pg_pr)
        v_align = pg_pr.find(qn("w:vAlign"))
        if v_align is None:
            v_align = OxmlElement("w:vAlign")
            pg_pr.append(v_align)
        v_align.set(qn("w:val"), "top")

    def _set_document_background(self, color_hex: str = "1E1E1E") -> None:
        """
        Set the document page background color in the .docx (stored in the file).
        Use PAGE_WHITE ("FFFFFF") for print-style so the document renders as white
        paper regardless of editor or app theme (e.g. Cursor/Word dark mode).
        Requires w:background on section and w:displayBackgroundShape in settings.
        """
        if self.doc is None:
            return
        color_hex = color_hex.replace("#", "").upper()
        section = self.doc.sections[0]
        sect_pr = section._sectPr
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), color_hex)
        sect_pr.insert(0, bg)
        try:
            settings_el = self.doc.settings.element
            if settings_el.find(qn("w:displayBackgroundShape")) is None:
                disp = OxmlElement("w:displayBackgroundShape")
                settings_el.insert(0, disp)
        except Exception:
            pass

    def _apply_font_to_run(self, run, font_name: str = "Arial") -> None:
        """
        Apply font name to a text run.
        
        Args:
            run: The text run to apply font to
            font_name: Name of the font to use (default: "Arial")
        """
        run.font.name = font_name

    def _add_section_divider(self) -> None:
        """
        Add a faint horizontal divider line below a major section title.
        Creates a paragraph with a bottom border to visually separate sections.
        """
        if self.doc is None:
            return
        
        divider_para = self.doc.add_paragraph()
        divider_para.paragraph_format.space_before = Pt(0)  # No space before (tight with content above)
        divider_para.paragraph_format.space_after = Pt(12)  # Space after (before next section title)
        
        # Add bottom border to create divider line
        p_pr = divider_para._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        
        # Create bottom border (gray, 20% opacity = light gray)
        bottom_border = OxmlElement("w:bottom")
        bottom_border.set(qn("w:val"), "single")
        bottom_border.set(qn("w:sz"), "6")  # Thin line
        bottom_border.set(qn("w:space"), "1")
        bottom_border.set(qn("w:color"), "CCCCCC")  # Light gray (approximately 20% opacity effect)
        p_bdr.append(bottom_border)

    def _clear_paragraph_borders(self, paragraph) -> None:
        """Remove any paragraph borders (e.g. line under title) so no line is shown."""
        p_pr = paragraph._element.find(qn("w:pPr"))
        if p_pr is not None:
            p_bdr = p_pr.find(qn("w:pBdr"))
            if p_bdr is not None:
                p_pr.remove(p_bdr)

    def _add_banner_header(self) -> None:
        """
        Add the standard banner image to the document header section.
        This method should be called at the start of _add_header() in subclasses.
        """
        banner_path = self._get_banner_path()
        if banner_path:
            # Get the document section and its header
            section = self.doc.sections[0]
            header = section.header
            
            # Clear any existing content in the header
            header.paragraphs[0].clear()
            
            # Add the banner image to the header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = header_para.add_run()
            # Calculate width to fit within margins (page width minus left and right margins)
            # Standard page width is 8.5", minus left (0.75") and right (0.75") margins = 7"
            run.add_picture(banner_path, width=Inches(7.0))
        else:
            logger.warning("Bulletin_banner.jpg not found - banner will be omitted from report")
