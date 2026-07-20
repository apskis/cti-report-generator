"""
Weekly CTI Report Generator.

Generates weekly threat intelligence reports matching the branded template.
"""

import logging
from datetime import timedelta
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.reports.base import BaseReportGenerator, BrandColors, FontSizes
from src.reports.registry import register_report_generator

logger = logging.getLogger(__name__)


@register_report_generator("weekly")
class WeeklyReportGenerator(BaseReportGenerator):
    """
    Weekly CTI Report Generator.

    Generates reports matching the CTI_Weekly_Report_Template_Example.docx format.

    Template structure:
    - Report ID header (CTI-WK-YYYY-WW)
    - Title: "Cyber Threat Intelligence Weekly Report" (orange, 18pt bold)
    - Date range subtitle
    - Executive Summary
    - This Week at a Glance (metric cards)
    - Vulnerability Exposure (CVE table with weeks tracked)
    - Sector Threat Activity (threat actor table)
    - Exploitation Indicators
    - Recommended Actions
    - Footer with contact info and data sources
    """

    @property
    def report_type(self) -> str:
        return "weekly"

    @property
    def filename_prefix(self) -> str:
        return "CTI_Weekly_Report"

    def generate(self, analysis_result: dict[str, Any]) -> Document:
        """
        Generate the weekly report document.

        Args:
            analysis_result: Dictionary containing:
                - executive_summary: str
                - statistics: dict with counts
                - cve_analysis: list of CVE dicts
                - apt_activity: list of APT dicts
                - recommendations: list of strings
                - exploitation_indicators: list of strings (optional)

        Returns:
            Document object
        """
        try:
            logger.info("Generating Weekly CTI Report")
            self.doc = Document()

            # Page setup: margins, Letter 8.5x11", header/footer distance, paragraph spacing
            self._configure_page_settings()

            # Print-style: white page so document displays correctly in light and dark mode.
            self._set_document_background(BrandColors.PAGE_WHITE)
            normal_style = self.doc.styles["Normal"]
            normal_style.font.color.rgb = BrandColors.TEXT_DARK
            normal_style.font.name = "Arial"

            # Calculate date range (Monday to Sunday of current week)
            self._calculate_date_range()

            # Store period_end for filename generation (use current week number, not lookback start)
            self._report_week_start = self.period_end

            # Add sections in order
            self._add_header()
            self._add_week_at_glance(analysis_result)
            self._add_executive_summary(analysis_result)
            self._add_vulnerability_exposure(analysis_result)
            self._add_sector_threat_activity(analysis_result)
            self._add_active_campaigns(analysis_result)
            self._add_industry_incidents(analysis_result)
            self._add_recommended_actions(analysis_result)
            self._add_resources_section(analysis_result)
            self._add_footer()

            logger.info("Weekly CTI Report generated successfully")
            return self.doc

        except Exception as e:
            logger.error(f"Error generating weekly report: {str(e)}", exc_info=True)
            raise

    def _calculate_date_range(self) -> None:
        """Calculate the reporting period based on actual data lookback window."""
        from src.core.config import collector_config

        today = self.created_at
        lookback_days = collector_config.nvd_lookback_days
        self.period_end = today
        self.period_start = today - timedelta(days=lookback_days)
        self.lookback_days = lookback_days

    def _add_header(self) -> None:
        """Add banner, report code in header bar; then title block on single line."""
        week_num = self._get_week_number()
        year = self._get_year()
        report_id = f"CTI-WK-{year}-{week_num:02d}"
        date_range = (
            f"{self.lookback_days}-Day Lookback | "
            f"{self.period_start.strftime('%B %d')} to {self.period_end.strftime('%B %d, %Y')}"
        )

        # Banner in header
        self._add_banner_header()

        # Same horizontal line as bottom of header: report code only (gray monospace), right-aligned
        section = self.doc.sections[0]
        header = section.header
        id_para = header.add_paragraph()
        id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        id_para.paragraph_format.space_before = Pt(2)
        id_para.paragraph_format.space_after = Pt(0)
        id_run = id_para.add_run(report_id)
        id_run.font.size = FontSizes.FOOTNOTE
        id_run.font.color.rgb = BrandColors.GRAY_LIGHT
        id_run.font.name = "Consolas"

        # Body: main title, 18pt, orange — use Normal style so no line/border under title
        title_para = self.doc.add_paragraph()
        title_para.style = self.doc.styles["Normal"]
        title_para.paragraph_format.space_before = Pt(4)
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.line_spacing = 0.9  # Tighter line height
        title_run = title_para.add_run("Cyber Threat Intelligence Weekly Report")
        title_run.font.size = FontSizes.TITLE  # 18pt
        title_run.font.bold = True
        title_run.font.color.rgb = BrandColors.ORANGE_DESIGN
        title_run.font.name = "Arial"
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._clear_paragraph_borders(title_para)  # No line under the title

        # Subtitle directly below with minimal vertical spacing (dark gray on white)
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(date_range)
        date_run.font.size = FontSizes.BODY_SMALL
        date_run.font.color.rgb = BrandColors.GRAY_LIGHT
        date_run.font.name = "Arial"
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.paragraph_format.space_before = Pt(2)
        date_para.paragraph_format.space_after = Pt(6)

        # Minimal gap before Executive Summary
        self.doc.add_paragraph()

    def _add_executive_summary(self, analysis_result: dict[str, Any]) -> None:
        """Add summary: bold orange heading, white body, tight spacing."""
        logger.info("Adding Summary section")

        # Heading 1 style: 14pt, orange
        heading_para = self.doc.add_paragraph()
        heading_para.style = self.doc.styles["Heading 1"]
        heading_para.paragraph_format.space_before = Pt(6)
        heading_para.paragraph_format.space_after = Pt(4)
        heading_run = heading_para.add_run("Summary")
        heading_run.font.bold = True
        heading_run.font.color.rgb = BrandColors.ORANGE_DESIGN
        heading_run.font.size = FontSizes.HEADING_1  # 14pt
        heading_run.font.name = "Arial"

        # Body: left-aligned, dark text on white, tight paragraph spacing
        summary = analysis_result.get("executive_summary", "No executive summary available.")

        # Add trend context if available
        cve_trends = analysis_result.get("cve_trends", {})
        actor_trends = analysis_result.get("actor_trends", {})

        if cve_trends or actor_trends:
            # Prepend trend summary to executive summary
            trend_intro = []

            if cve_trends and cve_trends.get("trend_summary"):
                trend_intro.append(f"Trend Analysis: {cve_trends['trend_summary']}")

            if actor_trends and actor_trends.get("trend_summary"):
                trend_intro.append(f"Actor Trends: {actor_trends['trend_summary']}")

            if trend_intro:
                summary = "\n\n".join(trend_intro) + "\n\n" + summary

        para = self.doc.add_paragraph(summary)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.5
        for run in para.runs:
            run.font.size = FontSizes.BODY
            run.font.color.rgb = BrandColors.TEXT_DARK
            run.font.name = "Arial"

    def _add_week_at_glance(self, analysis_result: dict[str, Any]) -> None:
        """Add 'This Week at a Glance': bold orange heading; space before subtitle."""
        logger.info("Adding This Week at a Glance section")

        # Heading 1 style: "This Week at a Glance", 14pt, orange
        glance_para = self.doc.add_paragraph()
        glance_para.style = self.doc.styles["Heading 1"]
        glance_para.paragraph_format.space_before = Pt(0)
        glance_para.paragraph_format.space_after = Pt(0)
        r1 = glance_para.add_run("This Week ")
        r1.font.bold = True
        r1.font.color.rgb = BrandColors.ORANGE_DESIGN
        r1.font.size = FontSizes.HEADING_1  # 14pt
        r1.font.name = "Arial"
        r2 = glance_para.add_run("at a Glance")
        r2.font.bold = True
        r2.font.underline = False
        r2.font.color.rgb = BrandColors.ORANGE_DESIGN
        r2.font.size = FontSizes.HEADING_1  # 14pt
        r2.font.name = "Arial"

        # Subtitle: italic, light gray, smaller; extra space between title and this paragraph
        subtitle = self.doc.add_paragraph()
        subtitle.paragraph_format.space_before = Pt(10)
        subtitle.paragraph_format.space_after = Pt(4)
        sub_run = subtitle.add_run(
            f"Threat intelligence metrics from Intel471, CrowdStrike, and OSINT "
            f"({self.period_start.strftime('%B %d')} to {self.period_end.strftime('%B %d, %Y')})."
        )
        sub_run.font.size = Pt(9)
        sub_run.font.italic = True
        sub_run.font.color.rgb = BrandColors.GRAY_LIGHT
        sub_run.font.name = "Arial"

        # Calculate deterministic statistics from actual data
        stats = self._calculate_statistics(analysis_result)
        logger.info(f"Calculated stats: {stats}")

        # Four threat intelligence metrics in 2x2 grid
        # Row 1: Threat Actors, Active Campaigns
        # Row 2: Exploited CVEs, Peer Incidents
        metrics_row1 = [
            (str(stats.get("threat_actors", 0)), "Threat Actors", "APT groups active this week"),
            (str(stats.get("active_campaigns", 0)), "Active Campaigns", "Operations from underground intel"),
        ]
        metrics_row2 = [
            (str(stats.get("exploited_cves", 0)), "Exploited CVEs", "Known exploitation activity"),
            (str(stats.get("peer_incidents", 0)), "Peer Incidents", "Company breaches observed"),
        ]
        # Two separate tables: 2 boxes, then space, then 2 boxes (2x2 grid)
        self._create_metric_cards(metrics_row1)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(20)  # Vertical gap between the two rows of boxes
        self._create_metric_cards(metrics_row2)

        # Gap before summary section
        spacer2 = self.doc.add_paragraph()
        spacer2.paragraph_format.space_before = Pt(0)
        spacer2.paragraph_format.space_after = Pt(12)

    def _create_metric_cards(self, metrics: list[tuple]) -> None:
        """Create a row of metric cards with light gray backgrounds per spec."""
        table = self.doc.add_table(rows=1, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for col_idx, (number, title, subtitle) in enumerate(metrics):
            cell = table.rows[0].cells[col_idx]

            # Light gray background per CTI_Weekly_Report_Template_Spec.json (F5F5F5)
            self._set_cell_shading(cell, BrandColors.METRIC_CARD_BG)
            # Add border to make card visible
            self._set_cell_borders(cell, "CCCCCC", "4")

            # Clear and build content
            cell.paragraphs[0].clear()

            # 1. Colored number: Red (threats/exploits), Orange (campaigns), Green (N/A for now); 22pt
            num_para = cell.paragraphs[0]
            num_run = num_para.add_run(number)
            num_run.font.size = Pt(22)
            num_run.font.bold = True
            num_run.font.name = "Arial"
            if title in ("Threat Actors", "Exploited CVEs", "Peer Incidents"):
                num_run.font.color.rgb = BrandColors.RED_HIGH_RISK  # High-risk / alert
            elif title == "Active Campaigns":
                num_run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Ongoing operations
            else:
                num_run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Default
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_para.paragraph_format.space_after = Pt(2)

            # Title (bold, explicit black color)
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = FontSizes.BODY_SMALL
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Explicit black
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(2)

            # Subtitle (smaller gray)
            sub_para = cell.add_paragraph()
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.size = FontSizes.FOOTNOTE
            sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            sub_run.font.name = "Arial"
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(0)
            sub_para.paragraph_format.space_after = Pt(8)  # Spacing between last line and box border

    def _format_exposure_cell(self, cve: dict[str, Any]) -> str:
        """Format Exposure column as count of servers/databases/endpoints from Rapid7 only. Never use vulnerability description."""
        # Debug: log what we received for this CVE
        cve_id = cve.get("cve_id", "Unknown")

        # Only use fields that represent asset/server/database counts (never description)
        raw = cve.get("exposure") or cve.get("exposure_summary") or cve.get("asset_count") or cve.get("affected_assets")

        logger.debug(
            f"CVE {cve_id} exposure fields - exposure: {cve.get('exposure')}, "
            f"asset_count: {cve.get('asset_count')}, raw: {raw}"
        )

        if isinstance(raw, str) and raw.strip():
            s = raw.strip()
            # Use only if it looks like a count: "N servers", "N databases", "N endpoints", or "Production"
            lower = s.lower()
            if any(
                lower.endswith(x)
                for x in (
                    "servers",
                    "server",
                    "databases",
                    "database",
                    "endpoints",
                    "endpoint",
                    "systems",
                    "system",
                    "workstations",
                    "workstation",
                    "cloud servers",
                    "cloud instances",
                )
            ):
                return s[:50]
            if s in ("Production", "N/A", "—", "-"):
                return s
            parts = s.split()
            if len(parts) >= 2 and parts[0].isdigit():
                return s[:50]
        # Build from numeric fields (Rapid7 / pipeline can supply these)
        for key, label in (
            ("server_count", "server"),
            ("servers", "server"),
            ("asset_count", "server"),
            ("host_count", "server"),
            ("hosts", "server"),
            ("database_count", "database"),
            ("databases", "database"),
            ("endpoint_count", "endpoint"),
            ("endpoints", "endpoint"),
            ("affected_asset_count", "server"),
        ):
            val = cve.get(key)
            if val is not None:
                try:
                    n = int(val)
                    return f"{n} {label}{'s' if n != 1 else ''}"
                except (TypeError, ValueError):
                    pass
        return "N/A"

    def _style_heading_1(self, paragraph) -> None:
        """Apply Heading 1 look: orange, 14pt, Arial to all runs in the paragraph."""
        for run in paragraph.runs:
            run.font.color.rgb = BrandColors.ORANGE_DESIGN
            run.font.size = FontSizes.HEADING_1
            run.font.name = "Arial"

    @staticmethod
    def _extract_count(exposure_str: str) -> int:
        """Extract numeric count from exposure string like '7 systems'."""
        try:
            return int(str(exposure_str).split()[0])
        except (ValueError, IndexError, AttributeError):
            return 0

    def _calculate_statistics(self, analysis_result: dict[str, Any]) -> dict[str, int]:
        """
        Calculate threat intelligence statistics from actual data.

        Weekly Tactical Report metrics (no environment/Rapid7 data):
        - Threat Actors: Unique APT actors tracked this week
        - Active Campaigns: Campaign names from Intel471 tags
        - Exploited CVEs: CVEs with exploitation evidence
        - Peer Incidents: Company breaches from OSINT
        """
        cve_analysis = analysis_result.get("cve_analysis", [])
        apt_activity = analysis_result.get("apt_activity", [])

        # Metric 1: Threat Actors (unique actor names)
        threat_actors_count = len(apt_activity)

        # Metric 2: Active Campaigns (from AI-structured campaign data)
        active_campaigns = analysis_result.get("active_campaigns", [])
        active_campaigns_count = len(active_campaigns)

        # Metric 3: Exploited CVEs (CVEs with exploitation evidence)
        exploited_cves_count = sum(
            1 for cve in cve_analysis if cve.get("actively_exploited") or cve.get("targeted_by_actors")
        )

        # Metric 4: Peer Incidents (from industry_incidents - includes both OSINT and Intel471)
        industry_incidents = analysis_result.get("industry_incidents", [])
        peer_incidents_count = len(industry_incidents)

        return {
            "threat_actors": threat_actors_count,
            "active_campaigns": active_campaigns_count,
            "exploited_cves": exploited_cves_count,
            "peer_incidents": peer_incidents_count,
            # Legacy fallbacks for backwards compatibility
            "total_cves": len(cve_analysis),
            "critical_count": sum(1 for cve in cve_analysis if cve.get("severity", "").upper() == "CRITICAL"),
        }

    def _calculate_resolved_count(self, current_cves: list[dict[str, Any]]) -> int:
        """
        Calculate resolved CVEs by comparing with previous week.
        Reads last week's report from reports/ directory if available.
        """
        try:
            import os
            from datetime import timedelta

            # Calculate last week's date
            last_week = self.created_at - timedelta(days=7)
            last_week_filename = f"CTI_Weekly_Report_{last_week.strftime('%Y-%m-%d')}.docx"

            # Check in both reports/ and reports_test/ directories
            for report_dir in ["reports", "reports_test", "."]:
                last_week_path = os.path.join(report_dir, last_week_filename)
                if os.path.exists(last_week_path):
                    # We can't easily read DOCX, so check for a JSON cache instead
                    json_cache = last_week_path.replace(".docx", "_cves.json")
                    if os.path.exists(json_cache):
                        import json

                        with open(json_cache) as f:
                            last_week_cves = set(json.load(f))

                        # CVEs that were in last week but not this week
                        current_cve_ids = set(cve.get("cve_id") for cve in current_cves)
                        resolved = len(last_week_cves - current_cve_ids)
                        logger.info(f"Resolved CVEs: {resolved} (comparing with {last_week_filename})")
                        return resolved

            logger.debug("No previous week's data found for comparison")
            return 0

        except Exception as e:
            logger.debug(f"Could not calculate resolved count: {e}")
            return 0

    def _save_cve_cache(self, cve_analysis: list[dict[str, Any]], output_path: str):
        """Save CVE IDs to JSON cache for next week's comparison."""
        try:
            import json

            cve_ids = [cve.get("cve_id") for cve in cve_analysis if cve.get("cve_id")]
            cache_path = output_path.replace(".docx", "_cves.json")
            with open(cache_path, "w") as f:
                json.dump(cve_ids, f)
            logger.debug(f"Saved CVE cache to {cache_path}")
        except Exception as e:
            logger.debug(f"Could not save CVE cache: {e}")

    def _group_cves_by_technology(self, cve_analysis: list[dict[str, Any]]) -> tuple:
        """
        Group CVEs by technology/product family for cleaner reporting.

        Returns:
            tuple: (grouped_items, individual_cves, all_cves)
        """
        from collections import defaultdict

        # Define grouping patterns
        groups = defaultdict(list)
        individual_cves = []

        # Minimum CVEs needed to form a group (reduced to 2 for more aggressive grouping)
        MIN_GROUP_SIZE = 2

        # Debug: log first few products to see what we're working with
        if cve_analysis:
            sample_products = [cve.get("affected_product", "?")[:50] for cve in cve_analysis[:5]]
            logger.info(f"Sample products for grouping: {sample_products}")

        for cve in cve_analysis:
            product = cve.get("affected_product", "").strip()
            product_lower = product.lower()

            # Group WordPress products (plugins, themes, core all together)
            if "wordpress" in product_lower or "wp plugin" in product_lower or "wp theme" in product_lower:
                groups["WordPress Products"].append(cve)
            # Group by vendor for other common products
            elif product.startswith("Microsoft "):
                groups["Microsoft Products"].append(cve)
            elif product.startswith("Apache "):
                groups["Apache Products"].append(cve)
            elif product.startswith("VMware ") or product.startswith("VMWare "):
                groups["VMware Products"].append(cve)
            elif product.startswith("Oracle "):
                groups["Oracle Products"].append(cve)
            elif product.startswith("Cisco "):
                groups["Cisco Products"].append(cve)
            elif product.startswith("Adobe "):
                groups["Adobe Products"].append(cve)
            elif "linux kernel" in product_lower or "kernel" in product_lower:
                groups["Linux Kernel"].append(cve)
            elif "python" in product_lower:
                groups["Python Packages"].append(cve)
            elif "node" in product_lower or "npm" in product_lower:
                groups["Node.js/NPM Packages"].append(cve)
            else:
                # Don't group unique products
                individual_cves.append(cve)

        # Convert groups to summary items, only if they meet minimum size
        grouped_items = []
        for group_name, cves in groups.items():
            if len(cves) >= MIN_GROUP_SIZE:
                # Calculate aggregate metrics
                total_systems = sum(self._extract_count(cve.get("exposure", "0")) for cve in cves)
                max_weeks = max((cve.get("weeks_detected", 1) for cve in cves), default=1)

                logger.info(f"Grouping {len(cves)} CVEs into '{group_name}'")

                grouped_items.append(
                    {
                        "is_group": True,
                        "group_name": group_name,
                        "cve_count": len(cves),
                        "cves": cves,
                        "exposure": f"{total_systems} systems" if total_systems > 0 else "Multiple",
                        "weeks_detected": max_weeks,
                    }
                )
            else:
                # If group is too small, treat as individual CVEs
                individual_cves.extend(cves)

        # Sort individual CVEs by severity/exposure (show most critical first)
        individual_cves.sort(key=lambda x: (-self._extract_count(x.get("exposure", "0")), -x.get("weeks_detected", 1)))

        # Limit individual CVEs to top 5 most critical
        if len(individual_cves) > 5:
            logger.info(f"Limiting individual CVEs to top 5 (total: {len(individual_cves)})")
            individual_cves = individual_cves[:5]

        return grouped_items, individual_cves, cve_analysis

    def _is_actively_exploited(self, cve: dict[str, Any]) -> tuple:
        """
        Check if a CVE is actively exploited.
        Returns (is_exploited: bool, reason: str)
        """
        # Check boolean exploited flag (from enrichment)
        if cve.get("exploited", False):
            return (True, "Confirmed exploitation")

        # Check in_cisa_kev flag (from enrichment)
        if cve.get("in_cisa_kev", False):
            return (True, "CISA KEV")

        # Check targeted_by_actors field (from AI analysis)
        targeted_by = str(cve.get("targeted_by_actors", "")).strip()
        if targeted_by and targeted_by.lower() not in ("", "none", "unknown", "n/a"):
            return (True, f"Targeted by: {targeted_by}")

        # Check exploited_by field (from enrichment)
        exploited_by = str(cve.get("exploited_by", "")).upper()

        if exploited_by and exploited_by != "NONE KNOWN":
            # Check for CISA KEV
            if "CISA KEV" in exploited_by:
                return (True, "CISA KEV")

            # Check for ransomware
            if "RANSOMWARE" in exploited_by:
                return (True, "Ransomware campaigns")

            # Check for active exploitation mention
            if "ACTIVE EXPLOITATION" in exploited_by:
                return (True, "Active exploitation")

            # Check for threat actors
            threat_actors = ["APT", "LAZARUS", "PANDA", "BEAR", "KITTEN", "DRAGON", "SPIDER", "GROUP"]
            for actor in threat_actors:
                if actor in exploited_by and "NONE" not in exploited_by:
                    return (True, "Threat actor activity")

        # Check known_ransomware field (from KEV enrichment)
        known_ransomware = str(cve.get("known_ransomware", "")).upper()
        if known_ransomware and known_ransomware not in ("UNKNOWN", ""):
            return (True, "Ransomware campaigns")

        return (False, "")

    def _add_vulnerability_exposure(self, analysis_result: dict[str, Any]) -> None:
        """Add Exploited Vulnerabilities section - intelligence-sourced CVEs only (no environment data)."""
        logger.info("Adding Exploited Vulnerabilities section")

        h = self.doc.add_heading("Exploited Vulnerabilities", level=1)
        self._style_heading_1(h)

        cve_analysis = analysis_result.get("cve_analysis", [])

        if not cve_analysis:
            self.doc.add_paragraph("No exploited vulnerabilities identified from threat intelligence this week.")
            self.doc.add_paragraph()
            return

        # Filter to only exploited CVEs (from threat intelligence sources)
        exploited_cves = [
            cve
            for cve in cve_analysis
            if cve.get("actively_exploited") or cve.get("targeted_by_actors") or cve.get("in_cisa_kev")
        ]

        if not exploited_cves:
            self.doc.add_paragraph("No exploited vulnerabilities identified from threat intelligence this week.")
            self.doc.add_paragraph()
            return

        logger.info(f"Exploited CVEs: {len(exploited_cves)}")

        # Add intro with actionable context
        intro = self.doc.add_paragraph()
        intro.space_before = Pt(0)
        intro.space_after = Pt(6)

        intro_run = intro.add_run(
            "CVEs with confirmed exploitation from threat intelligence (CISA KEV, threat actors, ransomware). "
        )
        intro_run.font.size = Pt(8)
        intro_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        intro_run.font.italic = True

        # Add actionable context
        action_run = intro.add_run(
            "Review vulnerability scan results and prioritize patching. "
            "Check affected systems for signs of exploitation."
        )
        action_run.font.size = Pt(8)
        action_run.font.color.rgb = BrandColors.TEXT_DARK
        action_run.font.bold = True

        # Create table: CVE ID | Affected Product | Exploitation Evidence | Source
        table = self.doc.add_table(rows=1, cols=4)
        headers = ["CVE ID", "Affected Product", "Exploitation Evidence", "Source"]
        header_cells = table.rows[0].cells
        Pt(7)

        # Header row: orange background, white bold text
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = BrandColors.WHITE
            self._set_cell_shading(cell, BrandColors.TABLE_HEADER_BG)
            self._set_cell_borders(cell, "CCCCCC")
            self._set_cell_margins_tight(cell, 2)

        # Add exploited CVEs (no grouping - straight list from threat intelligence)
        for cve in exploited_cves:
            row = table.add_row()
            cells = row.cells

            # Column 0: CVE ID (bold)
            cells[0].text = ""
            cve_para = cells[0].paragraphs[0]
            cve_run = cve_para.add_run(cve.get("cve_id", "N/A"))
            cve_run.font.size = FontSizes.SUBTITLE
            cve_run.font.color.rgb = BrandColors.TEXT_DARK
            cve_run.font.bold = True

            # Column 1: Affected Product
            cells[1].text = cve.get("affected_product", cve.get("product", "N/A"))[:60]

            # Column 2: Exploitation Evidence (source of exploitation intel)
            evidence = self._format_exploitation_evidence(cve)
            cells[2].text = evidence

            # Column 3: Source Citations - NORMAL SIZE with source names
            cells[3].text = ""
            cite_para = cells[3].paragraphs[0]
            source_citations = cve.get("source_citations", [])
            if isinstance(source_citations, list) and source_citations:
                # Build citation map to get numbers
                citation_map = self._build_citation_map(analysis_result)

                # Add each source with its number
                for idx, src in enumerate(source_citations):
                    if idx > 0:
                        # Add comma separator between sources
                        cite_para.add_run(", ")

                    # Get citation number
                    cite_num = citation_map.get(src, "?")

                    # Add [#] Source (normal size, not superscript)
                    cite_run = cite_para.add_run(f"[{cite_num}] {src}")
                    cite_run.font.size = FontSizes.SUBTITLE
                    cite_run.font.color.rgb = BrandColors.TEXT_DARK
            else:
                # Default to NVD
                cite_run = cite_para.add_run("[1] NVD")
                cite_run.font.size = FontSizes.SUBTITLE
                cite_run.font.color.rgb = BrandColors.TEXT_DARK

            # Styling for all columns - white background (no fill)
            for idx in range(4):
                self._clear_cell_shading(cells[idx])  # White/no fill
                self._set_cell_borders(cells[idx], "CCCCCC")
                for para in cells[idx].paragraphs:
                    for run in para.runs:
                        run.font.size = FontSizes.SUBTITLE
                        run.font.color.rgb = BrandColors.TEXT_DARK
                        if idx == 0:  # Bold CVE ID
                            run.font.bold = True

        # Set column widths (CVE ID narrower, Product wider, Evidence medium, Source narrow)
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(2.5)
        table.columns[2].width = Inches(1.8)
        table.columns[3].width = Inches(0.9)

        # Caption
        caption = self.doc.add_paragraph()
        caption.space_before = Pt(4)
        caption.space_after = Pt(8)
        caption_text = (
            f"Table: {len(exploited_cves)} exploited CVEs from threat intelligence sources "
            f"(Intel471, CrowdStrike, CISA KEV, OSINT)."
        )
        caption_run = caption.add_run(caption_text)
        caption_run.font.size = Pt(7)
        caption_run.font.italic = True
        caption_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        caption_run.font.name = "Arial"
        caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

        self.doc.add_paragraph()

    def _format_exploitation_evidence(self, cve: dict[str, Any]) -> str:
        """Format exploitation evidence column based on threat intelligence sources."""
        evidence_parts = []

        if cve.get("in_cisa_kev"):
            evidence_parts.append("CISA KEV")

        targeted_by = str(cve.get("targeted_by_actors", "")).strip()
        if targeted_by and targeted_by.lower() not in ("", "none", "unknown", "n/a"):
            evidence_parts.append(f"Actor: {targeted_by[:30]}")

        exploited_by = str(cve.get("exploited_by", "")).upper()
        if "RANSOMWARE" in exploited_by:
            evidence_parts.append("Ransomware")
        elif "ACTIVE EXPLOITATION" in exploited_by:
            evidence_parts.append("Active exploit")

        if not evidence_parts:
            evidence_parts.append("Exploitation confirmed")

        return "; ".join(evidence_parts[:2])  # Max 2 sources to keep readable

    def _add_sector_threat_activity(self, analysis_result: dict[str, Any]) -> None:
        """Add sector threat activity section with threat actor table."""
        logger.info("Adding Sector Threat Activity section")

        h = self.doc.add_heading("Sector Threat Activity", level=1)
        self._style_heading_1(h)

        # Intro text with actionable context
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run(
            "The following threat actors have been observed targeting organizations "
            "in the biotech, pharmaceutical, healthcare, manufacturing, and related sectors this week. "
        )
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        action_run = intro.add_run(
            "Review 'What to Monitor' column to identify relevant log sources and detection opportunities."
        )
        action_run.font.size = FontSizes.BODY_SMALL
        action_run.font.bold = True
        action_run.font.color.rgb = BrandColors.TEXT_DARK

        self.doc.add_paragraph()

        apt_activity = analysis_result.get("apt_activity", [])

        if apt_activity:
            # Create threat actor table per CTI_Weekly_Report_Template_Spec.json
            table = self.doc.add_table(rows=1, cols=4)

            # Header row: orange background, white text
            headers = ["Origin / Motivation", "Activity Observed", "What to Monitor", "Source"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                para = header_cells[i].paragraphs[0]
                para.runs[0].font.bold = True
                para.runs[0].font.size = FontSizes.SUBTITLE
                para.runs[0].font.color.rgb = BrandColors.WHITE
                self._set_cell_shading(header_cells[i], BrandColors.TABLE_HEADER_BG)
                self._set_cell_borders(header_cells[i], "CCCCCC")

            # Build citation map
            citation_map = self._build_citation_map(analysis_result)

            # Data rows: styled per spec
            for apt in apt_activity:
                row = table.add_row()
                cells = row.cells

                # Column 0: Origin/Motivation - light gray background
                country = apt.get("country", apt.get("origin", "Unknown"))
                motivation = apt.get("motivation", "Unknown")
                cells[0].text = f"{country}\n{motivation}"
                self._set_cell_shading(cells[0], BrandColors.SECTOR_ORIGIN_BG)
                self._set_cell_borders(cells[0], "CCCCCC")

                # Column 1: Activity Observed - no fill (inherit background)
                actor = apt.get("actor", apt.get("name", "Unknown"))
                activity = apt.get("activity", apt.get("description", ""))
                cells[1].text = f"{actor}, {activity}" if activity else actor
                self._clear_cell_shading(cells[1])
                self._set_cell_borders(cells[1], "CCCCCC")

                # Column 2: What to Monitor - light blue background
                monitoring = apt.get("what_to_monitor", apt.get("indicators", ""))
                if isinstance(monitoring, list):
                    monitoring = "; ".join(monitoring[:3])
                cells[2].text = monitoring
                self._set_cell_shading(cells[2], BrandColors.SECTOR_MONITOR_BG)
                self._set_cell_borders(cells[2], "CCCCCC")

                # Column 3: Source Citations - NORMAL SIZE [#] Source
                cells[3].text = ""
                cite_para = cells[3].paragraphs[0]
                source_citations = apt.get("source_citations", [])
                if isinstance(source_citations, list) and source_citations:
                    # Add each source with its number
                    for idx, src in enumerate(source_citations):
                        if idx > 0:
                            cite_para.add_run(", ")

                        # Get citation number
                        cite_num = citation_map.get(src, "?")

                        # Add [#] Source (normal size, not superscript)
                        cite_run = cite_para.add_run(f"[{cite_num}] {src}")
                        cite_run.font.size = FontSizes.SUBTITLE
                        cite_run.font.color.rgb = BrandColors.TEXT_DARK
                else:
                    # Default if no source_citations
                    cite_run = cite_para.add_run("Unknown")
                    cite_run.font.size = FontSizes.SUBTITLE
                    cite_run.font.color.rgb = BrandColors.TEXT_DARK

                self._clear_cell_shading(cells[3])
                self._set_cell_borders(cells[3], "CCCCCC")

                for cell in cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = FontSizes.SUBTITLE
                            run.font.color.rgb = BrandColors.TEXT_DARK

            # Set column widths
            table.columns[0].width = Inches(1.3)  # Origin/Motivation
            table.columns[1].width = Inches(2.5)  # Activity Observed
            table.columns[2].width = Inches(2.0)  # What to Monitor
            table.columns[3].width = Inches(0.9)  # Source
        else:
            self.doc.add_paragraph("No threat actor activity data available.")

        self.doc.add_paragraph()

    def _add_active_campaigns(self, analysis_result: dict[str, Any]) -> None:
        """Add Active Campaigns section with campaign details table."""
        logger.info("Adding Active Campaigns section")

        h = self.doc.add_heading("Active Campaigns", level=1)
        self._style_heading_1(h)

        # Intro text
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run(
            "Coordinated attack campaigns observed this week from threat intelligence sources. "
            "Campaigns represent sustained, targeted operations by threat actors with specific objectives."
        )
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

        campaigns = analysis_result.get("active_campaigns", [])

        if campaigns:
            # Create campaigns table
            # Columns: Campaign Name | Threat Actor(s) | Objective | Targets | TTPs | Timeline | Source
            table = self.doc.add_table(rows=1, cols=7)

            # Header row: orange background, white text
            headers = ["Campaign Name", "Threat Actor(s)", "Objective", "Targets", "TTPs", "Timeline", "Source"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                para = header_cells[i].paragraphs[0]
                para.runs[0].font.bold = True
                para.runs[0].font.size = FontSizes.SUBTITLE
                para.runs[0].font.color.rgb = BrandColors.WHITE
                self._set_cell_shading(header_cells[i], BrandColors.TABLE_HEADER_BG)
                self._set_cell_borders(header_cells[i], "CCCCCC")

            # Build citation map
            citation_map = self._build_citation_map(analysis_result)

            # Data rows
            for campaign in campaigns:
                row = table.add_row()
                cells = row.cells

                # Column 0: Campaign Name
                campaign_name = campaign.get("campaign_name", "Unknown Campaign")
                cells[0].text = campaign_name
                self._clear_cell_shading(cells[0])
                self._set_cell_borders(cells[0], "CCCCCC")

                # Column 1: Threat Actor(s) - handle list of up to 4, or "Multiple actors"
                actors = campaign.get("threat_actors", [])
                if isinstance(actors, list):
                    if len(actors) > 4:
                        actor_text = "Multiple actors"
                    else:
                        actor_text = ", ".join(actors) if actors else "Multiple actors"
                else:
                    actor_text = str(actors) if actors else "Multiple actors"
                cells[1].text = actor_text
                self._clear_cell_shading(cells[1])
                self._set_cell_borders(cells[1], "CCCCCC")

                # Column 2: Objective
                objective = campaign.get("objective", "")
                cells[2].text = objective
                self._clear_cell_shading(cells[2])
                self._set_cell_borders(cells[2], "CCCCCC")

                # Column 3: Targets
                targets = campaign.get("targets", "")
                cells[3].text = targets
                self._clear_cell_shading(cells[3])
                self._set_cell_borders(cells[3], "CCCCCC")

                # Column 4: TTPs - join list
                ttps = campaign.get("ttps", [])
                if isinstance(ttps, list):
                    ttps_text = ", ".join(ttps[:4])  # Limit to 4 TTPs for readability
                else:
                    ttps_text = str(ttps) if ttps else ""
                cells[4].text = ttps_text
                self._clear_cell_shading(cells[4])
                self._set_cell_borders(cells[4], "CCCCCC")

                # Column 5: Timeline
                timeline = campaign.get("timeline", "Active this week")
                cells[5].text = timeline
                self._clear_cell_shading(cells[5])
                self._set_cell_borders(cells[5], "CCCCCC")

                # Column 6: Source Citations - NORMAL SIZE [#] Source
                cells[6].text = ""
                cite_para = cells[6].paragraphs[0]
                source_citations = campaign.get("sources", [])
                if isinstance(source_citations, list) and source_citations:
                    # Add each source with its number
                    for idx, src in enumerate(source_citations):
                        if idx > 0:
                            cite_para.add_run(", ")

                        # Get citation number - must be a real API source
                        cite_num = citation_map.get(src)

                        if cite_num:
                            # Add [#] Source (normal size, not superscript)
                            cite_run = cite_para.add_run(f"[{cite_num}] {src}")
                            cite_run.font.size = FontSizes.SUBTITLE
                            cite_run.font.color.rgb = BrandColors.TEXT_DARK
                        else:
                            # No valid citation - log warning and show in gray
                            logger.warning(
                                f"Campaign '{campaign.get('campaign_name')}' cites unverifiable source: {src}"
                            )
                            cite_run = cite_para.add_run(f"{src}")
                            cite_run.font.size = FontSizes.SUBTITLE
                            cite_run.font.color.rgb = BrandColors.GRAY_MEDIUM  # Gray to indicate issue
                else:
                    # Default if no source_citations
                    cite_run = cite_para.add_run("Unknown")
                    cite_run.font.size = FontSizes.SUBTITLE
                    cite_run.font.color.rgb = BrandColors.GRAY_MEDIUM

                self._clear_cell_shading(cells[6])
                self._set_cell_borders(cells[6], "CCCCCC")

                # Style all cells
                for cell in cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = FontSizes.SUBTITLE
                            run.font.color.rgb = BrandColors.TEXT_DARK

            # Set column widths for readability
            # Campaign Name: wider, Actors: medium, others: proportional
            table.columns[0].width = Inches(1.3)  # Campaign Name
            table.columns[1].width = Inches(1.0)  # Threat Actor(s)
            table.columns[2].width = Inches(1.3)  # Objective
            table.columns[3].width = Inches(1.1)  # Targets
            table.columns[4].width = Inches(1.0)  # TTPs
            table.columns[5].width = Inches(0.8)  # Timeline
            table.columns[6].width = Inches(0.9)  # Source
        else:
            # No campaigns observed
            no_campaigns = self.doc.add_paragraph()
            no_campaigns_run = no_campaigns.add_run(
                "No named campaigns observed this week. See Sector Threat Activity for general threat actor behavior."
            )
            no_campaigns_run.font.size = FontSizes.BODY
            no_campaigns_run.font.italic = True
            no_campaigns_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        self.doc.add_paragraph()

    def _add_industry_incidents(self, analysis_result: dict[str, Any]) -> None:
        """Add Peer Incidents section - company breaches from OSINT sources."""
        logger.info("Adding Peer Incidents section")

        h = self.doc.add_heading("Peer Incidents", level=1)
        self._style_heading_1(h)

        # Intro text with actionable context
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run(
            "Notable security incidents affecting organizations in similar sectors or technology stacks, "
            "observed from underground intelligence, public breach disclosures, and security news sources. "
        )
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        action_run = intro.add_run(
            "Review attack vectors to identify relevant log sources: authentication logs for Credential Harvesting, "
            "web logs for Remote Code Execution, network logs for C2 Communication."
        )
        action_run.font.size = FontSizes.BODY_SMALL
        action_run.font.bold = True
        action_run.font.color.rgb = BrandColors.TEXT_DARK

        self.doc.add_paragraph()

        # Build citation map for consistent numbering
        citation_map = self._build_citation_map(analysis_result)

        # Build reverse lookup: map AI's original citation numbers to OSINT titles
        osint_sources = analysis_result.get("osint_sources_used", [])
        ai_citation_to_title = {}
        for i, source in enumerate(osint_sources, 1):
            if isinstance(source, dict):
                title = source.get("title", "")
                # AI assigns citation numbers starting from 1
                ai_citation_to_title[i] = title

        # Collect incidents from multiple sources:
        incidents = []

        # 1. Get AI-identified incidents from OSINT (structured output)
        ai_incidents = analysis_result.get("industry_incidents", [])
        if ai_incidents:
            logger.info(f"Found {len(ai_incidents)} incidents from AI analysis")
            incidents.extend(ai_incidents)

        # 2. Add Intel471 breach alerts directly (may not be in AI output)
        intel471_breaches = self._extract_intel471_breaches(analysis_result)
        if intel471_breaches:
            logger.info(f"Adding {len(intel471_breaches)} Intel471 breach alerts")
            incidents.extend(intel471_breaches)

        # 3. Fallback: extract from OSINT if AI didn't provide structured incidents
        if not incidents:
            logger.warning("No industry_incidents from AI or Intel471, falling back to extraction from OSINT titles")
            incidents = self._extract_breach_incidents(osint_sources)

        if incidents:
            # Create table: Company | Incident Type | Date | Source
            table = self.doc.add_table(rows=1, cols=4)

            # Header row: orange background, white text
            headers = ["Organization", "Incident Type", "Date", "Source"]
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                para = header_cells[i].paragraphs[0]
                para.runs[0].font.bold = True
                para.runs[0].font.size = FontSizes.SUBTITLE
                para.runs[0].font.color.rgb = BrandColors.WHITE
                self._set_cell_shading(header_cells[i], BrandColors.TABLE_HEADER_BG)
                self._set_cell_borders(header_cells[i], "CCCCCC")

            # Data rows
            for incident in incidents:
                row = table.add_row()
                cells = row.cells

                # Column 0: Organization name
                org = incident.get("organization", "Unknown")
                cells[0].text = org[:40]  # Truncate long names

                # Column 1: Incident type (using SOC-friendly attack vectors)
                incident_type = incident.get("incident_type", "Breach")
                # Map to VERIS-aligned attack vectors for SOC consumption
                incident_type = self._normalize_incident_type(incident_type)
                cells[1].text = incident_type

                # Column 2: Date
                date = incident.get("date", "Unknown")
                if hasattr(date, "strftime"):
                    date = date.strftime("%Y-%m-%d")
                elif isinstance(date, str) and len(date) > 10:
                    date = date[:10]  # Truncate to YYYY-MM-DD
                cells[2].text = date

                # Column 3: Source with citation number
                source_name = incident.get("source", "OSINT")
                ai_citation_num = incident.get("osint_citation_number")

                # Map AI's citation number to the correct final citation number
                final_citation_num = None
                if ai_citation_num and ai_citation_num in ai_citation_to_title:
                    # Look up the title from AI's citation number
                    title = ai_citation_to_title[ai_citation_num]
                    # Get the final citation number from the map
                    final_citation_num = citation_map.get(title)

                # Check if this is an Intel471 source
                intel471_citation = citation_map.get("Intel471")

                # Check for generic OSINT fallback
                citation_map.get("OSINT")

                # Column 3: Source - NORMAL SIZE [#] SourceName (not superscript in table column)
                cells[3].text = ""
                source_para = cells[3].paragraphs[0]

                if final_citation_num:
                    # Normal size: [#] SourceName
                    cite_run = source_para.add_run(f"[{final_citation_num}] {source_name[:20]}")
                    cite_run.font.size = FontSizes.SUBTITLE
                    cite_run.font.color.rgb = BrandColors.TEXT_DARK
                elif "Intel471" in source_name and intel471_citation:
                    # Normal size: [#] Intel471
                    cite_run = source_para.add_run(f"[{intel471_citation}] Intel471")
                    cite_run.font.size = FontSizes.SUBTITLE
                    cite_run.font.color.rgb = BrandColors.TEXT_DARK
                else:
                    # No valid citation - log warning and show source without number
                    logger.warning(f"Peer incident for {org} has no valid citation source: {source_name}")
                    source_text_run = source_para.add_run(source_name[:25])
                    source_text_run.font.size = FontSizes.SUBTITLE
                    source_text_run.font.color.rgb = BrandColors.GRAY_MEDIUM  # Gray to indicate missing citation

                # Styling for all columns
                for idx in range(4):
                    self._set_cell_borders(cells[idx], "CCCCCC")
                    for para in cells[idx].paragraphs:
                        for run in para.runs:
                            run.font.size = FontSizes.SUBTITLE
                            run.font.color.rgb = BrandColors.TEXT_DARK

            # Set column widths
            table.columns[0].width = Inches(2.0)  # Organization
            table.columns[1].width = Inches(1.5)  # Incident Type
            table.columns[2].width = Inches(1.0)  # Date
            table.columns[3].width = Inches(1.7)  # Source

            # Caption
            caption = self.doc.add_paragraph()
            caption.space_before = Pt(4)
            caption.space_after = Pt(8)
            caption_text = (
                f"Table: {len(incidents)} peer incidents from threat intelligence sources (Intel471, OSINT). "
                "Incidents may indicate threat actor targeting patterns or emerging attack vectors."
            )
            caption_run = caption.add_run(caption_text)
            caption_run.font.size = Pt(7)
            caption_run.font.italic = True
            caption_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            caption_run.font.name = "Arial"
            caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            self.doc.add_paragraph("No significant industry incidents observed this week.")

        self.doc.add_paragraph()

    def _extract_intel471_breaches(self, analysis_result: dict[str, Any]) -> list[dict[str, Any]]:
        """Extract breach incidents from raw Intel471 data.

        Intel471 breach alerts have threat_type='BREACH ALERT' or 'Breach Alert'.
        We need to access the raw data, not the AI's analysis.
        """
        incidents = []

        # The analysis_result doesn't have raw Intel471 data
        # We need to get it from somewhere else - for now, return empty
        # TODO: Pass raw Intel471 data to the report generator

        logger.debug("Intel471 breach extraction not yet implemented - raw data not available to report generator")
        return incidents

    def _extract_breach_incidents(self, osint_sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
        """Extract breach incidents from OSINT sources."""
        incidents = []

        # Keywords that indicate breach/incident articles
        breach_keywords = [
            "breach",
            "ransomware",
            "hack",
            "attack",
            "compromise",
            "data leak",
            "cyberattack",
            "cyber attack",
            "incident",
        ]

        for source in osint_sources:
            title = source.get("title", "").lower()
            url = source.get("url", "")
            date = source.get("date", source.get("published", "Unknown"))

            # Check if this is a breach-related article
            if any(keyword in title for keyword in breach_keywords):
                # Try to extract organization name from title
                org_name = self._extract_organization_from_title(source.get("title", ""))

                # Determine incident type (VERIS-aligned attack vectors)
                incident_type = "Credential Harvesting"  # Default
                if "ransomware" in title:
                    incident_type = "Ransomware"
                elif "ddos" in title or "denial of service" in title:
                    incident_type = "DDoS"
                elif "data leak" in title or "data breach" in title:
                    incident_type = "Data Exfiltration"
                elif "supply chain" in title:
                    incident_type = "Supply Chain Attack"
                elif "phishing" in title:
                    incident_type = "Phishing"
                elif "malware" in title:
                    incident_type = "Malware Infection"
                elif "exploit" in title or "vulnerability" in title:
                    incident_type = "Remote Code Execution"

                # Extract source domain from URL
                source_name = "OSINT"
                if url:
                    try:
                        from urllib.parse import urlparse

                        domain = urlparse(url).netloc
                        # Remove www. prefix
                        if domain.startswith("www."):
                            domain = domain[4:]
                        source_name = domain[:25]  # Truncate long domains
                    except (ValueError, AttributeError):
                        source_name = "OSINT"

                # Format date if it's a datetime
                if hasattr(date, "strftime"):
                    date = date.strftime("%Y-%m-%d")
                elif isinstance(date, str) and len(date) > 10:
                    date = date[:10]  # Truncate to YYYY-MM-DD

                incidents.append(
                    {
                        "organization": org_name,
                        "incident_type": incident_type,
                        "date": date,
                        "source": source_name,
                        "url": url,
                    }
                )

        # Limit to top 15 most recent
        return incidents[:15]

    def _extract_organization_from_title(self, title: str) -> str:
        """Extract organization name from article title."""
        # Common patterns in breach headlines:
        # "Company X suffers data breach"
        # "Ransomware attack hits Organization Y"
        # "Hackers breach Company Z"

        # Remove common prefixes/suffixes
        title = title.replace("'s ", " ").replace("'s ", " ")

        # Split on common verbs/keywords
        split_words = [
            " suffers ",
            " confirms ",
            " discloses ",
            " reports ",
            " hit by ",
            " targeted by ",
            " attacked by ",
            " breached by ",
            " says ",
            " warns ",
            " announces ",
        ]

        for word in split_words:
            if word in title.lower():
                # Take the part before the verb
                org_part = title.lower().split(word)[0]
                # Take last few words as org name
                words = org_part.strip().split()
                if len(words) >= 2:
                    # Take last 2-3 words as organization name
                    org_name = " ".join(words[-3:]) if len(words) >= 3 else " ".join(words[-2:])
                    return org_name.title()[:40]  # Capitalize and truncate

        # If no pattern match, take first N words
        words = title.split()
        if len(words) >= 2:
            return " ".join(words[:3]).title()[:40]

        return "Unknown Organization"

    def _normalize_incident_type(self, incident_type: str) -> str:
        """Normalize incident types to SOC-friendly VERIS-aligned attack vectors."""
        incident_type_lower = incident_type.lower()

        # Map common incident types to VERIS attack vectors
        if "ransomware" in incident_type_lower:
            return "Ransomware"
        elif "credential" in incident_type_lower or "harvesting" in incident_type_lower:
            return "Credential Harvesting"
        elif "phishing" in incident_type_lower:
            return "Phishing"
        elif "malware" in incident_type_lower:
            return "Malware Infection"
        elif "ddos" in incident_type_lower or "denial" in incident_type_lower:
            return "DDoS"
        elif "data leak" in incident_type_lower or "exfiltration" in incident_type_lower:
            return "Data Exfiltration"
        elif "supply chain" in incident_type_lower:
            return "Supply Chain Attack"
        elif "exploit" in incident_type_lower or "rce" in incident_type_lower or "remote code" in incident_type_lower:
            return "Remote Code Execution"
        elif "injection" in incident_type_lower or "sqli" in incident_type_lower:
            return "Code Injection"
        elif "c2" in incident_type_lower or "command and control" in incident_type_lower:
            return "C2 Communication"
        elif "brute" in incident_type_lower or "force" in incident_type_lower:
            return "Brute Force"
        elif "breach" in incident_type_lower:
            return "Unauthorized Access"
        else:
            return incident_type  # Return as-is if no mapping

    def _add_recommended_actions(self, analysis_result: dict[str, Any]) -> None:
        """Add recommended actions section."""
        logger.info("Adding Recommended Actions section")

        h = self.doc.add_heading("Recommended Actions", level=1)
        self._style_heading_1(h)

        recommendations = analysis_result.get("recommendations", [])

        if recommendations:
            for rec in recommendations:
                para = self.doc.add_paragraph(rec, style="List Bullet")
                for run in para.runs:
                    run.font.size = FontSizes.BODY_SMALL

                    # Bold important/urgent recommendations
                    lower_rec = rec.lower()
                    if any(word in lower_rec for word in ["urgent", "critical", "immediate", "persistent"]):
                        run.font.bold = True
        else:
            default_recs = [
                "Review scan results for exposed vulnerabilities; validate asset ownership and remediation timelines",
                "Brief development teams on current threat campaigns",
                "Verify endpoint protection has latest behavioral IOAs enabled",
                "Confirm SIEM is receiving logs from affected systems",
            ]
            for rec in default_recs:
                para = self.doc.add_paragraph(rec, style="List Bullet")
                for run in para.runs:
                    run.font.size = FontSizes.BODY_SMALL

        self.doc.add_paragraph()

    def _build_citation_map(self, analysis_result: dict[str, Any]) -> dict[str, int]:
        """Build a mapping of OSINT source titles/URLs to their final citation numbers."""
        citation_map = {}
        citation_counter = 1

        # Collect all API sources actually cited in analysis
        api_sources_used = set()

        # Check CVE analysis for source citations
        cve_analysis = analysis_result.get("cve_analysis", [])
        for cve in cve_analysis:
            if isinstance(cve, dict):
                sources = cve.get("source_citations", [])
                if isinstance(sources, list):
                    api_sources_used.update(sources)

        # Check APT activity for source citations
        apt_activity = analysis_result.get("apt_activity", [])
        for apt in apt_activity:
            if isinstance(apt, dict):
                sources = apt.get("source_citations", [])
                if isinstance(sources, list):
                    api_sources_used.update(sources)

        # Build primary sources list based on what was actually used
        primary_sources = []

        # Always include NVD (CVE database)
        primary_sources.append(("NVD", "NIST National Vulnerability Database (NVD)"))

        # Include CISA KEV if any CVEs are in KEV
        if any(cve.get("in_cisa_kev") for cve in cve_analysis if isinstance(cve, dict)):
            primary_sources.append(("CISA KEV", "CISA Known Exploited Vulnerabilities (KEV) Catalog"))

        # Include Intel471 if cited
        if "Intel471" in api_sources_used:
            primary_sources.append(("Intel471", "Intel471 Titan threat intelligence platform"))

        # Include CrowdStrike if cited
        if "CrowdStrike" in api_sources_used:
            primary_sources.append(("CrowdStrike", "CrowdStrike Falcon Intelligence"))

        # Include ThreatQ if cited (usually disabled)
        if "ThreatQ" in api_sources_used:
            primary_sources.append(("ThreatQ", "ThreatQ threat intelligence management platform"))

        # Map API sources
        for short_name, _full_name in primary_sources:
            citation_map[short_name] = citation_counter
            citation_counter += 1

        # Map OSINT sources - create mapping from title/URL to citation number
        osint_sources = analysis_result.get("osint_sources_used", [])
        for source in osint_sources:
            if isinstance(source, dict):
                title = source.get("title", "")
                url = source.get("url", "")
                if title:
                    # Map by title (primary key)
                    citation_map[title] = citation_counter
                if url:
                    # Also map by URL as fallback
                    citation_map[url] = citation_counter
                citation_counter += 1

        return citation_map

    def _add_resources_section(self, analysis_result: dict[str, Any]) -> None:
        """Add Resources section listing threat intelligence sources that actually provided data."""
        logger.info("Adding Resources section")

        h = self.doc.add_heading("Resources & Intelligence Sources", level=1)
        self._style_heading_1(h)

        # Intro text
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run("This report was compiled using the following intelligence sources:")
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        # Build citation map (used by Industry Incidents table too)
        citation_map = self._build_citation_map(analysis_result)

        # Collect all API sources actually cited in analysis
        api_sources_used = set()

        # Check CVE analysis for source citations
        cve_analysis = analysis_result.get("cve_analysis", [])
        for cve in cve_analysis:
            if isinstance(cve, dict):
                sources = cve.get("source_citations", [])
                if isinstance(sources, list):
                    api_sources_used.update(sources)

        # Check APT activity for source citations
        apt_activity = analysis_result.get("apt_activity", [])
        for apt in apt_activity:
            if isinstance(apt, dict):
                sources = apt.get("source_citations", [])
                if isinstance(sources, list):
                    api_sources_used.update(sources)

        # Build primary sources list based on what was actually used
        primary_sources = []

        # Always include NVD (CVE database)
        primary_sources.append(("NVD", "NIST National Vulnerability Database (NVD)"))

        # Include CISA KEV if any CVEs are in KEV
        if any(cve.get("in_cisa_kev") for cve in cve_analysis if isinstance(cve, dict)):
            primary_sources.append(("CISA KEV", "CISA Known Exploited Vulnerabilities (KEV) Catalog"))

        # Include Intel471 if cited
        if "Intel471" in api_sources_used:
            primary_sources.append(("Intel471", "Intel471 Titan threat intelligence platform"))

        # Include CrowdStrike if cited
        if "CrowdStrike" in api_sources_used:
            primary_sources.append(("CrowdStrike", "CrowdStrike Falcon Intelligence"))

        # Include ThreatQ if cited (usually disabled)
        if "ThreatQ" in api_sources_used:
            primary_sources.append(("ThreatQ", "ThreatQ threat intelligence management platform"))

        # Add numbered citations for API sources
        for short_name, full_name in primary_sources:
            para = self.doc.add_paragraph(style="List Bullet")

            # Get citation number from map - NORMAL SIZE BOLD (not superscript)
            citation_num = citation_map.get(short_name, "?")
            cite_run = para.add_run(f"[{citation_num}] ")
            cite_run.font.bold = True
            cite_run.font.size = FontSizes.BODY_SMALL
            cite_run.font.color.rgb = BrandColors.TEXT_DARK

            # Add source name
            source_run = para.add_run(full_name)
            source_run.font.size = FontSizes.BODY_SMALL

        # OSINT sources (if any were listed by the AI)
        osint_sources = analysis_result.get("osint_sources_used", [])
        if osint_sources:
            self.doc.add_paragraph()
            osint_heading = self.doc.add_paragraph()
            osint_heading_run = osint_heading.add_run("Open Source Intelligence (OSINT) Sources:")
            osint_heading_run.font.size = FontSizes.BODY_SMALL
            osint_heading_run.font.bold = True
            osint_heading_run.font.color.rgb = BrandColors.TEXT_DARK

            for source in osint_sources:
                # Handle both old string format and new dict format
                if isinstance(source, dict):
                    title = source.get("title", "")
                    url = source.get("url", "")
                    relevance = source.get("relevance", "")

                    if title and url:
                        # Create paragraph with bullet
                        para = self.doc.add_paragraph(style="List Bullet")

                        # Get citation number from map - NORMAL SIZE BOLD (not superscript)
                        citation_num = citation_map.get(title, citation_map.get(url, "?"))
                        cite_run = para.add_run(f"[{citation_num}] ")
                        cite_run.font.bold = True
                        cite_run.font.size = FontSizes.FOOTNOTE
                        cite_run.font.color.rgb = BrandColors.TEXT_DARK

                        # Add hyperlink for the title
                        self._add_hyperlink(para, title, url)

                        # Add relevance note if provided
                        if relevance:
                            rel_run = para.add_run(f" - {relevance}")
                            rel_run.font.size = FontSizes.FOOTNOTE
                            rel_run.font.color.rgb = BrandColors.GRAY_MEDIUM
                            rel_run.font.italic = True

                        # Set font size for the whole paragraph
                        for run in para.runs:
                            if run.font.size is None:
                                run.font.size = FontSizes.BODY_SMALL
                    elif title:
                        # No URL, just show title
                        para = self.doc.add_paragraph(title, style="List Bullet")
                        for run in para.runs:
                            run.font.size = FontSizes.BODY_SMALL
                elif isinstance(source, str) and source.strip():
                    # Legacy string format
                    para = self.doc.add_paragraph(source, style="List Bullet")
                    for run in para.runs:
                        run.font.size = FontSizes.BODY_SMALL

        self.doc.add_paragraph()

    def _add_ai_disclaimer(self) -> None:
        """Add AI-generated disclaimer box with intelligence sources."""
        logger.info("Adding AI disclaimer box")

        # Add a small gap before the disclaimer
        self.doc.add_paragraph()

        # Create a single-cell table for the box
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Light blue/gray background for the info box
        self._set_cell_shading(cell, "E8F4F8")  # Light blue-gray
        self._set_cell_borders(cell, "B0D4E3", "8")  # Slightly darker blue border

        # Clear default paragraph
        cell.paragraphs[0].clear()

        # Title in the box
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run("AI-Generated Intelligence Report")
        title_run.font.size = FontSizes.BODY_SMALL
        title_run.font.bold = True
        title_run.font.color.rgb = BrandColors.TEXT_DARK
        title_run.font.name = "Arial"
        title_para.paragraph_format.space_after = Pt(4)

        # Body text
        body_para = cell.add_paragraph()
        body_text = (
            "This report was generated using AI-powered analysis to curate and correlate threat intelligence. "
            "Vulnerability and threat actor data are sourced from the intelligence platforms listed in the Resources section above. "
            "The analysis identifies threats actively exploited in the wild or targeting the manufacturing, biotech, "
            "medical device, clinical, and genomics industries. All CVEs listed have been detected in our environment "
            "through automated scanning."
        )
        body_run = body_para.add_run(body_text)
        body_run.font.size = Pt(8)
        body_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        body_run.font.name = "Arial"
        body_para.paragraph_format.space_before = Pt(0)
        body_para.paragraph_format.space_after = Pt(6)
        body_para.paragraph_format.line_spacing = 1.15

        self.doc.add_paragraph()

    def _add_footer(self) -> None:
        """Add footer with contact info and ServiceNow hyperlink."""
        logger.info("Adding footer")

        # Contact info
        contact = self.doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add text: "Questions or suspicious activity: secops@illumina.com | "
        contact_run = contact.add_run("Questions or suspicious activity: secops@illumina.com | ")
        contact_run.font.size = FontSizes.BODY_SMALL
        contact_run.font.bold = True
        contact_run.font.color.rgb = BrandColors.TEXT_DARK

        # Add ServiceNow as hyperlink
        servicenow_url = "https://illumina.service-now.com/esp?id=sc_cat_item&sys_id=2498318b4fee53c0f628d0af0310c75d"
        self._add_hyperlink(contact, "ServiceNow", servicenow_url)

        # Apply font styling to hyperlink
        for run in contact.runs:
            if run != contact_run:  # Style the hyperlink run
                run.font.size = FontSizes.BODY_SMALL
                run.font.bold = True

        self.doc.add_paragraph()
