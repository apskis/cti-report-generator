"""
Quarterly Strategic CTI Report Generator.

Generates quarterly strategic threat intelligence briefs for leadership.
"""

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.core.config import customer_profile
from src.core.reporting_period import ReportingPeriod, current_quarter, make_period
from src.reports.base import BaseReportGenerator, BrandColors, FontSizes
from src.reports.registry import register_report_generator

logger = logging.getLogger(__name__)


class RiskLevel:
    """Risk level constants for quarterly assessments."""

    HIGH = "HIGH"
    MEDIUM = "MEDIUM"
    LOW = "LOW"
    UNCHANGED = "Unchanged"
    INCREASED = "↑"
    DECREASED = "↓"


@register_report_generator("quarterly")
class QuarterlyReportGenerator(BaseReportGenerator):
    """
    Quarterly Strategic CTI Report Generator.

    Generates leadership-focused strategic threat briefs matching
    the CTI_Quarterly_Strategic_Report_Example.docx format.

    Template structure:
    - Report ID header (CTI-QTR-YYYY-QN)
    - Title: "Cyber Threat Intelligence" / "Quarterly Strategic Brief"
    - Quarter date range
    - Executive Summary (strategic overview)
    - Quarterly Risk Assessment (risk metric cards)
    - Industry Breach Landscape (incidents table)
    - Geopolitical Threat Landscape (by nation-state)
    - Looking Ahead (next quarter outlook)
    - Recommendations for Leadership
    - Footer with sources
    """

    @property
    def report_type(self) -> str:
        return "quarterly"

    @property
    def filename_prefix(self) -> str:
        return "CTI_Quarterly_Strategic_Brief"

    def get_filename(self, report_week_start: datetime = None) -> str:
        """Name quarterly reports by their calendar quarter, not the ISO week.

        A quarterly brief covers a whole quarter, so ``Week30`` is meaningless in the
        name. Use the pinned reporting period to produce e.g.
        ``CTI_Quarterly_Strategic_Brief_Q2_2026.docx``. Falls back to the base
        week-based name only if no period was resolved (should not happen in practice).
        """
        period = getattr(self, "reporting_period", None)
        if period is None:
            return super().get_filename(report_week_start=report_week_start)
        mock_suffix = "_MOCK" if self.use_mock_data else ""
        return f"{self.filename_prefix}_Q{period.quarter}_{period.year}{mock_suffix}.docx"

    def generate(self, analysis_result: dict[str, Any]) -> Document:
        """
        Generate the quarterly strategic report document.

        Args:
            analysis_result: Dictionary containing:
                - executive_summary: str (strategic overview)
                - risk_assessment: dict with risk levels
                - breach_landscape: dict with incident statistics
                - incidents_by_type: list of incident type breakdowns
                - geopolitical_threats: dict by country (china, russia, north_korea, etc.)
                - looking_ahead: dict with outlook, initiatives, watch_items
                - recommendations: list of strategic recommendations

        Returns:
            Document object
        """
        try:
            logger.info("Generating Quarterly Strategic CTI Report")
            self.doc = Document()

            # Set default font to Arial for the entire document
            self._set_default_font("Arial")

            # Print-style: white page with dark body text so the document displays
            # correctly in both light and dark mode (parity with the weekly report).
            self._set_document_background(BrandColors.PAGE_WHITE)
            normal_style = self.doc.styles["Normal"]
            normal_style.font.color.rgb = BrandColors.TEXT_DARK

            # Configure page settings (margins, header/footer distances, paragraph spacing)
            self._configure_page_settings()

            # Calculate quarter info
            self._calculate_quarter_info()

            # Make the reporting period authoritative over any quarter labels the AI
            # guessed, so the cards/table never disagree with the chosen quarter.
            self._apply_reporting_period_labels(analysis_result)

            # Ground the four breach stat cards in real, date-stamped incidents (VCDB/HHS/
            # HIBP) for this quarter when a dataset was supplied.
            self._apply_breach_dataset_grounding(analysis_result)

            # Fallback: if the Ransomware card is still N/A, fill it from the observed
            # incident-type count BEFORE computing prior-quarter percentages.
            self._derive_ransomware_count(analysis_result)

            # If a real prior-quarter baseline exists in history, use its stat values for
            # the cards' "prior" column and quarter-over-quarter percentages.
            self._apply_prior_quarter_stats(analysis_result)

            # Add sections in order
            self._add_header()
            self._add_executive_summary(analysis_result)
            self._add_risk_assessment(analysis_result)
            self._add_geopolitical_landscape(analysis_result)
            self._add_breach_landscape(analysis_result)
            self._add_looking_ahead(analysis_result)
            self._add_recommendations(analysis_result)
            # Sources/References: the [N] markers here LABEL the reference entries, so keep them at
            # normal size — superscript is only for inline citations in the body. Record this
            # section's paragraphs so the citation pass below skips them.
            _pids_before = {id(p._p) for p in self.doc.paragraphs}
            self._add_sources(analysis_result)
            self._citation_skip_p_ids = {id(p._p) for p in self.doc.paragraphs} - _pids_before
            self._add_footer()

            # Document-wide pass: render every inline citation marker ([1], [3][4]) as a raised
            # superscript, wherever it appears (summary, tables, cards, incidents, etc.) EXCEPT the
            # Sources list recorded above.
            self._subscript_all_citations()

            logger.info("Quarterly Strategic CTI Report generated successfully")
            return self.doc

        except Exception as e:
            logger.error(f"Error generating quarterly report: {str(e)}", exc_info=True)
            raise

    def _set_cell_left_border(self, cell, color_hex: str, size: str = "4") -> None:
        """
        Apply a left border to a table cell (used for accent borders on geo cards and stat cards).

        Args:
            cell: The table cell to apply left border to
            color_hex: Border color in hex format (e.g., "E65100" for orange)
            size: Border size as string (e.g., "12" for thicker, "4" for thin)
        """
        tc_pr = cell._element.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        left_border = tc_borders.find(qn("w:left"))
        if left_border is None:
            left_border = OxmlElement("w:left")
            tc_borders.append(left_border)

        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), size)
        left_border.set(qn("w:space"), "0")
        left_border.set(qn("w:color"), color_hex)

    def set_reporting_period(self, period: ReportingPeriod) -> None:
        """Pin the report to a specific calendar quarter (year + quarter).

        When set, every quarter/year label, the reporting window, and the historical
        key derive from this single choice instead of being guessed from today's date.
        """
        self.reporting_period = period

    def set_breach_dataset(self, records: list[dict[str, Any]]) -> None:
        """Provide date-stamped breach incidents (VCDB/HHS/HIBP) used to ground the stat cards."""
        self.breach_dataset = records or []

    def _apply_breach_dataset_grounding(self, analysis_result: dict[str, Any]) -> None:
        """Ground the breach stat cards in real, date-stamped incidents for the period.

        Counts and records-exposed come straight from the deduplicated dataset incidents
        that fall inside the reporting quarter; the dollar figure is an explicit estimate
        (records x IBM per-record cost). Only overrides when the dataset actually has
        incidents for this quarter — otherwise the AI's values (or honest N/A) stand.
        """
        breach = analysis_result.get("breach_landscape")
        breach = breach if isinstance(breach, dict) else None

        def _mark_ungrounded(reason: str) -> None:
            # Don't let unvalidated AI breach numbers read as authoritative fact. Stamp an honest
            # marker; a compliant model already emits N/A for feed-underivable figures, so this
            # just labels the estimate rather than blanking the section (C1).
            if breach is not None and breach.get("stat_cards"):
                breach["stat_methodology"] = (
                    "Breach figures below are AI estimates from the threat feeds and are NOT "
                    f"grounded in a date-stamped breach dataset this quarter ({reason}). Treat "
                    "counts and impact as indicative, not authoritative."
                )

        # Ledger stash: what the quarter's metrics were computed from, persisted by
        # _save_current_risk_assessment so next quarter reads a real prior value + basis.
        self._last_grounded_metrics = None
        self._last_grounding_mode = "ungrounded"
        self._last_dataset_incidents = 0

        records = getattr(self, "breach_dataset", None)
        period = getattr(self, "reporting_period", None)
        if not records or period is None:
            _mark_ungrounded("no breach dataset available")
            return
        from src.core.breach_metrics import (
            apply_metrics_to_stat_cards,
            build_incidents_by_type,
            compute_breach_metrics,
            scope_breaches,
        )
        from src.core.config import collector_config

        scoped = scope_breaches(records, period.start, period.end)
        self._last_dataset_incidents = len(scoped)
        metrics = compute_breach_metrics(
            scoped, default_cost_per_breach_usd=collector_config.breach_cost_per_breach_usd
        )
        mode = apply_metrics_to_stat_cards(analysis_result, metrics)
        if mode == "none":
            self._last_grounding_mode = "none"
            _mark_ungrounded("dataset had no incidents in the reporting window")
            return

        # Grounded (full/enrich): record the canonical metrics + mode for the ledger.
        self._last_grounded_metrics = metrics
        self._last_grounding_mode = mode

        n = metrics["total_incidents"]
        logger.info(
            f"Grounded breach stat cards ({mode}) from {n} dataset incidents "
            f"(records: {metrics['records_exposed']:,}, est. impact: ${metrics['est_impact_millions']}M)"
        )
        if breach is None:
            return

        if mode == "enrich":
            # Dataset lags the live feeds — kept live counts; Est. Impact rescaled to the live
            # count. Records Exposed is NOT grounded in this mode (dataset records would
            # undercount, the live feed carries no records count) -> force N/A rather than leave
            # an unvalidated AI number rendering as fact.
            for card in breach.get("stat_cards") or []:
                if isinstance(card, dict) and "record" in str(card.get("label", "")).lower():
                    card["value"] = "N/A"
            logger.info(
                f"Dataset had fewer incidents ({n}) than the live feeds already found; "
                "kept live counts, rescaled Est. Impact, and set Records Exposed to N/A (dataset lag)."
            )
        else:
            # Authoritative dataset -> ground the named incident examples from it too.
            grounded_types = build_incidents_by_type(scoped)
            if grounded_types:
                breach["incidents_by_type"] = grounded_types

        breach["stat_methodology"] = (
            "Est. Total Impact estimated per incident as the sector-weighted average total "
            "cost per breach (IBM Cost of a Data Breach; e.g. healthcare ~$9.8M, "
            "manufacturing ~$4.7M), summed over the quarter's incidents. Counts and records "
            "exposed are from date-stamped industry breach disclosures (VCDB/HHS)."
        )

    def _calculate_quarter_info(self) -> None:
        """Establish the reporting period as an EXACT calendar quarter.

        Uses the explicitly pinned period (``set_reporting_period``) when present;
        otherwise defaults to the calendar quarter that contains the report date. The
        window is the quarter's real start/end dates, not a rolling 90-day span, so the
        label and the covered data always agree.
        """
        period = getattr(self, "reporting_period", None)
        if period is None:
            year, quarter = current_quarter(self.created_at.date())
            period = make_period(year, quarter)
            self.reporting_period = period

        self.quarter = period.quarter
        self.period_start = datetime.combine(period.start, datetime.min.time())
        self.period_end = datetime.combine(period.end, datetime.min.time())
        self.lookback_days = (period.end - period.start).days

    def _get_year(self) -> int:
        """Year of the pinned reporting period (falls back to the report date's year)."""
        period = getattr(self, "reporting_period", None)
        return period.year if period is not None else super()._get_year()

    def _apply_reporting_period_labels(self, analysis_result: dict[str, Any]) -> None:
        """Overwrite AI-guessed quarter labels with the authoritative reporting period.

        The model often labels the breach cards/table with a quarter it inferred from the
        data (which drifts from the chosen quarter). Force the current/prior-quarter labels
        everywhere they render so the report is internally consistent with the period the
        user selected.
        """
        period = getattr(self, "reporting_period", None)
        if period is None:
            return
        current_label = period.label
        prior_label = period.prior.label

        breach = analysis_result.get("breach_landscape")
        if isinstance(breach, dict):
            breach["current_quarter_label"] = current_label
            breach["prior_quarter_label"] = prior_label
            for card in breach.get("stat_cards") or []:
                if isinstance(card, dict):
                    card["prior_label"] = prior_label

    def persist_baseline_from_analysis(self, analysis_result: dict[str, Any], period: ReportingPeriod) -> None:
        """Persist a quarter's risk levels + breach stats to history WITHOUT rendering a doc.

        Used by prior-quarter backfill: it seeds ``history[period.key]`` from a scoped
        analysis so a later report for the following quarter finds a real baseline to
        compare against (trend arrows + stat-card prior column) instead of showing N/A.
        """
        self.set_reporting_period(period)
        self._calculate_quarter_info()
        # Ground the stat cards the same way the rendered report does, so the stored
        # baseline matches what the next quarter compares against.
        self._apply_breach_dataset_grounding(analysis_result)
        self._derive_ransomware_count(analysis_result)
        risk = analysis_result.get("risk_assessment") or {}
        self._save_current_risk_assessment(
            {
                "nation_state": str(risk.get("nation_state") or "MEDIUM").upper(),
                "ransomware": str(risk.get("ransomware") or "MEDIUM").upper(),
                "supply_chain": str(risk.get("supply_chain") or "MEDIUM").upper(),
                "insider": str(risk.get("insider") or "LOW").upper(),
            },
            breach_stats=self._extract_breach_stats(analysis_result),
        )

    def _get_historical_file_path(self) -> Path:
        """Get the path to the historical data JSON file.

        The directory is overridable via ``QUARTERLY_HISTORY_DIR`` so tests/CI can
        redirect the write out of the repo tree instead of the generator writing
        ``data/historical/`` into the working directory as a render side effect.
        """
        import os

        data_dir = Path(os.environ.get("QUARTERLY_HISTORY_DIR", "data/historical"))
        data_dir.mkdir(parents=True, exist_ok=True)
        return data_dir / "quarterly_risk_history.json"

    def set_history_store(self, store: Any) -> None:
        """Persist the quarter ledger to a durable store (``load``/``save``) instead of the local
        file. The deployed Function injects a blob-backed store so QoQ history survives the
        ephemeral container; local runs leave this unset and use the JSON file."""
        self.history_store = store

    def _load_historical_data(self) -> dict[str, Any]:
        """Load the quarter ledger from the injected store, or the local JSON file."""
        store = getattr(self, "history_store", None)
        if store is not None:
            return store.load()
        file_path = self._get_historical_file_path()
        if not file_path.exists():
            return {}

        try:
            with open(file_path) as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Failed to load historical data: {e}")
            return {}

    def _save_historical_data(self, history: dict[str, Any]) -> None:
        """Save the quarter ledger to the injected store, or the local JSON file."""
        store = getattr(self, "history_store", None)
        if store is not None:
            store.save(history)
            return
        file_path = self._get_historical_file_path()
        try:
            with open(file_path, "w") as f:
                json.dump(history, f, indent=2)
            logger.info(f"Saved historical data to {file_path}")
        except Exception as e:
            logger.error(f"Failed to save historical data: {e}")

    def _get_quarter_key(self, year: int, quarter: int) -> str:
        """Get a string key for a specific quarter."""
        return f"{year}-Q{quarter}"

    def _calculate_previous_quarter(self, year: int, quarter: int) -> tuple:
        """Calculate the previous quarter's year and number."""
        if quarter == 1:
            return (year - 1, 4)
        else:
            return (year, quarter - 1)

    def _save_current_risk_assessment(
        self, risk_data: dict[str, Any], breach_stats: dict[str, str] | None = None
    ) -> None:
        """Save the current quarter's risk assessment (and breach stats) to history.

        The 4 risk levels feed next quarter's threat-level trend arrows; ``breach_stats``
        (the stat-card values keyed by label) feed next quarter's stat-card "prior" column
        and quarter-over-quarter percentages, so the comparison is grounded in a real prior
        quarter instead of showing N/A.
        """
        history = self._load_historical_data()

        year = self._get_year()
        quarter_key = self._get_quarter_key(year, self.quarter)

        # Store the risk assessment for this quarter
        entry = {
            "timestamp": self.created_at.isoformat(),
            "year": year,
            "quarter": self.quarter,
            "nation_state": risk_data.get("nation_state", "MEDIUM"),
            "ransomware": risk_data.get("ransomware", "MEDIUM"),
            "supply_chain": risk_data.get("supply_chain", "MEDIUM"),
            "insider": risk_data.get("insider", "LOW"),
        }
        if breach_stats:
            entry["breach_stats"] = breach_stats
        # Ledger: canonical grounded metrics (stable keys) + the basis they were computed on, so
        # next quarter's QoQ reads real numbers and can tell a trend from a basis change.
        from src.core import quarter_ledger

        metrics_block = quarter_ledger.metrics_block(getattr(self, "_last_grounded_metrics", None))
        if metrics_block:
            entry["metrics"] = metrics_block
        entry["basis"] = quarter_ledger.basis_block(
            getattr(self, "_last_grounding_mode", "ungrounded"),
            getattr(self, "_last_dataset_incidents", 0),
            self.created_at.isoformat(),
        )
        history[quarter_key] = entry

        self._save_historical_data(history)
        logger.info(f"Saved risk assessment for {quarter_key}")

    @staticmethod
    def _extract_breach_stats(analysis_result: dict[str, Any]) -> dict[str, str]:
        """Pull the breach stat-card current values, keyed by label, for history storage."""
        breach = analysis_result.get("breach_landscape") or {}
        stats: dict[str, str] = {}
        for card in breach.get("stat_cards") or []:
            if isinstance(card, dict):
                label = str(card.get("label", "")).strip()
                if label:
                    stats[label] = str(card.get("value", "")).strip()
        return stats

    @staticmethod
    def _parse_stat_number(value: str) -> float | None:
        """Parse a stat value like ``"20"``, ``"1,000"``, ``"$1.2M"``, ``"5.0M"`` to a float.

        Returns ``None`` for non-numeric values (e.g. ``"N/A"``) so callers can skip a
        percentage they cannot compute honestly.
        """
        if value is None:
            return None
        s = str(value).strip().upper()
        if s in ("", "N/A", "NA", "—", "-"):
            return None
        s = s.replace(",", "").replace("$", "").strip()
        multiplier = 1.0
        if s.endswith("M"):
            multiplier, s = 1.0, s[:-1]  # keep the magnitude as-is (both sides in "M")
        elif s.endswith("B"):
            multiplier, s = 1000.0, s[:-1]
        try:
            return float(s) * multiplier
        except ValueError:
            return None

    def _apply_prior_quarter_stats(self, analysis_result: dict[str, Any]) -> None:
        """Overwrite stat-card prior values/percentages from the stored prior quarter.

        When a real prior-quarter baseline exists in history, its stat values become the
        cards' ``prior_value`` and drive ``change_pct`` — so "Q1 2026: N/A" becomes a
        genuine number and the delta is computed, not guessed. Cards whose prior value is
        missing or non-numeric keep an honest ``N/A`` with no fabricated sign.
        """
        breach = analysis_result.get("breach_landscape")
        if not isinstance(breach, dict):
            return
        year = self._get_year()
        prev_year, prev_quarter = self._calculate_previous_quarter(year, self.quarter)
        prev_key = self._get_quarter_key(prev_year, prev_quarter)
        from src.core import quarter_ledger

        prev_entry = self._load_historical_data().get(prev_key) or {}
        has_prior = bool(prev_entry.get("breach_stats") or prev_entry.get("metrics"))

        na_tokens = {"", "n/a", "na", "none", "null"}
        for card in breach.get("stat_cards") or []:
            if not isinstance(card, dict):
                continue
            label = str(card.get("label", "")).strip()
            # Resilient lookup: label-keyed display value first, then the canonical metrics block
            # (so a relabeled or metrics-only prior entry still yields a prior value).
            prior_value = quarter_ledger.prior_value_for_label(prev_entry, label)
            if prior_value not in (None, ""):
                card["prior_value"] = prior_value
                card["change_pct"] = self._compute_change_pct(card.get("value", ""), prior_value)
            elif str(card.get("prior_value", "")).strip().lower() in na_tokens:
                # No stored baseline AND no real prior value behind the delta -> force N/A, so a
                # model-invented "(+25%)" with nothing behind it cannot render as fact (C2). A
                # legitimate delta always carries a real prior_value and is left as-is.
                card["prior_value"] = "N/A"
                card["change_pct"] = "N/A"
        if has_prior:
            logger.info(f"Applied prior-quarter ({prev_key}) breach stats to stat cards")

    @classmethod
    def _compute_change_pct(cls, current: str, prior: str) -> str:
        """Honest quarter-over-quarter change label from a current + prior stat value.

        - current is non-numeric (N/A) -> "N/A" (can't compute)
        - prior is non-numeric (no baseline) -> "N/A"
        - both zero -> "0%"
        - prior zero, current > 0 -> "New" (a percentage from zero is undefined/misleading)
        - otherwise -> signed percentage, e.g. "+150%" / "-40%"
        """
        curr_n = cls._parse_stat_number(current)
        prev_n = cls._parse_stat_number(prior)
        if curr_n is None or prev_n is None:
            return "N/A"
        if prev_n == 0:
            return "0%" if curr_n == 0 else "New"
        change = ((curr_n - prev_n) / prev_n) * 100
        if int(change) == 0:
            return "0%"
        return f"+{int(change)}%" if change > 0 else f"{int(change)}%"

    @staticmethod
    def _derive_ransomware_count(analysis_result: dict[str, Any]) -> None:
        """Fill the Ransomware stat card from the observed incident-type counts.

        The model sometimes leaves the Ransomware stat card as N/A even though it broke
        out a "Ransomware" row in incidents_by_type. That row is a real count grounded in
        the breach data, so use it as the card's value — turning an N/A into a genuine
        number (and enabling a real quarter-over-quarter percentage). This does not
        fabricate: it reuses the analysis's own observed count. If no ransomware incident
        type was observed, the card is left untouched.
        """
        breach = analysis_result.get("breach_landscape")
        if not isinstance(breach, dict):
            return
        incidents = breach.get("incidents_by_type") or analysis_result.get("incidents_by_type") or []
        total = 0
        found = False
        for row in incidents:
            if isinstance(row, dict) and "ransom" in str(row.get("type", "")).lower():
                try:
                    total += int(row.get("current_count", 0) or 0)
                    found = True
                except (ValueError, TypeError):
                    continue
        if not found:
            return
        for card in breach.get("stat_cards") or []:
            if isinstance(card, dict) and "ransomware" in str(card.get("label", "")).lower():
                # Fallback only — don't override a value already grounded from the dataset.
                if QuarterlyReportGenerator._parse_stat_number(card.get("value", "")) is None:
                    card["value"] = str(total)
                return

    def _compare_with_previous_quarter(self, current_risk: str, previous_risk: str) -> str:
        """Compare current risk level with previous quarter and return trend indicator."""
        risk_values = {"LOW": 1, "MEDIUM": 2, "HIGH": 3}

        current_val = risk_values.get(current_risk.upper(), 2)
        previous_val = risk_values.get(previous_risk.upper(), 2)

        if current_val > previous_val:
            return "↑"
        elif current_val < previous_val:
            return "↓"
        else:
            return "Unchanged"

    def _get_previous_quarter(self) -> str:
        """Get the previous quarter string (e.g., 'Q4 2025')."""
        if self.quarter == 1:
            return f"Q4 {self._get_year() - 1}"
        return f"Q{self.quarter - 1} {self._get_year()}"

    def _add_header(self) -> None:
        """Add report header with banner image, ID, title, and date range."""
        year = self._get_year()

        # Add banner image at the top
        self._add_banner_header()

        # Report ID (e.g., CTI-QTR-2026-Q1) - positioned at top-right
        report_id = f"CTI-QTR-{year}-Q{self.quarter}"
        id_para = self.doc.add_paragraph()
        id_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        id_run = id_para.add_run(report_id)
        id_run.font.name = "Arial"
        id_run.font.size = FontSizes.SUBTITLE
        id_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        # Main title - Quarterly Strategic Brief (styled like Book Title, orange, 20pt, centered)
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run("Quarterly Strategic Brief")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(20)  # Font size 20pt
        title_run.font.bold = True
        title_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned
        # Reduce spacing after title
        title_para.paragraph_format.space_after = Pt(0)

        # Subtitle - Reporting period based on actual lookback
        date_range = (
            f"{self.lookback_days}-Day Lookback | "
            f"{self.period_start.strftime('%B %d')} to {self.period_end.strftime('%B %d, %Y')}"
        )

        subtitle_para = self.doc.add_paragraph(date_range, style="Subtitle")
        for run in subtitle_para.runs:
            run.font.name = "Arial"
            run.font.color.rgb = BrandColors.GRAY_DARK
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-aligned
        # Reduce spacing after subtitle
        subtitle_para.paragraph_format.space_after = Pt(0)

        # Spacer after cover page
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    def _add_executive_summary(self, analysis_result: dict[str, Any]) -> None:
        """Add executive summary section."""
        logger.info("Adding Executive Summary section")

        # Executive Summary heading - Heading 1
        summary_heading = self.doc.add_heading("Executive Summary", level=1)
        for run in summary_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        # Add space after heading
        summary_heading.paragraph_format.space_after = Pt(6)

        # Get summary paragraphs
        summary = analysis_result.get("executive_summary", "")
        # The AI may emit a list of paragraphs (or a non-string); coerce so .split/.strip below
        # never crash. A list is joined into paragraphs; anything else becomes str().
        if isinstance(summary, list):
            summary = "\n\n".join(str(p) for p in summary if p)
        elif not isinstance(summary, str):
            summary = str(summary or "")
        if not summary:
            summary = self._generate_default_executive_summary(analysis_result)

        # Split into paragraphs if it's a long string
        paragraphs = summary.split("\n\n") if "\n\n" in summary else [summary]

        for para_text in paragraphs:
            if para_text.strip():
                para = self.doc.add_paragraph(para_text.strip())
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = FontSizes.BODY

        # Spacer after executive summary
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    def _generate_default_executive_summary(self, analysis_result: dict[str, Any]) -> str:
        """Generate a default executive summary from available data."""
        stats = analysis_result.get("breach_landscape", {})
        total_incidents = stats.get("total_incidents", 0)
        # geopolitical_threats is a list of per-country entries in the current format;
        # tolerate the legacy dict form ({"actors": [...]}) without crashing.
        geo = analysis_result.get("geopolitical_threats", [])
        apt_groups = len(geo) if isinstance(geo, list) else len(geo.get("actors", []))

        return f"""The threat landscape for the genomics, life sciences, and precision manufacturing sectors \
remained elevated throughout Q{self.quarter} {self._get_year()}, with {total_incidents} publicly disclosed \
breaches affecting peer organizations in the industry.

No direct threats to the organization were identified this quarter; however, the threat actors, techniques, \
and vulnerabilities observed are consistent with those historically used against genomics companies. \
{apt_groups} threat actor groups were observed targeting the sector with varying levels of sophistication."""

    def _add_risk_assessment(self, analysis_result: dict[str, Any]) -> None:
        """Add quarterly risk assessment section with risk cards."""
        logger.info("Adding Quarterly Risk Assessment section")

        # Quarterly Risk Assessment heading - Heading 1
        risk_heading = self.doc.add_heading("Quarterly Risk Assessment", level=1)
        for run in risk_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        # Add space before and after heading
        risk_heading.paragraph_format.space_before = Pt(12)
        risk_heading.paragraph_format.space_after = Pt(6)

        # Add explanatory paragraph for risk ratings
        explanation = self.doc.add_paragraph()
        explanation.paragraph_format.space_before = Pt(0)
        explanation.paragraph_format.space_after = Pt(12)
        explanation.paragraph_format.line_spacing = 1.15

        exp_text = (
            "Risk levels are assessed based on observed threat actor activity, peer organization incidents, "
            "and potential business impact. "
        )
        exp_run = explanation.add_run(exp_text)
        exp_run.font.name = "Arial"
        exp_run.font.size = Pt(11)
        exp_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add HIGH definition
        high_run = explanation.add_run("HIGH")
        high_run.font.name = "Arial"
        high_run.font.size = Pt(11)
        high_run.font.bold = True
        high_run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)  # Red

        high_def = explanation.add_run(" indicates active targeting with multiple incidents; ")
        high_def.font.name = "Arial"
        high_def.font.size = Pt(11)
        high_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add MEDIUM definition
        medium_run = explanation.add_run("MEDIUM")
        medium_run.font.name = "Arial"
        medium_run.font.size = Pt(11)
        medium_run.font.bold = True
        medium_run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)  # Orange

        medium_def = explanation.add_run(" reflects ongoing activity with moderate incident levels; ")
        medium_def.font.name = "Arial"
        medium_def.font.size = Pt(11)
        medium_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add LOW definition
        low_run = explanation.add_run("LOW")
        low_run.font.name = "Arial"
        low_run.font.size = Pt(11)
        low_run.font.bold = True
        low_run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)  # Green

        low_def = explanation.add_run(" represents minimal observed activity or limited sector-specific targeting.")
        low_def.font.name = "Arial"
        low_def.font.size = Pt(11)
        low_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # `or {}` guards a present-but-null value (the AI can emit `"risk_assessment": null`,
        # which `.get(k, {})` would NOT protect against).
        risk_data = analysis_result.get("risk_assessment") or {}

        # Load historical data to compare with previous quarter
        history = self._load_historical_data()
        year = self._get_year()
        prev_year, prev_quarter = self._calculate_previous_quarter(year, self.quarter)
        prev_quarter_key = self._get_quarter_key(prev_year, prev_quarter)
        previous_assessment = history.get(prev_quarter_key, {})

        # Get current risk levels. `(x or default)` coerces present-but-null values too.
        # str(...) so a non-string risk value (e.g. an int) doesn't crash .upper(); mirrors the
        # persistence path. `or DEFAULT` still handles present-but-null.
        current_nation_state = str(risk_data.get("nation_state") or RiskLevel.HIGH).upper()
        current_ransomware = str(risk_data.get("ransomware") or RiskLevel.HIGH).upper()
        current_supply_chain = str(risk_data.get("supply_chain") or RiskLevel.MEDIUM).upper()
        current_insider = str(risk_data.get("insider") or RiskLevel.LOW).upper()

        # Get AI's trend assessments from analysis
        ai_nation_state_trend = risk_data.get("nation_state_trend", RiskLevel.UNCHANGED)
        ai_ransomware_trend = risk_data.get("ransomware_trend", RiskLevel.UNCHANGED)
        ai_supply_chain_trend = risk_data.get("supply_chain_trend", RiskLevel.UNCHANGED)
        ai_insider_trend = risk_data.get("insider_trend", RiskLevel.UNCHANGED)

        # Calculate trends by comparing with previous quarter
        # ONLY use historical comparison if we have prior quarter data AND the AI didn't provide trends
        # Otherwise, trust the AI's analysis which considers breach statistics
        if (
            previous_assessment
            and ai_nation_state_trend == RiskLevel.UNCHANGED
            and ai_ransomware_trend == RiskLevel.UNCHANGED
        ):
            # Historical comparison mode - calculate trends from stored risk levels
            nation_state_trend = self._compare_with_previous_quarter(
                current_nation_state, previous_assessment.get("nation_state", "MEDIUM")
            )
            ransomware_trend = self._compare_with_previous_quarter(
                current_ransomware, previous_assessment.get("ransomware", "MEDIUM")
            )
            supply_chain_trend = self._compare_with_previous_quarter(
                current_supply_chain, previous_assessment.get("supply_chain", "MEDIUM")
            )
            insider_trend = self._compare_with_previous_quarter(
                current_insider, previous_assessment.get("insider", "LOW")
            )
            logger.info(
                f"Using historical comparison with {prev_quarter_key}: trends calculated from stored risk levels"
            )
        else:
            # Use AI's trend assessment (which considers breach statistics, not just risk level changes)
            nation_state_trend = ai_nation_state_trend
            ransomware_trend = ai_ransomware_trend
            supply_chain_trend = ai_supply_chain_trend
            insider_trend = ai_insider_trend
            logger.info("Using AI's trend assessment based on breach statistics and threat intelligence")

        # Save current assessment for future comparisons (risk levels + breach stats so
        # next quarter's prior column and QoQ percentages are grounded in a real baseline).
        self._save_current_risk_assessment(
            {
                "nation_state": current_nation_state,
                "ransomware": current_ransomware,
                "supply_chain": current_supply_chain,
                "insider": current_insider,
            },
            breach_stats=self._extract_breach_stats(analysis_result),
        )

        risks = [
            (
                "Nation-State Espionage",
                current_nation_state,
                nation_state_trend,
                None,  # No percentage data available
            ),
            (
                "Ransomware & Extortion",
                current_ransomware,
                ransomware_trend,
                None,  # Could calculate from breach data if available
            ),
            (
                "Supply Chain Compromise",
                current_supply_chain,
                supply_chain_trend,
                None,  # No percentage data available
            ),
            (
                "Insider Threat",
                current_insider,
                insider_trend,
                None,  # No percentage data available
            ),
        ]

        # Create horizontal table (1 row, 4 columns)
        table = self.doc.add_table(rows=1, cols=len(risks))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT  # Left align table like the example

        # Set table width to 100%
        tbl = table._element
        # Get or create tblPr element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)  # Insert at the beginning

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "5000")  # 5000 = 100% in Word's 50ths of a percent
        tbl_w.set(qn("w:type"), "pct")  # Percentage type

        # Fill table cells with risk data
        for i, (category, level, trend, calculated_pct) in enumerate(risks):
            cell = table.rows[0].cells[i]
            cell.paragraphs[0].clear()

            # Set cell vertical alignment to center for better alignment
            tc_pr = cell._element.get_or_add_tcPr()
            v_align = tc_pr.find(qn("w:vAlign"))
            if v_align is None:
                v_align = OxmlElement("w:vAlign")
                tc_pr.append(v_align)
            v_align.set(qn("w:val"), "center")  # Center vertically

            # Set cell padding: 120 twips all sides
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)

            for margin_type in ["top", "bottom", "left", "right"]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), "120")  # 120 twips
                margin.set(qn("w:type"), "dxa")

            # Determine background and text color based on risk level
            level_upper = level.upper()
            if level_upper == "HIGH":
                bg_color = "FEE2E2"  # Light red
                text_color = RGBColor(0x99, 0x1B, 0x1B)  # Dark red
            elif level_upper == "MEDIUM":
                bg_color = "FEF3C7"  # Light amber
                text_color = RGBColor(0x92, 0x40, 0x0E)  # Dark amber
            elif level_upper == "LOW":
                bg_color = "D1FAE5"  # Light green
                text_color = RGBColor(0x06, 0x5F, 0x46)  # Dark green
            else:
                # Default for unexpected values
                bg_color = "FFFFFF"  # White
                text_color = BrandColors.GRAY_MEDIUM
                logger.warning(f"Unexpected risk level '{level}' - using default styling")

            # Set cell background
            self._set_cell_shading(cell, bg_color)

            # Set cell borders (BORDER_GRAY, size 4)
            self._set_cell_borders(cell, color_hex="CCCCCC", size="4")

            # Category name - centered, single line
            cat_para = cell.paragraphs[0]
            cat_para.paragraph_format.space_before = Pt(0)
            cat_para.paragraph_format.space_after = Pt(0)
            # Replace any newlines with spaces to keep on single line
            category_single_line = category.replace("\n", " ")
            cat_run = cat_para.add_run(category_single_line)
            cat_run.font.name = "Arial"
            cat_run.font.size = Pt(10)
            cat_run.font.bold = True
            cat_run.font.color.rgb = BrandColors.TEXT_DARK
            cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Risk level - large, bold, colored
            level_para = cell.add_paragraph()
            level_para.paragraph_format.space_before = Pt(0)
            level_para.paragraph_format.space_after = Pt(0)
            level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            level_run = level_para.add_run(level)
            level_run.font.name = "Arial"
            level_run.font.size = Pt(20)
            level_run.font.bold = True
            level_run.font.color.rgb = text_color

            # Trend vs previous quarter - format with percentage if available
            prev_quarter_short = f"Q{self.quarter - 1 if self.quarter > 1 else 4}"
            trend_para = cell.add_paragraph()
            trend_para.paragraph_format.space_before = Pt(0)
            trend_para.paragraph_format.space_after = Pt(0)
            trend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Parse trend to extract percentage or determine display
            trend_text = str(trend).strip()
            trend_display = None

            # Prefer calculated percentage if available
            if calculated_pct:
                trend_display = f"vs {prev_quarter_short}: {calculated_pct}"
            # Check if trend contains percentage
            elif "%" in trend_text:
                trend_display = f"vs {prev_quarter_short}: {trend_text}"
            elif trend_text.upper() == "UNCHANGED":
                trend_display = f"vs {prev_quarter_short}: Unchanged"
            elif trend_text in ["↑", "INCREASED", "INCREASE"]:
                trend_display = f"vs {prev_quarter_short}: Increased"
            elif trend_text in ["↓", "DECREASED", "DECREASE"]:
                trend_display = f"vs {prev_quarter_short}: Decreased"
            else:
                # Try to extract percentage from trend string
                percentage_match = re.search(r"([+-]?\d+(?:\.\d+)?%)", trend_text)
                if percentage_match:
                    percentage = percentage_match.group(1)
                    trend_display = f"vs {prev_quarter_short}: {percentage}"
                else:
                    trend_display = f"vs {prev_quarter_short}: Unchanged"

            trend_run = trend_para.add_run(trend_display)
            trend_run.font.name = "Arial"
            trend_run.font.size = Pt(8)
            trend_run.font.italic = True
            trend_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        # Spacer after risk assessment
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _as_bullets(value: Any) -> list[str]:
        """Normalize a bullet field to a list of non-empty strings.

        Accepts a list (each item stringified) or a single string (kept whole, not sliced into
        characters). Anything else -> empty list.
        """
        if isinstance(value, str):
            return [value] if value.strip() else []
        if isinstance(value, list):
            return [str(v).strip() for v in value if v is not None and str(v).strip()]
        return []

    @staticmethod
    def _change_pct_color(change_pct: str) -> RGBColor:
        """Color for a QoQ delta: red for an increase (+), green for a decrease (-), gray for
        0% / N/A / unparseable. Implements the contract the strategic prompt advertises — the
        comparison line was previously hardcoded red, so a good decrease and an honest N/A both
        rendered as alarm-red.
        """
        red = RGBColor(0xDC, 0x35, 0x45)
        green = RGBColor(0x1A, 0x7F, 0x37)
        gray = RGBColor(0x6B, 0x72, 0x80)
        t = (change_pct or "").strip().lower()
        if t in {"", "n/a", "na", "none", "null"}:
            return gray
        if t[0] == "-" or t[0] == "−":  # ASCII hyphen or Unicode minus
            return green
        if t[0] == "+":
            return red
        m = re.search(r"-?\d+(?:\.\d+)?", t)
        if not m:
            return gray
        val = float(m.group())
        return red if val > 0 else green if val < 0 else gray

    def _add_breach_landscape(self, analysis_result: dict[str, Any]) -> None:
        """Add industry breach landscape section."""
        logger.info("Adding Industry Breach Landscape section")

        # COMPONENT 1 — Section heading
        breach_heading = self.doc.add_heading("Industry Breach Landscape", level=1)
        for run in breach_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        breach_heading.paragraph_format.space_before = Pt(12)
        breach_heading.paragraph_format.space_after = Pt(6)

        # Get breach landscape data
        breach_data = analysis_result.get("breach_landscape") or {}

        # If missing, render unavailable message and return
        if not breach_data:
            logger.warning("breach_landscape missing from analysis_result")
            unavailable_para = self.doc.add_paragraph()
            unavailable_run = unavailable_para.add_run("Breach landscape data unavailable for this reporting period.")
            unavailable_run.font.name = "Arial"
            unavailable_run.font.size = Pt(10)
            unavailable_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return

        # COMPONENT 2 — Italic scope note
        scope_note = breach_data.get("scope_note", "")
        if scope_note:
            scope_para = self.doc.add_paragraph()
            scope_para.paragraph_format.space_before = Pt(0)
            scope_para.paragraph_format.space_after = Pt(6)
            scope_run = scope_para.add_run(scope_note)
            scope_run.font.name = "Arial"
            scope_run.font.size = Pt(9)
            scope_run.font.italic = True
            scope_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Spacer after scope note
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # COMPONENT 3 — Stat cards. Render whatever the AI provides, up to 4 (previously
        # only an exact count of 4 rendered, so 1-3 or 5 cards silently produced nothing).
        stat_cards = breach_data.get("stat_cards")
        stat_cards = [c for c in stat_cards if isinstance(c, dict)] if isinstance(stat_cards, list) else []
        if stat_cards:
            num_cards = min(len(stat_cards), 4)
            stat_cards = stat_cards[:num_cards]
            # Create single-row table sized to the number of cards
            table = self.doc.add_table(rows=1, cols=num_cards)
            table.autofit = False
            table.style = None

            # Set column widths to divide the ~6.5in content width evenly
            col_width = Inches(6.5 / num_cards)
            for col in table.columns:
                col.width = col_width

            for i, card in enumerate(stat_cards):
                cell = table.rows[0].cells[i]
                cell.paragraphs[0].clear()

                # Set cell background to light blue-gray (#E8F4F8 or similar light blue)
                self._set_cell_shading(cell, "E8F4F8")

                # Apply thin gray borders on all sides (no accent border)
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)

                for border_name in ["top", "left", "right", "bottom"]:
                    border = tc_borders.find(qn(f"w:{border_name}"))
                    if border is None:
                        border = OxmlElement(f"w:{border_name}")
                        tc_borders.append(border)
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "4")
                    border.set(qn("w:color"), "D0D0D0")  # Light gray border

                # Paragraph 1 — large display value (black, bold)
                value_para = cell.paragraphs[0]
                value_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                value_para.paragraph_format.space_after = Pt(2)
                value_run = value_para.add_run(str(card.get("value", "")))
                value_run.font.name = "Arial"
                value_run.font.size = Pt(24)
                value_run.font.bold = True
                value_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

                # Paragraph 2 — label (black, bold)
                label_para = cell.add_paragraph()
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label_para.paragraph_format.space_after = Pt(4)
                label_run = label_para.add_run(str(card.get("label", "")))
                label_run.font.name = "Arial"
                label_run.font.size = Pt(10)
                label_run.font.bold = True
                label_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

                # Paragraph 3 — prior quarter comparison (italic; color by direction, single line)
                prior_label = str(card.get("prior_label", "")).strip()
                prior_value = str(card.get("prior_value", "")).strip()
                change_pct = str(card.get("change_pct", "")).strip()

                prior_para = cell.add_paragraph()
                prior_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Compose the comparison text, degrading gracefully when no real
                # prior-quarter data is available (Q9: prior_value/change_pct == "N/A").
                # A fabricated "(+25%)" must never appear just to fill the slot.
                na_tokens = {"", "n/a", "na", "none", "null"}
                has_prior = prior_value.lower() not in na_tokens
                has_change = change_pct.lower() not in na_tokens
                if not has_prior and not has_change:
                    comparison_text = f"{prior_label}: N/A" if prior_label else "No prior-quarter comparison"
                elif has_prior and has_change:
                    comparison_text = f"{prior_label}: {prior_value} ({change_pct})"
                elif has_prior:
                    comparison_text = f"{prior_label}: {prior_value}"
                else:
                    comparison_text = f"{prior_label}: {change_pct}"
                comparison_run = prior_para.add_run(comparison_text)
                comparison_run.font.name = "Arial"
                comparison_run.font.size = Pt(8)
                comparison_run.font.italic = True
                comparison_run.font.bold = False
                # Color by direction (red up / green down / gray flat|N/A). Gray whenever there is
                # no real change value to signal on.
                comparison_run.font.color.rgb = self._change_pct_color(change_pct if has_change else "")

        # Spacer after stat cards
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # COMPONENT 4 — Incidents by type. Accept it nested under breach_landscape
        # (current AI schema) OR at the top level of the analysis (docstring/tests).
        incidents = breach_data.get("incidents_by_type") or analysis_result.get("incidents_by_type", [])
        # Guard against the AI emitting a list of strings (weekly-schema habit) — the row loop
        # below calls incident.get(...), which would crash on a str. Mirrors the stat-card and
        # geopolitical guards.
        incidents = [i for i in incidents if isinstance(i, dict)]
        current_quarter_label = breach_data.get("current_quarter_label", "Current")
        prior_quarter_label = breach_data.get("prior_quarter_label", "Prior")

        # Subheading
        incidents_heading = self.doc.add_paragraph()
        incidents_heading.paragraph_format.space_before = Pt(10)
        incidents_heading.paragraph_format.space_after = Pt(4)
        heading_run = incidents_heading.add_run("Incidents by Type")
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(11)
        heading_run.font.bold = True
        heading_run.font.color.rgb = BrandColors.ORANGE_PRIMARY

        if incidents:
            # Create table with header row + data rows
            table = self.doc.add_table(rows=1 + len(incidents), cols=4)
            table.autofit = False
            table.style = None

            # Set column widths
            table.columns[0].width = Inches(1.39)
            table.columns[1].width = Inches(0.83)
            table.columns[2].width = Inches(0.83)
            table.columns[3].width = Inches(3.45)

            # Header row
            headers = ["Incident Type", current_quarter_label, prior_quarter_label, "Notable Example"]
            header_cells = table.rows[0].cells

            for i, header_text in enumerate(headers):
                cell = header_cells[i]
                cell.paragraphs[0].clear()

                # Set background to #E65100 (orange)
                self._set_cell_shading(cell, "E65100")

                # Header text
                header_run = cell.paragraphs[0].add_run(header_text)
                header_run.font.name = "Arial"
                header_run.font.size = Pt(10)
                header_run.font.bold = True
                header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Cell padding
                tc_pr = cell._element.get_or_add_tcPr()
                tc_mar = tc_pr.find(qn("w:tcMar"))
                if tc_mar is None:
                    tc_mar = OxmlElement("w:tcMar")
                    tc_pr.append(tc_mar)

                for margin_type in ["top", "bottom", "left", "right"]:
                    margin = tc_mar.find(qn(f"w:{margin_type}"))
                    if margin is None:
                        margin = OxmlElement(f"w:{margin_type}")
                        tc_mar.append(margin)
                    if margin_type in ["top", "bottom"]:
                        margin.set(qn("w:w"), "60")
                    else:  # left, right
                        margin.set(qn("w:w"), "80")
                    margin.set(qn("w:type"), "dxa")

            # Data rows
            for row_idx, incident in enumerate(incidents):
                row = table.rows[row_idx + 1]
                cells = row.cells

                # Alternate row backgrounds
                if row_idx % 2 == 0:
                    bg_color = "FFFFFF"  # even rows
                else:
                    bg_color = "F3F4F6"  # odd rows

                for cell in cells:
                    self._set_cell_shading(cell, bg_color)

                # Col 0 — incident type (bold black)
                cells[0].paragraphs[0].clear()
                type_run = cells[0].paragraphs[0].add_run(str(incident.get("type", "")))
                type_run.font.name = "Arial"
                type_run.font.size = Pt(10)
                type_run.font.bold = True
                type_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Col 1 — current_count (large bold black)
                cells[1].paragraphs[0].clear()
                current_run = cells[1].paragraphs[0].add_run(str(incident.get("current_count", "0")))
                current_run.font.name = "Arial"
                current_run.font.size = Pt(14)  # Larger size for emphasis
                current_run.font.bold = True
                current_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Col 2 — prior_count (regular black, not bold). Schema uses `prev_count`;
                # accept `prior_count` too for back-compat, and coerce ints to str.
                cells[2].paragraphs[0].clear()
                prior_run = cells[2].paragraphs[0].add_run(
                    str(incident.get("prev_count", incident.get("prior_count", "0")))
                )
                prior_run.font.name = "Arial"
                prior_run.font.size = Pt(10)
                prior_run.font.bold = False
                prior_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Col 3 — notable_example (italic gray)
                cells[3].paragraphs[0].clear()
                example_run = cells[3].paragraphs[0].add_run(str(incident.get("notable_example", "")))
                example_run.font.name = "Arial"
                example_run.font.size = Pt(9)
                example_run.font.italic = True
                example_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # Gray
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Apply gray borders and padding to all data cells
                for cell in cells:
                    tc_pr = cell._element.get_or_add_tcPr()

                    # Borders
                    tc_borders = tc_pr.find(qn("w:tcBorders"))
                    if tc_borders is None:
                        tc_borders = OxmlElement("w:tcBorders")
                        tc_pr.append(tc_borders)

                    for border_name in ["top", "left", "bottom", "right"]:
                        border = tc_borders.find(qn(f"w:{border_name}"))
                        if border is None:
                            border = OxmlElement(f"w:{border_name}")
                            tc_borders.append(border)
                        border.set(qn("w:val"), "single")
                        border.set(qn("w:sz"), "4")
                        border.set(qn("w:color"), "D1D5DB")

                    # Padding
                    tc_mar = tc_pr.find(qn("w:tcMar"))
                    if tc_mar is None:
                        tc_mar = OxmlElement("w:tcMar")
                        tc_pr.append(tc_mar)

                    for margin_type in ["top", "bottom", "left", "right"]:
                        margin = tc_mar.find(qn(f"w:{margin_type}"))
                        if margin is None:
                            margin = OxmlElement(f"w:{margin_type}")
                            tc_mar.append(margin)
                        if margin_type in ["top", "bottom"]:
                            margin.set(qn("w:w"), "60")
                        else:  # left, right
                            margin.set(qn("w:w"), "80")
                        margin.set(qn("w:type"), "dxa")

        # COMPONENT 5 — Common factors
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        common_factors = breach_data.get("common_factors", "")
        if not isinstance(common_factors, str):
            common_factors = str(common_factors or "")
        if common_factors:
            # Subheading
            factors_heading = self.doc.add_paragraph()
            factors_heading.paragraph_format.space_before = Pt(8)
            factors_heading.paragraph_format.space_after = Pt(4)
            factors_heading_run = factors_heading.add_run("Common Factors Across Incidents")
            factors_heading_run.font.name = "Arial"
            factors_heading_run.font.size = Pt(11)
            factors_heading_run.font.bold = True
            factors_heading_run.font.color.rgb = BrandColors.ORANGE_PRIMARY

            # Body paragraph
            factors_para = self.doc.add_paragraph()
            factors_para.paragraph_format.space_after = Pt(6)
            factors_run = factors_para.add_run(common_factors)
            factors_run.font.name = "Arial"
            factors_run.font.size = Pt(10)
            factors_run.font.bold = False
            factors_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

            # If the narrative cites bare percentages, add a short caveat so an AI estimate is not
            # read as a hard measured statistic — these figures are not dataset-grounded.
            if re.search(r"\d\s*%|\bpercent\b", common_factors, re.IGNORECASE):
                caveat = self.doc.add_paragraph()
                caveat.paragraph_format.space_after = Pt(6)
                caveat_run = caveat.add_run(
                    "Percentages above reflect analyst assessment of this quarter's incidents, "
                    "not a measured dataset statistic."
                )
                caveat_run.font.name = "Arial"
                caveat_run.font.size = Pt(8)
                caveat_run.font.italic = True
                caveat_run.font.color.rgb = BrandColors.GRAY_MEDIUM

        # Spacer after breach landscape
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        logger.info("Breach landscape section added")

    def _add_geopolitical_landscape(self, analysis_result: dict[str, Any]) -> None:
        """Add geopolitical threat landscape section with dynamic card table."""
        logger.info("Adding Geopolitical Threat Landscape section")

        geo_heading = self.doc.add_heading("Geopolitical Threat Landscape", level=1)
        for run in geo_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)  # Font size 14pt
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange color
        # Add space before and after heading
        geo_heading.paragraph_format.space_before = Pt(12)
        geo_heading.paragraph_format.space_after = Pt(6)

        # Subtitle
        subtitle = self.doc.add_paragraph()
        sub_run = subtitle.add_run(
            f"Nation-state activity assessed for direct relevance to {customer_profile.name}'s assets, operations, and competitive position — Q{self.quarter} {self._get_year()}."
        )
        sub_run.font.name = "Arial"
        sub_run.font.size = FontSizes.SUBTITLE
        sub_run.font.italic = True
        sub_run.font.color.rgb = BrandColors.GRAY_MEDIUM
        subtitle.paragraph_format.keep_with_next = True  # Keep with table

        # Add explanatory paragraph for threat levels
        explanation = self.doc.add_paragraph()
        explanation.paragraph_format.space_before = Pt(6)
        explanation.paragraph_format.space_after = Pt(12)
        explanation.paragraph_format.line_spacing = 1.15

        exp_text = (
            "Threat levels reflect the combination of actor capability, demonstrated intent, and targeting frequency. "
        )
        exp_run = explanation.add_run(exp_text)
        exp_run.font.name = "Arial"
        exp_run.font.size = Pt(11)
        exp_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add HIGH definition
        high_run = explanation.add_run("HIGH")
        high_run.font.name = "Arial"
        high_run.font.size = Pt(11)
        high_run.font.bold = True
        high_run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)  # Red

        high_def = explanation.add_run(" indicates systematic sector targeting with confirmed intrusions; ")
        high_def.font.name = "Arial"
        high_def.font.size = Pt(11)
        high_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add MEDIUM definition
        medium_run = explanation.add_run("MEDIUM")
        medium_run.font.name = "Arial"
        medium_run.font.size = Pt(11)
        medium_run.font.bold = True
        medium_run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)  # Orange

        medium_def = explanation.add_run(" reflects opportunistic targeting or reconnaissance activity; ")
        medium_def.font.name = "Arial"
        medium_def.font.size = Pt(11)
        medium_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Add LOW definition
        low_run = explanation.add_run("LOW")
        low_run.font.name = "Arial"
        low_run.font.size = Pt(11)
        low_run.font.bold = True
        low_run.font.color.rgb = RGBColor(0x28, 0xA7, 0x45)  # Green

        low_def = explanation.add_run(" represents limited capability or minimal sector-specific interest.")
        low_def.font.name = "Arial"
        low_def.font.size = Pt(11)
        low_def.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Spacer after subtitle
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # Get geopolitical threats from AI analysis (should be a list of dicts)
        geopolitical_list = analysis_result.get("geopolitical_threats", [])

        # If it's a dict (old format), skip rendering - log warning
        if isinstance(geopolitical_list, dict):
            logger.warning("geopolitical_threats is a dict (old format), expected list. Skipping geopolitical section.")
            no_data_para = self.doc.add_paragraph()
            no_data_run = no_data_para.add_run(
                "No significant nation-state threat activity identified in this reporting period."
            )
            no_data_run.font.name = "Arial"
            no_data_run.font.size = FontSizes.BODY
            no_data_run.font.italic = True
            no_data_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return

        # Cap at 4 countries max
        if len(geopolitical_list) > 4:
            logger.warning(f"AI returned {len(geopolitical_list)} countries, capping at 4 for readability")
            geopolitical_list = geopolitical_list[:4]

        # Validate and filter geopolitical entries
        valid_entries = []
        for entry in geopolitical_list:
            # Skip non-dict entries (the AI can emit a bare list of country strings) — this
            # is the quarterly analogue of the weekly "'IOC' object has no attribute get" crash.
            if not isinstance(entry, dict):
                logger.warning(f"Skipping non-dict geopolitical entry: {entry!r}")
                continue
            # Extract fields for validation (support both 'name' and 'country' fields)
            name = entry.get("name", "").strip() if entry.get("name") else ""
            country = entry.get("country", "").strip() if entry.get("country") else ""
            display_name = entry.get("display_name", "").strip() if entry.get("display_name") else ""
            level = entry.get("threat_level", "").strip() if entry.get("threat_level") else ""
            if not level:
                level = entry.get("level", "").strip() if entry.get("level") else ""
            relevance = entry.get("relevance", [])
            activity = entry.get("activity", [])
            risk = entry.get("risk", [])

            # Check if we have at least one valid identifier (name, country, or display_name)
            # At least one must be present and not "Unknown"
            has_valid_name = (
                (name and name.upper() != "UNKNOWN")
                or (country and country.upper() != "UNKNOWN")
                or (display_name and display_name.upper() != "UNKNOWN")
            )

            # Check if level is missing or empty
            level_invalid = not level

            # Check if all three bullet lists are empty
            all_bullets_empty = (
                (not relevance or len(relevance) == 0)
                and (not activity or len(activity) == 0)
                and (not risk or len(risk) == 0)
            )

            # Skip if any validation fails
            if not has_valid_name or level_invalid or all_bullets_empty:
                logger.warning(f"Skipping geopolitical entry with insufficient data: {entry}")
                continue

            valid_entries.append(entry)

        # Update list to only valid entries
        geopolitical_list = valid_entries

        # If no valid entries remain after filtering, render insufficient data message
        if not geopolitical_list:
            no_data_para = self.doc.add_paragraph()
            no_data_run = no_data_para.add_run(
                "Insufficient geopolitical threat data returned for this reporting period. Review ThreatAnalystAgent output."
            )
            no_data_run.font.name = "Arial"
            no_data_run.font.size = FontSizes.BODY
            no_data_run.font.italic = True
            no_data_run.font.color.rgb = BrandColors.GRAY_MEDIUM
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return

        # Calculate column count and widths dynamically based on valid countries
        num_countries = len(geopolitical_list)

        # Column widths based on count (updated to exact specifications)
        if num_countries == 1:
            col_width = Inches(6.5)  # Full content width
        elif num_countries == 2:
            col_width = Inches(3.25)
        elif num_countries == 3:
            col_width = Inches(2.167)
        else:  # 4 countries
            col_width = Inches(1.625)

        # Define geopolitical card color scheme (local to this section only)
        GEO_HEADER_BG = "1E2D3D"  # dark navy charcoal — header row bg
        GEO_METRICS_BG = "F8FAFC"  # near white — metrics strip bg
        GEO_BULLET_BG_A = "FFFFFF"  # white — relevance and risk rows bg
        GEO_BULLET_BG_B = "F3F4F6"  # light gray — activity row bg
        GEO_LABEL_COLOR = BrandColors.ORANGE_PRIMARY  # orange text for section labels

        # Create table with 5 rows (header, metrics, relevance, activity, risk) and N columns (one per country)
        table = self.doc.add_table(rows=5, cols=num_countries)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = None  # Remove any default styling
        table.autofit = False

        # Set table width to content width
        tbl = table._element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)

        # Set column widths
        for col in table.columns:
            col.width = col_width

        # Prevent page breaks within card rows
        for row in table.rows:
            tr = row._tr
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = OxmlElement("w:trPr")
                tr.insert(0, tr_pr)
            cant_split = OxmlElement("w:cantSplit")
            cant_split.set(qn("w:val"), "1")
            tr_pr.append(cant_split)

        # Populate each country card (column)
        for col_idx, country_data in enumerate(geopolitical_list):
            if not isinstance(country_data, dict):
                continue
            # Extract country data (support both 'name' and 'country' fields)
            country_name = country_data.get("name", "")
            if not country_name:
                country_name = country_data.get("country", "Unknown")
            display_name = country_data.get("display_name") or country_name  # `or` guards present-null

            # Truncate display name if longer than 20 characters
            if len(display_name) > 20:
                truncate_pos = display_name.rfind(" ", 0, 20)
                if truncate_pos > 0:
                    display_name = display_name[:truncate_pos] + "..."
                else:
                    display_name = display_name[:17] + "..."
                logger.debug(f"Geo card name truncated: '{country_name}' -> '{display_name}'")

            # The AI schema emits "level"; older/fallback paths use "threat_level". Accept
            # both (was always falling through to "MEDIUM" because only threat_level was read).
            threat_level = str(country_data.get("threat_level") or country_data.get("level") or "MEDIUM").upper()
            primary_vector = country_data.get("vector", "Multiple vectors")
            # Present-but-null "exposure" returns None; the .get default only fires on a missing
            # key. str(... or "MEDIUM") mirrors the threat_level guard above (was a NoneType crash).
            exposure = str(country_data.get("exposure") or "MEDIUM").upper()
            # Coerce each bullet field to a list of strings. If the AI emits a single string,
            # slicing/iterating it later would render one character per bullet ("• C", "• h"); a
            # list of non-dicts would crash on concatenation. _as_bullets normalizes both.
            relevance_bullets = self._as_bullets(country_data.get("relevance"))
            activity_bullets = self._as_bullets(country_data.get("activity"))
            risk_bullets = self._as_bullets(country_data.get("risk"))

            # ============================================================
            # ROW 1 — HEADER ROW (dark charcoal background, white text)
            # ============================================================
            header_cell = table.rows[0].cells[col_idx]
            header_cell.width = col_width
            header_cell.paragraphs[0].clear()

            # Set background to dark charcoal
            self._set_cell_shading(header_cell, GEO_HEADER_BG)

            # Set all borders to match the header background color (dark charcoal) for seamless look
            tc_pr = header_cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for border_name in ["top", "left", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), GEO_HEADER_BG)  # Match cell background

            # Set cell padding: 100 top/bottom, 100 left, 80 right (twips)
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_type, value in [("top", "100"), ("bottom", "100"), ("left", "100"), ("right", "80")]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")

            # Country name paragraph
            name_para = header_cell.paragraphs[0]
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_para.paragraph_format.space_after = Pt(3)
            name_run = name_para.add_run(display_name)
            name_run.font.name = "Arial"
            name_run.font.size = Pt(10)
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White

            # Threat level paragraph (two runs on same line)
            level_para = header_cell.add_paragraph()
            level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            level_para.paragraph_format.space_after = Pt(0)

            # Run A: "THREAT LEVEL  " label
            level_label_run = level_para.add_run("THREAT LEVEL  ")
            level_label_run.font.name = "Arial"
            level_label_run.font.size = Pt(7)
            level_label_run.font.bold = False
            level_label_run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)  # Muted blue-gray

            # Run B: threat level value (colored based on level)
            level_value_run = level_para.add_run(threat_level)
            level_value_run.font.name = "Arial"
            level_value_run.font.size = Pt(9)
            level_value_run.font.bold = True
            if threat_level == "HIGH":
                level_value_run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x6B)  # Soft red
            elif threat_level == "MEDIUM":
                level_value_run.font.color.rgb = RGBColor(0xFF, 0xD1, 0x66)  # Soft amber
            else:  # LOW
                level_value_run.font.color.rgb = RGBColor(0x06, 0xD6, 0xA0)  # Soft green

            # ============================================================
            # ROW 2 — METRICS STRIP (near white background)
            # ============================================================
            metrics_cell = table.rows[1].cells[col_idx]
            metrics_cell.width = col_width
            metrics_cell.paragraphs[0].clear()

            # Set background to near white
            self._set_cell_shading(metrics_cell, GEO_METRICS_BG)

            # Set borders to match cell background for seamless look
            tc_pr = metrics_cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for border_name in ["top", "left", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), GEO_METRICS_BG)  # Match cell background

            # Set cell padding: 60 top/bottom, 100 left, 80 right (twips)
            tc_pr = metrics_cell._element.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_type, value in [("top", "60"), ("bottom", "60"), ("left", "100"), ("right", "80")]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")

            # Line 1 — Primary vector
            vector_para = metrics_cell.paragraphs[0]
            vector_para.paragraph_format.space_after = Pt(2)
            vector_label_run = vector_para.add_run("Primary vector  ")
            vector_label_run.font.name = "Arial"
            vector_label_run.font.size = Pt(7.5)
            vector_label_run.font.bold = True
            vector_label_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            vector_value_run = vector_para.add_run(primary_vector)
            vector_value_run.font.name = "Arial"
            vector_value_run.font.size = Pt(7.5)
            vector_value_run.font.bold = False
            vector_value_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

            # Line 2 — organization exposure
            exposure_para = metrics_cell.add_paragraph()
            exposure_para.paragraph_format.space_after = Pt(0)
            exposure_label_run = exposure_para.add_run(f"{customer_profile.name} exposure  ")
            exposure_label_run.font.name = "Arial"
            exposure_label_run.font.size = Pt(7.5)
            exposure_label_run.font.bold = True
            exposure_label_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            exposure_value_run = exposure_para.add_run(exposure)
            exposure_value_run.font.name = "Arial"
            exposure_value_run.font.size = Pt(7.5)
            exposure_value_run.font.bold = True
            if exposure == "CRITICAL":
                exposure_value_run.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
            elif exposure in ("HIGH", "MEDIUM"):
                exposure_value_run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            else:  # LOW
                exposure_value_run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)

            # ============================================================
            # ROW 3 — RELEVANCE TO ILLUMINA (white background)
            # ============================================================
            relevance_cell = table.rows[2].cells[col_idx]
            relevance_cell.width = col_width
            relevance_cell.paragraphs[0].clear()

            # Set background to white
            self._set_cell_shading(relevance_cell, GEO_BULLET_BG_A)

            # Set borders to match cell background for seamless look
            tc_pr = relevance_cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for border_name in ["top", "left", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), GEO_BULLET_BG_A)  # Match cell background

            # Set cell padding: 80 top/bottom, 100 left, 80 right (twips)
            tc_pr = relevance_cell._element.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_type, value in [("top", "80"), ("bottom", "80"), ("left", "100"), ("right", "80")]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")

            # Section label paragraph
            label_para = relevance_cell.paragraphs[0]
            label_para.paragraph_format.space_after = Pt(3)
            label_run = label_para.add_run("RELEVANCE TO ILLUMINA")
            label_run.font.name = "Arial"
            label_run.font.size = Pt(7)
            label_run.font.bold = True
            label_run.font.color.rgb = GEO_LABEL_COLOR

            # Bullets (max 2, truncate at 120 chars)
            bullets_to_render = relevance_bullets[:2]
            if len(relevance_bullets) > 2:
                logger.debug(f"Geo card bullets capped at 2 for section 'RELEVANCE', country '{display_name}'")

            for bullet_text in bullets_to_render:
                # Truncate bullet to 120 characters
                if len(bullet_text) > 120:
                    truncate_pos = bullet_text.rfind(" ", 0, 120)
                    if truncate_pos > 0:
                        bullet_text = bullet_text[:truncate_pos] + "..."
                    else:
                        bullet_text = bullet_text[:117] + "..."
                    logger.debug(f"Geo bullet truncated for '{display_name}': '{bullet_text[:40]}...'")

                # Use bullet character prefix (not List Bullet style)
                bullet_para = relevance_cell.add_paragraph()
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_para.add_run("\u2022  " + bullet_text)
                bullet_run.font.name = "Arial"
                bullet_run.font.size = Pt(7.5)
                bullet_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

            # ============================================================
            # ROW 4 — Q2 ACTIVITY (light gray background)
            # ============================================================
            activity_cell = table.rows[3].cells[col_idx]
            activity_cell.width = col_width
            activity_cell.paragraphs[0].clear()

            # Set background to light gray
            self._set_cell_shading(activity_cell, GEO_BULLET_BG_B)

            # Set borders to match cell background for seamless look
            tc_pr = activity_cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for border_name in ["top", "left", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), GEO_BULLET_BG_B)  # Match cell background

            # Set cell padding: 80 top/bottom, 100 left, 80 right (twips)
            tc_pr = activity_cell._element.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_type, value in [("top", "80"), ("bottom", "80"), ("left", "100"), ("right", "80")]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")

            # Section label paragraph
            label_para = activity_cell.paragraphs[0]
            label_para.paragraph_format.space_after = Pt(3)
            label_run = label_para.add_run(f"Q{self.quarter} ACTIVITY")
            label_run.font.name = "Arial"
            label_run.font.size = Pt(7)
            label_run.font.bold = True
            label_run.font.color.rgb = GEO_LABEL_COLOR

            # Bullets (max 2, truncate at 120 chars)
            bullets_to_render = activity_bullets[:2]
            if len(activity_bullets) > 2:
                logger.debug(f"Geo card bullets capped at 2 for section 'ACTIVITY', country '{display_name}'")

            for bullet_text in bullets_to_render:
                # Truncate bullet to 120 characters
                if len(bullet_text) > 120:
                    truncate_pos = bullet_text.rfind(" ", 0, 120)
                    if truncate_pos > 0:
                        bullet_text = bullet_text[:truncate_pos] + "..."
                    else:
                        bullet_text = bullet_text[:117] + "..."
                    logger.debug(f"Geo bullet truncated for '{display_name}': '{bullet_text[:40]}...'")

                # Use bullet character prefix (not List Bullet style)
                bullet_para = activity_cell.add_paragraph()
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_para.add_run("\u2022  " + bullet_text)
                bullet_run.font.name = "Arial"
                bullet_run.font.size = Pt(7.5)
                bullet_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

            # ============================================================
            # ROW 5 — RISK TO ILLUMINA (white background)
            # ============================================================
            risk_cell = table.rows[4].cells[col_idx]
            risk_cell.width = col_width
            risk_cell.paragraphs[0].clear()

            # Set background to white
            self._set_cell_shading(risk_cell, GEO_BULLET_BG_A)

            # Set borders to match cell background for seamless look
            tc_pr = risk_cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for border_name in ["top", "left", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), GEO_BULLET_BG_A)  # Match cell background

            # Set cell padding: 80 top/bottom, 100 left, 80 right (twips)
            tc_pr = risk_cell._element.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_type, value in [("top", "80"), ("bottom", "80"), ("left", "100"), ("right", "80")]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), value)
                margin.set(qn("w:type"), "dxa")

            # Section label paragraph
            label_para = risk_cell.paragraphs[0]
            label_para.paragraph_format.space_after = Pt(3)
            label_run = label_para.add_run("RISK TO ILLUMINA")
            label_run.font.name = "Arial"
            label_run.font.size = Pt(7)
            label_run.font.bold = True
            label_run.font.color.rgb = GEO_LABEL_COLOR

            # Bullets (max 2, truncate at 120 chars)
            bullets_to_render = risk_bullets[:2]
            if len(risk_bullets) > 2:
                logger.debug(f"Geo card bullets capped at 2 for section 'RISK', country '{display_name}'")

            for bullet_text in bullets_to_render:
                # Truncate bullet to 120 characters
                if len(bullet_text) > 120:
                    truncate_pos = bullet_text.rfind(" ", 0, 120)
                    if truncate_pos > 0:
                        bullet_text = bullet_text[:truncate_pos] + "..."
                    else:
                        bullet_text = bullet_text[:117] + "..."
                    logger.debug(f"Geo bullet truncated for '{display_name}': '{bullet_text[:40]}...'")

                # Use bullet character prefix (not List Bullet style)
                bullet_para = risk_cell.add_paragraph()
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_run = bullet_para.add_run("\u2022  " + bullet_text)
                bullet_run.font.name = "Arial"
                bullet_run.font.size = Pt(7.5)
                bullet_run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)

        # Spacer after geopolitical landscape
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    # DEPRECATED: Old country section rendering - replaced by dynamic card table in _add_geopolitical_landscape()
    # Keeping these methods for reference but they are no longer called
    #
    # def _add_country_section(self, country: str, data: Dict[str, Any]) -> None:
    #     """Add a country-specific threat section."""
    #     ...
    #
    # def _get_default_strategic_context(self, country: str) -> str:
    #     """Get default strategic context for a country."""
    #     ...
    #
    # def _get_default_activity(self, country: str) -> str:
    #     """Get default activity description for a country."""
    #     ...
    #
    # def _get_default_implications(self, country: str) -> str:
    #     """Get default business implications for a country."""
    #     ...

    def _add_looking_ahead(self, analysis_result: dict[str, Any]) -> None:
        """Add looking ahead section for next quarter."""
        logger.info("Adding Looking Ahead section")

        # Get looking ahead data (`or {}` guards a present-but-null value)
        looking_ahead = analysis_result.get("looking_ahead") or {}

        # If missing or watch_items is empty, render unavailable message
        watch_items = looking_ahead.get("watch_items", [])
        # Defensive: watch_items must be a list of {subject, detail} dicts. Older/
        # malformed analyses sometimes supply a bare string or a list of strings;
        # drop non-dict entries so one bad field cannot crash the whole report.
        if isinstance(watch_items, str) or not isinstance(watch_items, list):
            watch_items = []
        else:
            watch_items = [item for item in watch_items if isinstance(item, dict)]
        if not looking_ahead or not watch_items:
            logger.warning("looking_ahead missing or watch_items empty")

            # Still render heading with fallback quarter calculation
            next_quarter = self.quarter + 1 if self.quarter < 4 else 1
            next_year = self._get_year() if self.quarter < 4 else self._get_year() + 1

            looking_ahead_heading = self.doc.add_heading(f"Looking Ahead: Q{next_quarter} {next_year}", level=1)
            for run in looking_ahead_heading.runs:
                run.font.name = "Arial"
                run.font.size = Pt(14)
                run.font.color.rgb = BrandColors.ORANGE_PRIMARY
            looking_ahead_heading.paragraph_format.space_before = Pt(12)
            looking_ahead_heading.paragraph_format.space_after = Pt(6)

            # Unavailable message
            unavailable_para = self.doc.add_paragraph()
            unavailable_run = unavailable_para.add_run("No specific watch items identified for this reporting period.")
            unavailable_run.font.name = "Arial"
            unavailable_run.font.size = Pt(10)
            unavailable_run.font.italic = True
            unavailable_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return

        # Component 1 — Section heading with next_quarter_label
        next_quarter_label = looking_ahead.get("next_quarter_label", "")
        if not next_quarter_label:
            # Fallback calculation if AI didn't provide label
            next_quarter = self.quarter + 1 if self.quarter < 4 else 1
            next_year = self._get_year() if self.quarter < 4 else self._get_year() + 1
            next_quarter_label = f"Q{next_quarter} {next_year}"

        looking_ahead_heading = self.doc.add_heading(f"Looking Ahead: {next_quarter_label}", level=1)
        for run in looking_ahead_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        looking_ahead_heading.paragraph_format.space_before = Pt(12)
        looking_ahead_heading.paragraph_format.space_after = Pt(6)

        # Component 2 — Subheading (black, not orange)
        subheading_para = self.doc.add_paragraph()
        subheading_para.paragraph_format.space_before = Pt(6)
        subheading_para.paragraph_format.space_after = Pt(2)
        subheading_run = subheading_para.add_run("Specific Watch Items")
        subheading_run.font.name = "Arial"
        subheading_run.font.size = Pt(11)
        subheading_run.font.bold = True
        subheading_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

        # Component 3 — Italic note (gray)
        note_para = self.doc.add_paragraph()
        note_para.paragraph_format.space_before = Pt(0)
        note_para.paragraph_format.space_after = Pt(6)
        note_run = note_para.add_run("Named, specific items — not generic monitoring reminders.")
        note_run.font.name = "Arial"
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # Gray

        # Component 4 — Numbered watch item list
        for i, item in enumerate(watch_items):
            subject = str(item.get("subject", ""))  # str() so a non-string subject can't crash add_run
            detail = item.get("detail", "")

            # Try to use 'List Number' style, fall back to manual numbering
            try:
                item_para = self.doc.add_paragraph(style="List Number")
            except KeyError:
                # 'List Number' style not available, use manual numbering
                item_para = self.doc.add_paragraph()
                # Add manual number prefix
                num_run = item_para.add_run(f"{i + 1}.  ")
                num_run.font.name = "Arial"
                num_run.font.size = Pt(10)
                num_run.font.bold = False
                num_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

            item_para.paragraph_format.space_after = Pt(4)

            # Run 1 — subject (bold, dark navy)
            subject_run = item_para.add_run(subject)
            subject_run.font.name = "Arial"
            subject_run.font.size = Pt(10)
            subject_run.font.bold = True
            subject_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark navy

            # Run 2 — detail (regular, black)
            detail_run = item_para.add_run(f" {detail}")
            detail_run.font.name = "Arial"
            detail_run.font.size = Pt(10)
            detail_run.font.bold = False
            detail_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

        # Spacer after looking ahead
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        logger.info("Looking ahead section added")

    def _add_recommendations(self, analysis_result: dict[str, Any]) -> None:
        """Add recommendations section."""
        logger.info("Adding Recommendations section")

        # Component 1 — Section heading
        rec_heading = self.doc.add_heading("Recommendations", level=1)
        for run in rec_heading.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = BrandColors.ORANGE_DESIGN  # Orange heading
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY
        rec_heading.paragraph_format.space_before = Pt(12)
        rec_heading.paragraph_format.space_after = Pt(6)

        # Get recommendations data. Accept both the {"items": [...], "intro_note": ...} object
        # and a bare list of recommendations (list of dicts, or the old list-of-strings) — a bare
        # list previously fell through to "unavailable" and silently dropped real recommendations.
        recommendations = analysis_result.get("recommendations", {})
        if isinstance(recommendations, list):
            items = recommendations
            intro_note = ""
        elif isinstance(recommendations, dict):
            items = recommendations.get("items", [])
            intro_note = recommendations.get("intro_note", "")
        else:
            items = []
            intro_note = ""
        # Normalize each item to a dict so the render loop's .get(...) never crashes on a
        # non-dict (a list-of-strings becomes a body-only box).
        items = [it if isinstance(it, dict) else {"body": str(it)} for it in items if it]

        # If missing or items is empty, render unavailable message
        if not items:
            logger.warning("recommendations missing or items empty")
            unavailable_para = self.doc.add_paragraph()
            unavailable_run = unavailable_para.add_run("No recommendations generated for this reporting period.")
            unavailable_run.font.name = "Arial"
            unavailable_run.font.size = Pt(10)
            unavailable_run.font.italic = True
            unavailable_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return

        # Component 2 — Italic intro note (computed above; absent for a bare-list payload)
        if intro_note:
            intro_para = self.doc.add_paragraph()
            intro_para.paragraph_format.space_before = Pt(0)
            intro_para.paragraph_format.space_after = Pt(6)
            intro_run = intro_para.add_run(intro_note)
            intro_run.font.name = "Arial"
            intro_run.font.size = Pt(9)
            intro_run.font.italic = True
            intro_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Component 3 — Recommendation boxes
        for index, item in enumerate(items, start=1):
            title = str(item.get("title", ""))
            body = str(item.get("body", ""))

            # Create single-cell table for box
            table = self.doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.style = None

            # Set table width to full content width (6.5 inches)
            table.columns[0].width = Inches(6.5)

            cell = table.rows[0].cells[0]
            cell.paragraphs[0].clear()

            # Set cell background to #FFF3E0 (light orange)
            self._set_cell_shading(cell, "FFF3E0")

            # Apply borders via XML
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)

            # Left border: style SINGLE, size 18, color "E65100" (orange)
            left_border = tc_borders.find(qn("w:left"))
            if left_border is None:
                left_border = OxmlElement("w:left")
                tc_borders.append(left_border)
            left_border.set(qn("w:val"), "single")
            left_border.set(qn("w:sz"), "18")
            left_border.set(qn("w:color"), "E65100")

            # Top, right, bottom borders: style SINGLE, size 4, color "D1D5DB"
            for border_name in ["top", "right", "bottom"]:
                border = tc_borders.find(qn(f"w:{border_name}"))
                if border is None:
                    border = OxmlElement(f"w:{border_name}")
                    tc_borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), "D1D5DB")

            # Cell padding via w:tcMar XML: 80 top/bottom, 120 left/right
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)

            for margin_type in ["top", "bottom", "left", "right"]:
                margin = tc_mar.find(qn(f"w:{margin_type}"))
                if margin is None:
                    margin = OxmlElement(f"w:{margin_type}")
                    tc_mar.append(margin)
                if margin_type in ["top", "bottom"]:
                    margin.set(qn("w:w"), "80")
                else:  # left, right
                    margin.set(qn("w:w"), "120")
                margin.set(qn("w:type"), "dxa")

            # Title paragraph with underline guard
            title_para = cell.paragraphs[0]
            title_para.paragraph_format.space_after = Pt(6)

            # Prefix run: "{index}.  "
            prefix_run = title_para.add_run(f"{index}.  ")
            prefix_run.font.name = "Arial"
            prefix_run.font.size = Pt(11)
            prefix_run.font.bold = True
            prefix_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
            prefix_run.font.underline = False

            # Split title on first space
            if " " in title:
                first_word = title.split(" ", 1)[0]
                remainder = title.split(" ", 1)[1]
            else:
                first_word = title
                remainder = ""

            # Apply underline guard
            if first_word.isalpha():
                # First word is purely alphabetic - underline it
                first_word_run = title_para.add_run(first_word)
                first_word_run.font.name = "Arial"
                first_word_run.font.size = Pt(11)
                first_word_run.font.bold = True
                first_word_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
                first_word_run.font.underline = True

                # Remainder (if exists)
                if remainder:
                    remainder_run = title_para.add_run(f" {remainder}")
                    remainder_run.font.name = "Arial"
                    remainder_run.font.size = Pt(11)
                    remainder_run.font.bold = True
                    remainder_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
                    remainder_run.font.underline = False
            else:
                # First word not purely alphabetic - no underline, render entire title
                logger.debug(f"Rec title underline skipped — first word not purely alphabetic: '{first_word}'")
                full_title_run = title_para.add_run(title)
                full_title_run.font.name = "Arial"
                full_title_run.font.size = Pt(11)
                full_title_run.font.bold = True
                full_title_run.font.color.rgb = BrandColors.ORANGE_PRIMARY
                full_title_run.font.underline = False

            # Body paragraph
            body_para = cell.add_paragraph()
            body_para.paragraph_format.space_after = Pt(0)
            body_run = body_para.add_run(body)
            body_run.font.name = "Arial"
            body_run.font.size = Pt(10)
            body_run.font.bold = False
            body_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

            # Add spacer paragraph after box
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(8)

        logger.info("Recommendations section added")

    def _add_sources(self, analysis_result: dict[str, Any]) -> None:
        """Add Resources & Intelligence Sources section with numbered citations."""
        logger.info("Adding Resources & Intelligence Sources section")

        # Heading
        h = self.doc.add_heading("Resources & Intelligence Sources", level=1)
        for run in h.runs:
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.color.rgb = BrandColors.ORANGE_PRIMARY

        # Intro text
        intro = self.doc.add_paragraph()
        intro_run = intro.add_run("This report was compiled using the following intelligence sources:")
        intro_run.font.name = "Arial"
        intro_run.font.size = FontSizes.BODY_SMALL
        intro_run.font.italic = True
        intro_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        intro.paragraph_format.space_after = Pt(6)

        # Primary intelligence sources (numbered [1]-[4])
        # Only list sources that are actually collected/used
        primary_sources = [
            "NIST National Vulnerability Database (NVD)",
            "CISA Known Exploited Vulnerabilities (KEV) Catalog",
            "Intel471 Titan threat intelligence platform",
            "CrowdStrike Falcon Intelligence",
        ]

        for idx, source in enumerate(primary_sources, start=1):
            para = self.doc.add_paragraph(style="List Bullet")
            # Number in bold
            number_run = para.add_run(f"[{idx}] ")
            number_run.font.name = "Arial"
            number_run.font.size = FontSizes.BODY_SMALL
            number_run.font.bold = True
            # Source name
            name_run = para.add_run(source)
            name_run.font.name = "Arial"
            name_run.font.size = FontSizes.BODY_SMALL

        # OSINT Sources heading
        osint_heading = self.doc.add_paragraph()
        osint_heading.paragraph_format.space_before = Pt(12)
        osint_heading.paragraph_format.space_after = Pt(6)
        osint_run = osint_heading.add_run("Open Source Intelligence (OSINT) Sources:")
        osint_run.font.name = "Arial"
        osint_run.font.size = FontSizes.BODY_SMALL
        osint_run.font.bold = True

        # Get OSINT sources from analysis_result
        osint_sources = analysis_result.get("osint_sources_used", [])

        if osint_sources:
            for idx, osint in enumerate(osint_sources, start=5):
                para = self.doc.add_paragraph(style="List Bullet")

                # Number in bold
                number_run = para.add_run(f"[{idx}] ")
                number_run.font.name = "Arial"
                number_run.font.size = FontSizes.BODY_SMALL
                number_run.font.bold = True

                # Title as hyperlink (blue, underlined)
                title = osint.get("title", "Untitled")
                url = osint.get("url", "")
                description = osint.get("description", "")

                if url:
                    # Add title as blue underlined text (hyperlink styling)
                    title_run = para.add_run(title)
                    title_run.font.name = "Arial"
                    title_run.font.size = FontSizes.BODY_SMALL
                    title_run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9C)  # Blue
                    title_run.font.underline = True

                    # Try to add actual hyperlink functionality
                    try:
                        r = title_run._element
                        hyperlink = OxmlElement("w:hyperlink")
                        hyperlink.set(
                            qn("r:id"),
                            self.doc.part.relate_to(
                                url,
                                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                                is_external=True,
                            ),
                        )
                        # Move the run element into the hyperlink
                        new_r = OxmlElement("w:r")
                        for child in list(r):
                            new_r.append(child)
                        hyperlink.append(new_r)
                        # Replace the run with the hyperlink
                        parent = r.getparent()
                        parent.replace(r, hyperlink)
                    except Exception as e:
                        logger.warning(f"Could not create hyperlink for {title}: {e}")
                        # Hyperlink creation failed, but blue underlined text still shows
                else:
                    # No URL, just show title in regular text
                    title_run = para.add_run(title)
                    title_run.font.name = "Arial"
                    title_run.font.size = FontSizes.BODY_SMALL

                # Description in gray italic
                if description:
                    desc_run = para.add_run(f" - {description}")
                    desc_run.font.name = "Arial"
                    desc_run.font.size = FontSizes.BODY_SMALL
                    desc_run.font.italic = True
                    desc_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        else:
            # No OSINT sources, show placeholder
            para = self.doc.add_paragraph(style="List Bullet")
            para_run = para.add_run("No OSINT sources were referenced in this report")
            para_run.font.name = "Arial"
            para_run.font.size = FontSizes.BODY_SMALL
            para_run.font.italic = True
            para_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # Spacer after sources
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    def _add_footer(self) -> None:
        """Add footer with contact info, sources, TLP classification, and page number."""
        logger.info("Adding footer")

        # Contact info
        contact = self.doc.add_paragraph()
        contact_run = contact.add_run(
            f"Questions or suspicious activity: {customer_profile.security_contact} | ServiceNow"
        )
        contact_run.font.name = "Arial"
        contact_run.font.size = FontSizes.BODY_SMALL
        contact_run.font.bold = True
        contact_run.font.color.rgb = BrandColors.GRAY_DARK  # Dark text
        contact_run.font.underline = False  # No underline
        contact.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left align

        # Spacer at end
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # Data sources - removed, now covered by Sources section

        # Add TLP classification and page number to document footer
        section = self.doc.sections[0]
        footer = section.footer

        # Clear any existing footer content
        if footer.paragraphs:
            footer.paragraphs[0].clear()
        else:
            footer.add_paragraph()

        # Create footer paragraph
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # "TLP:" in light gray, italic
        tlp_label = footer_para.add_run("TLP: ")
        tlp_label.font.name = "Arial"
        tlp_label.font.size = FontSizes.FOOTNOTE
        tlp_label.font.italic = True
        tlp_label.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray

        # "AMBER+" in orange, italic
        amber = footer_para.add_run("AMBER+")
        amber.font.name = "Arial"
        amber.font.size = FontSizes.FOOTNOTE
        amber.font.italic = True
        amber.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange

        # "STRICT" in orange, italic, no underline
        strict = footer_para.add_run("STRICT")
        strict.font.name = "Arial"
        strict.font.size = FontSizes.FOOTNOTE
        strict.font.italic = True
        strict.font.color.rgb = BrandColors.ORANGE_PRIMARY  # Orange
        strict.font.underline = False  # No underline

        # Pipe separator in light gray, italic
        pipe = footer_para.add_run(" | ")
        pipe.font.name = "Arial"
        pipe.font.size = FontSizes.FOOTNOTE
        pipe.font.italic = True
        pipe.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray

        # Page number in light gray, italic
        # Use simple text for now - user can manually add page number field in Word if needed
        # This avoids XML corruption issues
        page_text = footer_para.add_run("Page 1")
        page_text.font.name = "Arial"
        page_text.font.size = FontSizes.FOOTNOTE
        page_text.font.italic = True
        page_text.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # Light gray
