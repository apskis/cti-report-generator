"""
Report generator for CTI Weekly Reports.

Generates Word documents from threat intelligence analysis results
and uploads them to Azure Blob Storage.
"""
from docx import Document  # type: ignore
from docx.shared import Pt, RGBColor  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions  # type: ignore
from azure.core.credentials import AzureNamedKeyCredential  # type: ignore
from datetime import datetime, timedelta
from io import BytesIO
import logging
from typing import Dict, Any

from config import report_config

logger = logging.getLogger(__name__)


def generate_report(analysis_result: Dict[str, Any]) -> Document:
    """
    Generate a Word document from threat intelligence analysis results.

    Args:
        analysis_result: Dictionary containing analysis results with keys:
            - executive_summary
            - statistics
            - cve_analysis
            - apt_activity
            - recommendations

    Returns:
        Document object (python-docx Document)
    """
    try:
        logger.info("Generating Word document report")

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading('Cyber Threat Intelligence Weekly Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with current date and week number
        current_date = datetime.now()
        week_number = current_date.isocalendar()[1]
        subtitle = doc.add_paragraph(f'Week {week_number}, {current_date.strftime("%B %d, %Y")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(12)
        subtitle_format.italic = True

        # Add spacing
        doc.add_paragraph()

        # Section 1: Executive Summary
        _add_executive_summary(doc, analysis_result)

        # Section 2: Threat Landscape Overview (Statistics Table)
        _add_statistics_section(doc, analysis_result)

        # Section 3: Critical CVEs
        _add_cve_section(doc, analysis_result)

        # Section 4: APT Activity
        _add_apt_section(doc, analysis_result)

        # Section 5: Recommendations
        _add_recommendations_section(doc, analysis_result)

        logger.info("Word document generated successfully")
        return doc

    except Exception as e:
        logger.error(f"Error generating report: {str(e)}", exc_info=True)
        raise


def _add_executive_summary(doc: Document, analysis_result: Dict[str, Any]) -> None:
    """Add executive summary section to document."""
    logger.info("Adding Executive Summary section")
    doc.add_heading('Executive Summary', 1)
    exec_summary = analysis_result.get("executive_summary", "No executive summary available.")
    doc.add_paragraph(exec_summary)
    doc.add_paragraph()


def _add_statistics_section(doc: Document, analysis_result: Dict[str, Any]) -> None:
    """Add threat landscape statistics section to document."""
    logger.info("Adding Threat Landscape Overview section")
    doc.add_heading('Threat Landscape Overview', 1)

    stats = analysis_result.get("statistics", {})
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = report_config.table_style

    # Header row
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Count'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True

    # Add statistics rows
    metrics = [
        ("Total CVEs", stats.get("total_cves", 0)),
        ("Critical Count", stats.get("critical_count", 0)),
        ("High Count", stats.get("high_count", 0)),
        ("Exploited in Wild", stats.get("exploited_count", 0)),
        ("APT Groups", stats.get("apt_groups", 0)),
        ("P1 Count", stats.get("p1_count", 0)),
        ("P2 Count", stats.get("p2_count", 0)),
        ("P3 Count", stats.get("p3_count", 0))
    ]

    for metric, count in metrics:
        row_cells = stats_table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = str(count)

    doc.add_paragraph()


def _add_cve_section(doc: Document, analysis_result: Dict[str, Any]) -> None:
    """Add CVE analysis section to document."""
    logger.info("Adding Critical CVEs section")
    doc.add_heading('Critical CVEs', 1)

    cve_analysis = analysis_result.get("cve_analysis", [])
    if cve_analysis:
        cve_table = doc.add_table(rows=1, cols=4)
        cve_table.style = report_config.table_style

        # Header row
        header_cells = cve_table.rows[0].cells
        headers = ["CVE ID", "Severity", "Priority", "Exploited"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Add CVE rows with color coding
        for cve in cve_analysis:
            row_cells = cve_table.add_row().cells
            row_cells[0].text = cve.get("cve_id", "N/A")

            severity = cve.get("severity", "N/A")
            row_cells[1].text = severity

            priority = cve.get("priority", "N/A")
            row_cells[2].text = priority

            exploited = "Yes" if cve.get("exploited", False) else "No"
            row_cells[3].text = exploited

            # Color coding for severity
            _apply_severity_color(row_cells[1], severity)

            # P1 = Yellow background
            if priority == "P1":
                _apply_p1_highlighting(row_cells)
    else:
        doc.add_paragraph("No CVE data available.")

    doc.add_paragraph()


def _apply_severity_color(cell, severity: str) -> None:
    """Apply color coding to severity cell."""
    if not cell.paragraphs[0].runs:
        return

    run = cell.paragraphs[0].runs[0]
    if severity.upper() == "CRITICAL":
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    elif severity.upper() == "HIGH":
        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange


def _apply_p1_highlighting(row_cells) -> None:
    """Apply P1 yellow highlighting to row cells."""
    for cell in row_cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.find(qn('w:shd'))
        if shading_elm is None:
            shading_elm = shading.makeelement(qn('w:shd'))
            shading.append(shading_elm)
        shading_elm.set(qn('w:fill'), 'FFFF00')  # Yellow


def _add_apt_section(doc: Document, analysis_result: Dict[str, Any]) -> None:
    """Add APT activity section to document."""
    logger.info("Adding APT Activity section")
    doc.add_heading('APT Activity', 1)

    apt_activity = analysis_result.get("apt_activity", [])
    if apt_activity:
        for apt in apt_activity:
            actor = apt.get("actor", "Unknown")
            country = apt.get("country", "Unknown")
            motivation = apt.get("motivation", "Unknown")
            ttps = apt.get("ttps", [])

            # Format: Actor Name (Country) - Motivation - TTPs listed
            apt_text = f"{actor} ({country}) - {motivation}"
            if ttps:
                if isinstance(ttps, list):
                    ttps_str = ", ".join(ttps)
                else:
                    ttps_str = str(ttps)
                apt_text += f" - TTPs: {ttps_str}"

            doc.add_paragraph(apt_text, style='List Bullet')
    else:
        doc.add_paragraph("No APT activity data available.")

    doc.add_paragraph()


def _add_recommendations_section(doc: Document, analysis_result: Dict[str, Any]) -> None:
    """Add recommendations section to document."""
    logger.info("Adding Recommendations section")
    doc.add_heading('Recommendations', 1)

    recommendations = analysis_result.get("recommendations", [])
    if recommendations:
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
    else:
        doc.add_paragraph("No recommendations available.")


def upload_to_blob(
    document: Document,
    storage_account_name: str,
    storage_account_key: str,
    container_name: str | None = None
) -> str:
    """
    Upload a Word document to Azure Blob Storage and generate a SAS URL.

    Uses Azure SDK properly instead of building connection strings manually.

    Args:
        document: python-docx Document object
        storage_account_name: Azure Storage account name
        storage_account_key: Azure Storage account key
        container_name: Blob container name (default from config)

    Returns:
        Public SAS URL string
    """
    container_name = container_name or report_config.container_name

    try:
        logger.info(f"Uploading document to Azure Blob Storage: {storage_account_name}/{container_name}")

        # Save document to BytesIO buffer
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Create BlobServiceClient using account URL (safer than connection string)
        account_url = f"https://{storage_account_name}.blob.core.windows.net"
        credential = AzureNamedKeyCredential(storage_account_name, storage_account_key)
        blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)

        # Generate filename
        filename = f"CTI_Weekly_Report_{datetime.now().strftime('%Y-%m-%d')}.docx"

        # Get container client (create if doesn't exist)
        container_client = blob_service_client.get_container_client(container_name)
        try:
            container_client.create_container()
            logger.info(f"Created container: {container_name}")
        except Exception:
            # Container likely already exists
            pass

        # Upload blob
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=filename)
        blob_client.upload_blob(buffer.getvalue(), overwrite=True)
        logger.info(f"Uploaded blob: {filename}")

        # Generate SAS URL with configurable expiry
        sas_token = generate_blob_sas(
            account_name=storage_account_name,
            container_name=container_name,
            blob_name=filename,
            account_key=storage_account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(days=report_config.sas_expiry_days)
        )

        sas_url = f"{account_url}/{container_name}/{filename}?{sas_token}"
        logger.info(f"Generated SAS URL (valid for {report_config.sas_expiry_days} days)")

        return sas_url

    except Exception as e:
        logger.error(f"Error uploading to blob storage: {str(e)}", exc_info=True)
        raise


def create_and_upload_report(
    analysis_result: Dict[str, Any],
    storage_account_name: str,
    storage_account_key: str
) -> Dict[str, Any]:
    """
    Generate a Word document report and upload it to Azure Blob Storage.

    Args:
        analysis_result: Dictionary containing analysis results
        storage_account_name: Azure Storage account name
        storage_account_key: Azure Storage account key

    Returns:
        Dictionary with keys: filename, url, success
    """
    try:
        logger.info("Creating and uploading report")

        # Generate report
        document = generate_report(analysis_result)

        # Generate filename
        filename = f"CTI_Weekly_Report_{datetime.now().strftime('%Y-%m-%d')}.docx"

        # Upload to blob storage
        url = upload_to_blob(document, storage_account_name, storage_account_key)

        logger.info(f"Report created and uploaded successfully: {filename}")

        return {
            "filename": filename,
            "url": url,
            "success": True
        }

    except Exception as e:
        logger.error(f"Error creating and uploading report: {str(e)}", exc_info=True)
        return {
            "filename": None,
            "url": None,
            "success": False,
            "error": str(e)
        }
