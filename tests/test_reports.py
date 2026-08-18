"""
Tests for report generators.
"""

import pytest

from src.reports.base import BrandColors, FontSizes
from src.reports.quarterly_report import QuarterlyReportGenerator
from src.reports.registry import (
    REPORT_REGISTRY,
    get_report_generator,
    list_report_types,
)
from src.reports.weekly_report import WeeklyReportGenerator


@pytest.fixture(autouse=True)
def _isolate_quarterly_history(monkeypatch, tmp_path):
    """Keep the quarterly history-file side effect out of the working tree (Q25).

    Quarterly generate() persists risk history; redirect it to a temp dir so tests
    never write data/historical/ into the repo checkout.
    """
    monkeypatch.setenv("QUARTERLY_HISTORY_DIR", str(tmp_path / "quarterly_hist"))


def _get_document_text(doc):
    """Collect all text from document paragraphs and table cells."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


class TestBrandColors:
    """Tests for brand color constants."""

    def test_orange_primary_value(self):
        assert BrandColors.ORANGE_PRIMARY is not None

    def test_gray_colors_defined(self):
        assert BrandColors.GRAY_DARK is not None
        assert BrandColors.GRAY_MEDIUM is not None

    def test_severity_colors_defined(self):
        assert BrandColors.RED_CRITICAL is not None
        assert BrandColors.ORANGE_HIGH is not None


class TestFontSizes:
    """Tests for font size constants."""

    def test_title_size(self):
        assert FontSizes.TITLE.pt == 18

    def test_body_size(self):
        assert FontSizes.BODY.pt == 10.5

    def test_subtitle_size(self):
        assert FontSizes.SUBTITLE.pt == 9


class TestReportRegistry:
    """Tests for the report registry."""

    def test_weekly_report_registered(self):
        """WeeklyReportGenerator should be auto-registered on import."""
        assert "weekly" in REPORT_REGISTRY
        assert REPORT_REGISTRY["weekly"] == WeeklyReportGenerator

    def test_get_report_generator_weekly(self):
        """get_report_generator should return WeeklyReportGenerator for 'weekly'."""
        generator = get_report_generator("weekly")
        assert generator is not None
        assert isinstance(generator, WeeklyReportGenerator)

    def test_get_report_generator_case_insensitive(self):
        """get_report_generator should be case-insensitive."""
        generator = get_report_generator("WEEKLY")
        assert generator is not None
        assert isinstance(generator, WeeklyReportGenerator)

    def test_get_report_generator_unknown_type(self):
        """get_report_generator should return None for unknown types."""
        generator = get_report_generator("unknown_type")
        assert generator is None

    def test_list_report_types(self):
        """list_report_types should include 'weekly'."""
        types = list_report_types()
        assert "weekly" in types


class TestWeeklyReportGenerator:
    """Tests for the WeeklyReportGenerator."""

    @pytest.fixture
    def generator(self):
        """Create a fresh generator instance."""
        return WeeklyReportGenerator()

    @pytest.fixture
    def sample_analysis_result(self):
        """Sample analysis result for testing."""
        return {
            "executive_summary": "This week we identified 5 new vulnerabilities.",
            "statistics": {
                "total_cves": 10,
                "critical_count": 2,
                "high_count": 3,
                "exploited_count": 1,
                "apt_groups": 2,
                "new_this_week": 5,
                "persistent_count": 3,
                "resolved_count": 2,
            },
            "cve_analysis": [
                {
                    "cve_id": "CVE-2026-1234",
                    "affected_product": "TestApp",
                    "exposure": "Remote code execution",
                    "exploited_by": "APT29",
                    "risk": "CRITICAL",
                    "weeks_detected": 4,
                },
                {
                    "cve_id": "CVE-2026-5678",
                    "affected_product": "TestLib",
                    "exposure": "SQL injection",
                    "exploited_by": "None known",
                    "risk": "HIGH",
                    "weeks_detected": 1,
                },
            ],
            "apt_activity": [
                {
                    "actor": "APT29",
                    "country": "Russia",
                    "motivation": "Espionage",
                    "activity": "Targeting healthcare organizations",
                    "ttps": ["T1566", "T1059", "T1027"],
                    "what_to_monitor": "Phishing emails with healthcare themes",
                },
            ],
            "recommendations": [
                "Patch CVE-2026-1234 immediately",
                "Review access controls for TestApp",
                "Enable MFA for all admin accounts",
            ],
            "exploitation_indicators": [
                "CVE-2026-1234 (TestApp): Unusual outbound connections on port 443",
            ],
        }

    def test_report_type(self, generator):
        """report_type should return 'weekly'."""
        assert generator.report_type == "weekly"

    def test_filename_prefix(self, generator):
        """filename_prefix should return 'CTI_Weekly_Report'."""
        assert generator.filename_prefix == "CTI_Weekly_Report"

    def test_get_filename_format(self, generator):
        """get_filename should return properly formatted filename."""
        filename = generator.get_filename()
        assert filename.startswith("CTI_Weekly_Report_")
        assert filename.endswith(".docx")
        # Weekly reports are named by ISO year and week number, e.g. CTI_Weekly_Report_2026_Week30.docx
        assert "_Week" in filename
        assert str(generator.created_at.isocalendar()[0]) in filename

    def test_generate_creates_document(self, generator, sample_analysis_result):
        """generate should create a valid Document object."""
        doc = generator.generate(sample_analysis_result)
        assert doc is not None
        assert generator.doc is not None

    def test_generate_with_empty_data(self, generator):
        """generate should handle empty analysis result gracefully."""
        doc = generator.generate({})
        assert doc is not None

    def test_to_bytes_after_generate(self, generator, sample_analysis_result):
        """to_bytes should return bytes after generate is called."""
        generator.generate(sample_analysis_result)
        doc_bytes = generator.to_bytes()
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        # DOCX files start with PK (zip signature)
        assert doc_bytes[:2] == b"PK"

    def test_to_bytes_before_generate_raises(self, generator):
        """to_bytes should raise if generate wasn't called."""
        with pytest.raises(ValueError, match="Document not generated"):
            generator.to_bytes()

    def test_week_calculation(self, generator, sample_analysis_result):
        """Lookback period dates should be calculated correctly."""
        generator.generate(sample_analysis_result)
        assert hasattr(generator, "period_start")
        assert hasattr(generator, "period_end")
        delta = generator.period_end - generator.period_start
        assert delta.days == generator.lookback_days - 1

    def test_document_has_paragraphs(self, generator, sample_analysis_result):
        """Generated document should have paragraphs."""
        doc = generator.generate(sample_analysis_result)
        assert len(doc.paragraphs) > 0

    def test_document_has_tables(self, generator, sample_analysis_result):
        """Generated document should have tables (metric cards, CVE table, etc.)."""
        doc = generator.generate(sample_analysis_result)
        assert len(doc.tables) > 0

    def test_document_contains_title(self, generator, sample_analysis_result):
        """Document should contain the report title (may be in a paragraph or table cell)."""
        doc = generator.generate(sample_analysis_result)
        text_content = _get_document_text(doc)
        assert "Cyber Threat Intelligence Weekly Report" in text_content

    def test_document_contains_executive_summary(self, generator, sample_analysis_result):
        """Document should contain executive summary section."""
        doc = generator.generate(sample_analysis_result)
        text_content = _get_document_text(doc)
        assert "Summary" in text_content
        assert "This week we identified 5 new vulnerabilities" in text_content

    def test_document_omits_recommended_actions(self, generator, sample_analysis_result):
        """The weekly report no longer includes a Recommended Actions section."""
        doc = generator.generate(sample_analysis_result)
        text_content = _get_document_text(doc)
        assert "Recommended Actions" not in text_content


class TestBaseReportGenerator:
    """Tests for BaseReportGenerator utility methods."""

    @pytest.fixture
    def generator(self):
        return WeeklyReportGenerator()

    def test_get_week_number(self, generator):
        """_get_week_number should return valid ISO week number."""
        generator._calculate_date_range()
        week = generator._get_week_number()
        assert 1 <= week <= 53

    def test_get_year(self, generator):
        """_get_year should return a valid year."""
        generator._calculate_date_range()
        year = generator._get_year()
        assert year == generator.period_end.isocalendar()[0]

    def test_format_date_range(self, generator):
        """_format_date_range should produce readable date range."""
        date_range = generator._format_date_range()
        assert "to" in date_range
        # Should contain month name
        assert any(
            month in date_range
            for month in [
                "January",
                "February",
                "March",
                "April",
                "May",
                "June",
                "July",
                "August",
                "September",
                "October",
                "November",
                "December",
            ]
        )


class TestQuarterlyReportGenerator:
    """Tests for the QuarterlyReportGenerator."""

    @pytest.fixture
    def generator(self):
        """Create a fresh generator instance."""
        return QuarterlyReportGenerator()

    @pytest.fixture
    def sample_strategic_analysis(self):
        """Sample strategic analysis result for testing."""
        return {
            "executive_summary": "The threat landscape remained elevated throughout the quarter.",
            "risk_assessment": {
                "nation_state": "HIGH",
                "nation_state_trend": "↑",
                "ransomware": "HIGH",
                "ransomware_trend": "Unchanged",
                "supply_chain": "MEDIUM",
                "supply_chain_trend": "Unchanged",
                "insider": "LOW",
                "insider_trend": "Unchanged",
            },
            "breach_landscape": {
                "total_incidents": 47,
                "prev_total_incidents": 36,
                "total_impact_millions": 127,
                "prev_total_impact": 89,
                "ransomware_count": 18,
                "prev_ransomware": 12,
                "records_exposed_millions": 4.2,
                "prev_records": 2.8,
            },
            "incidents_by_type": [
                {
                    "type": "Ransomware",
                    "current_count": 18,
                    "prev_count": 12,
                    "notable_example": "Pharma manufacturer: 12-day production halt, FDA notification",
                },
                {
                    "type": "Data Theft / Exfiltration",
                    "current_count": 11,
                    "prev_count": 9,
                    "notable_example": "Genomics institute: 2.3M patient samples accessed",
                },
                {
                    "type": "Manufacturing / OT Disruption",
                    "current_count": 5,
                    "prev_count": 3,
                    "notable_example": "Medical device mfg: assembly line shutdown, 8-day recovery",
                },
                {
                    "type": "Business Email Compromise",
                    "current_count": 6,
                    "prev_count": 5,
                    "notable_example": "CRO: $3.8M fraudulent wire transfers",
                },
                {
                    "type": "Third-Party / Vendor",
                    "current_count": 4,
                    "prev_count": 4,
                    "notable_example": "Lab software vendor: credentials exposed for 200+ customers",
                },
                {
                    "type": "Unauthorized Access",
                    "current_count": 3,
                    "prev_count": 3,
                    "notable_example": "Biotech: former employee accessed IP post-termination",
                },
            ],
            "common_factors": "Unpatched systems (34%), compromised credentials (28%)",
            "geopolitical_threats": [
                {
                    "country": "China",
                    "threat_level": "HIGH",
                    "relevance": ["Strategic interest in biotech and genomics IP"],
                    "activity": ["APT41 conducted multiple intrusions"],
                    "risk": ["IP theft risk for proprietary research"],
                },
                {
                    "country": "Russia",
                    "threat_level": "MEDIUM",
                    "relevance": ["Ransomware ecosystem targeting manufacturing"],
                    "activity": ["Ransomware incidents increased 31%"],
                    "risk": ["Operational disruption risk"],
                },
                {
                    "country": "North Korea",
                    "threat_level": "MEDIUM",
                    "relevance": ["Dual-purpose revenue and espionage operations"],
                    "activity": ["LinkedIn social engineering campaigns"],
                    "risk": ["Credential compromise risk"],
                },
            ],
            "looking_ahead": {
                "threat_outlook": "Continued pressure from state-sponsored campaigns",
                "planned_initiatives": "Enhanced detection capabilities",
                "watch_items": [
                    {"subject": "Industry events", "detail": "Major industry events and announcements"},
                    {"subject": "Regulatory shifts", "detail": "New data-protection rules in key markets"},
                ],
            },
            "recommendations": {
                "intro_note": "Priority actions for the coming quarter.",
                "items": [
                    {"title": "Executive Awareness", "body": "Targeted security awareness for executives"},
                    {"title": "Vendor Risk Review", "body": "Evaluate vendor security posture"},
                ],
            },
        }

    def test_report_type(self, generator):
        """report_type should return 'quarterly'."""
        assert generator.report_type == "quarterly"

    def test_filename_prefix(self, generator):
        """filename_prefix should return 'CTI_Quarterly_Strategic_Brief'."""
        assert generator.filename_prefix == "CTI_Quarterly_Strategic_Brief"

    def test_get_filename_format(self, generator):
        """get_filename should return properly formatted filename."""
        filename = generator.get_filename()
        assert filename.startswith("CTI_Quarterly_Strategic_Brief_")
        assert filename.endswith(".docx")

    def test_get_filename_uses_quarter_not_week(self, generator):
        """A pinned period names the file by quarter (Q2_2026), never by ISO week."""
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        filename = generator.get_filename()
        assert filename == "CTI_Quarterly_Strategic_Brief_Q2_2026.docx"
        assert "Week" not in filename

    def test_quarterly_registered(self):
        """QuarterlyReportGenerator should be registered."""
        assert "quarterly" in REPORT_REGISTRY
        assert REPORT_REGISTRY["quarterly"] == QuarterlyReportGenerator

    def test_get_report_generator_quarterly(self):
        """get_report_generator should return QuarterlyReportGenerator for 'quarterly'."""
        generator = get_report_generator("quarterly")
        assert generator is not None
        assert isinstance(generator, QuarterlyReportGenerator)

    def test_generate_creates_document(self, generator, sample_strategic_analysis):
        """generate should create a valid Document object."""
        doc = generator.generate(sample_strategic_analysis)
        assert doc is not None
        assert generator.doc is not None

    def test_generate_with_empty_data(self, generator):
        """generate should handle empty analysis result gracefully."""
        doc = generator.generate({})
        assert doc is not None

    def test_to_bytes_after_generate(self, generator, sample_strategic_analysis):
        """to_bytes should return bytes after generate is called."""
        generator.generate(sample_strategic_analysis)
        doc_bytes = generator.to_bytes()
        assert isinstance(doc_bytes, bytes)
        assert len(doc_bytes) > 0
        # DOCX files start with PK (zip signature)
        assert doc_bytes[:2] == b"PK"

    def test_quarter_calculation(self, generator, sample_strategic_analysis):
        """Quarter and lookback period should be calculated correctly."""
        generator.generate(sample_strategic_analysis)
        assert 1 <= generator.quarter <= 4
        assert hasattr(generator, "period_start")
        assert hasattr(generator, "period_end")
        delta = generator.period_end - generator.period_start
        assert delta.days == generator.lookback_days

    def test_document_has_paragraphs(self, generator, sample_strategic_analysis):
        """Generated document should have paragraphs."""
        doc = generator.generate(sample_strategic_analysis)
        assert len(doc.paragraphs) > 0

    def test_document_has_tables(self, generator, sample_strategic_analysis):
        """Generated document should have tables (risk cards, breach stats, etc.)."""
        doc = generator.generate(sample_strategic_analysis)
        assert len(doc.tables) > 0

    def test_document_contains_title(self, generator, sample_strategic_analysis):
        """Document should contain the report title."""
        doc = generator.generate(sample_strategic_analysis)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Quarterly Strategic Brief" in text_content

    def test_document_contains_executive_summary(self, generator, sample_strategic_analysis):
        """Document should contain executive summary section."""
        doc = generator.generate(sample_strategic_analysis)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Executive Summary" in text_content
        assert "elevated" in text_content

    def test_document_contains_geopolitical_section(self, generator, sample_strategic_analysis):
        """Document should contain geopolitical threat landscape."""
        doc = generator.generate(sample_strategic_analysis)
        text_content = _get_document_text(doc)
        assert "Geopolitical Threat Landscape" in text_content
        # Should have country sections (rendered as a per-country table)
        assert "China" in text_content
        assert "Russia" in text_content

    def test_document_contains_recommendations(self, generator, sample_strategic_analysis):
        """Document should contain recommendations for leadership."""
        doc = generator.generate(sample_strategic_analysis)
        text_content = _get_document_text(doc)
        assert "Recommendations" in text_content
        assert "Executive Awareness" in text_content


# =============================================================================
# Quarterly robustness (Part 5): malformed AI output must degrade, not crash;
# plus regression tests for the display bugs (Q17 badge, Q18 trend, Q26 cards, Q23).
# =============================================================================


class TestQuarterlyRobustness:
    @pytest.fixture
    def generator(self):
        return QuarterlyReportGenerator()

    # ----- Part 5: malformed strategic analysis degrades without raising -----

    @pytest.mark.parametrize(
        "analysis",
        [
            {"risk_assessment": None},
            {"risk_assessment": {"nation_state": None, "ransomware": None}},
            {"geopolitical_threats": ["China", "Russia"]},
            {"looking_ahead": None},
            {
                "breach_landscape": {
                    "stat_cards": [
                        {"value": 20, "label": "Total", "prior_value": "N/A", "change_pct": "N/A"}
                    ],
                    "incidents_by_type": [
                        {"type": "Ransomware", "current_count": 12, "prev_count": 10, "notable_example": "Acme: hit"}
                    ],
                }
            },
            {"breach_landscape": {"stat_cards": [{"value": 1, "change_pct": "+high%"}]}},
            {"incidents_by_type": [{"type": "X", "current_count": 3, "notable_example": "Y: z"}]},
            # Aug 2026 review (T1.2) — sibling fields that were unguarded and crashed the .docx:
            {"breach_landscape": {"incidents_by_type": ["Ransomware at Covenant", "Hacking at X"]}},
            {"recommendations": ["Patch immediately", "Audit third-party vendors"]},
            {"recommendations": {"items": ["Patch immediately", "Audit vendors"]}},
            {"geopolitical_threats": [{"name": "China", "exposure": None, "threat_level": None}]},
        ],
    )
    def test_malformed_strategic_analysis_does_not_raise(self, generator, analysis):
        doc = generator.generate(analysis)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    # ----- T2.1: change_pct color contract (red up / green down / gray flat|N/A) -----

    def test_change_pct_color_by_direction(self, generator):
        from docx.shared import RGBColor

        red, green, gray = RGBColor(0xDC, 0x35, 0x45), RGBColor(0x1A, 0x7F, 0x37), RGBColor(0x6B, 0x72, 0x80)
        f = generator._change_pct_color
        assert f("+25%") == red and f("25%") == red
        assert f("-40%") == green and f("−40%") == green and f("-5") == green
        assert f("0%") == gray and f("N/A") == gray and f("") == gray and f("high") == gray

    # ----- T2.4: a fabricated QoQ delta with no prior value is scrubbed to N/A -----

    def test_fabricated_qoq_delta_without_prior_is_scrubbed(self, generator):
        # Model emitted a change_pct with no prior_value and there is no stored baseline in tests.
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"label": "Total Incidents", "value": "16", "prior_value": "N/A", "change_pct": "+25%"}]
            }
        }
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        assert "+25%" not in text  # the invented delta must not render

    # ----- T2.3: ungrounded breach stat cards are marked, not presented as authoritative -----

    def test_ungrounded_stat_cards_get_methodology_marker(self, generator):
        analysis = {
            "breach_landscape": {"stat_cards": [{"label": "Total Incidents", "value": "42", "change_pct": "N/A"}]}
        }
        # No breach_dataset set on the generator -> grounding cannot run.
        generator._apply_breach_dataset_grounding(analysis)
        note = analysis["breach_landscape"].get("stat_methodology", "")
        assert "NOT" in note and "grounded" in note

    # ----- Aug 2026 review (B2): a bare-list recommendations payload renders, not "unavailable" -----

    def test_bare_list_recommendations_render(self, generator):
        analysis = {"recommendations": ["Patch the ServiceNow instances this week", "Rotate exposed keys"]}
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        assert "Patch the ServiceNow instances this week" in text
        assert "No recommendations generated" not in text

    # ----- Q18: risk-card trend arrows map to a direction, not "Unchanged" -----

    def test_trend_arrows_render_direction(self, generator):
        analysis = {
            "risk_assessment": {
                "nation_state": "HIGH",
                "nation_state_trend": "↑",
                "ransomware": "MEDIUM",
                "ransomware_trend": "↓",
                "supply_chain": "LOW",
                "supply_chain_trend": "Unchanged",
                "insider": "LOW",
                "insider_trend": "Unchanged",
            }
        }
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        assert "Increased" in text
        assert "Decreased" in text

    # ----- Q17: geopolitical badge honors the AI's "level"/"threat_level" -----

    def test_threat_level_badge_reads_level_key(self, generator):
        analysis = {
            "geopolitical_threats": [
                {"name": "China", "level": "CRITICAL", "relevance": ["x"], "activity": ["y"], "risk": ["z"]}
            ]
        }
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        # The real assessed level must appear, not a hardcoded MEDIUM default.
        assert "CRITICAL" in text

    # ----- Q26: stat cards render for counts other than exactly 4 -----

    @pytest.mark.parametrize("num_cards", [1, 2, 3, 4])
    def test_stat_cards_render_for_varied_counts(self, generator, num_cards):
        cards = [
            {"value": str(i), "label": f"Metric {i}", "prior_value": "N/A", "change_pct": "N/A"}
            for i in range(num_cards)
        ]
        analysis = {"breach_landscape": {"stat_cards": cards, "incidents_by_type": []}}
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        for i in range(num_cards):
            assert f"Metric {i}" in text

    def test_na_prior_value_does_not_render_fabricated_percent(self, generator):
        cards = [{"value": "20", "label": "Total Incidents", "prior_label": "Q1 2026", "prior_value": "N/A", "change_pct": "N/A"}]
        analysis = {"breach_landscape": {"stat_cards": cards, "incidents_by_type": []}}
        doc = generator.generate(analysis)
        text = _get_document_text(doc)
        assert "Total Incidents" in text
        # No parenthetical percentage when there is no real prior data.
        assert "(N/A)" not in text

    # ----- Q23: inline citations are superscripted document-wide in quarterly -----

    def test_quarterly_citations_are_subscripted(self, generator):
        analysis = {"executive_summary": "Breach at Stadler Rail [1] and follow-up [2] noted."}
        doc = generator.generate(analysis)
        superscript_runs = [r.text for p in doc.paragraphs for r in p.runs if r.font.superscript]
        assert "[1]" in superscript_runs
        assert "[2]" in superscript_runs

    # ----- Explicit reporting period pins the quarter and overrides AI labels -----

    def test_explicit_period_sets_exact_window(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        generator.generate({"executive_summary": "x"})
        assert generator.quarter == 2
        assert generator._get_year() == 2026
        assert generator.period_start.date().isoformat() == "2026-04-01"
        assert generator.period_end.date().isoformat() == "2026-06-30"

    def test_explicit_period_overrides_ai_quarter_labels(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        analysis = {
            "executive_summary": "x",
            "breach_landscape": {
                "current_quarter_label": "Q3 2026",  # AI guessed wrong
                "prior_quarter_label": "Q2 2026",
                "stat_cards": [
                    {"value": "20", "label": "Total Incidents", "prior_label": "Q4 2025",
                     "prior_value": "N/A", "change_pct": "N/A"}
                ],
                "incidents_by_type": [],
            },
        }
        generator.generate(analysis)
        bl = analysis["breach_landscape"]
        assert bl["current_quarter_label"] == "Q2 2026"
        assert bl["prior_quarter_label"] == "Q1 2026"
        assert bl["stat_cards"][0]["prior_label"] == "Q1 2026"

    # ----- Prior-quarter baseline: history stores + drives stat-card prior column -----

    @pytest.mark.parametrize(
        "value,expected",
        [
            ("20", 20.0),
            ("1,000", 1000.0),
            ("$1.2M", 1.2),
            ("5.0M", 5.0),
            ("N/A", None),
            ("", None),
            ("—", None),
            ("high", None),
        ],
    )
    def test_parse_stat_number(self, generator, value, expected):
        assert generator._parse_stat_number(value) == expected

    def test_persist_baseline_writes_risk_and_breach_stats(self, generator):
        from src.core.reporting_period import make_period

        analysis = {
            "risk_assessment": {"nation_state": "high", "ransomware": "medium"},
            "breach_landscape": {
                "stat_cards": [
                    {"value": "10", "label": "Total Incidents"},
                    {"value": "3", "label": "Ransomware"},
                ]
            },
        }
        generator.persist_baseline_from_analysis(analysis, make_period(2026, "Q1"))
        history = generator._load_historical_data()
        assert "2026-Q1" in history
        assert history["2026-Q1"]["nation_state"] == "HIGH"
        assert history["2026-Q1"]["breach_stats"]["Total Incidents"] == "10"

    def test_prior_quarter_stats_fill_card_and_compute_pct(self, generator):
        from src.core.reporting_period import make_period

        # Seed Q1 2026 with a real baseline of 10 Total Incidents.
        generator.persist_baseline_from_analysis(
            {"breach_landscape": {"stat_cards": [{"value": "10", "label": "Total Incidents"}]}},
            make_period(2026, "Q1"),
        )
        # Now render Q2 2026 with 20 incidents — the prior column and % should come from Q1.
        generator.set_reporting_period(make_period(2026, "Q2"))
        analysis = {
            "breach_landscape": {
                "stat_cards": [
                    {"value": "20", "label": "Total Incidents", "prior_value": "N/A", "change_pct": "N/A"}
                ],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        card = analysis["breach_landscape"]["stat_cards"][0]
        assert card["prior_value"] == "10"  # real prior, not N/A
        assert card["change_pct"] == "+100%"  # 10 -> 20

    # ----- #1: Ransomware card is grounded in the observed incident-type count -----

    def test_ransomware_card_derived_from_incident_types(self, generator):
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"value": "N/A", "label": "Ransomware", "prior_value": "N/A", "change_pct": "N/A"}],
                "incidents_by_type": [
                    {"type": "Ransomware", "current_count": 5, "notable_example": "Acme: hit"},
                    {"type": "Ransomware & Extortion", "current_count": 2, "notable_example": "B: hit"},
                    {"type": "Data Exposure", "current_count": 3, "notable_example": "C: leak"},
                ],
            }
        }
        generator.generate(analysis)
        card = analysis["breach_landscape"]["stat_cards"][0]
        assert card["value"] == "7"  # 5 + 2 ransomware incidents, not N/A

    def test_ransomware_card_untouched_when_no_ransomware_type(self, generator):
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"value": "N/A", "label": "Ransomware"}],
                "incidents_by_type": [{"type": "Data Exposure", "current_count": 3, "notable_example": "C: leak"}],
            }
        }
        generator.generate(analysis)
        assert analysis["breach_landscape"]["stat_cards"][0]["value"] == "N/A"

    # ----- #2: honest zero-prior change labels -----

    @pytest.mark.parametrize(
        "current,prior,expected",
        [
            ("20", "10", "+100%"),
            ("6", "10", "-40%"),
            ("10", "10", "0%"),
            ("0", "0", "0%"),
            ("5", "0", "New"),   # prior zero, current nonzero -> New, not a bogus percent
            ("N/A", "10", "N/A"),  # current absent -> N/A
            ("5", "N/A", "N/A"),   # no baseline -> N/A
        ],
    )
    def test_compute_change_pct(self, generator, current, prior, expected):
        assert generator._compute_change_pct(current, prior) == expected

    def test_new_label_when_prior_zero(self, generator):
        from src.core.reporting_period import make_period

        generator.persist_baseline_from_analysis(
            {"breach_landscape": {"stat_cards": [{"value": "0", "label": "Ransomware"}]}},
            make_period(2026, "Q1"),
        )
        generator.set_reporting_period(make_period(2026, "Q2"))
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"value": "3", "label": "Ransomware", "prior_value": "N/A", "change_pct": "N/A"}],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        card = analysis["breach_landscape"]["stat_cards"][0]
        assert card["prior_value"] == "0"
        assert card["change_pct"] == "New"

    # ----- Breach-dataset grounding drives the stat cards from real incidents -----

    def test_breach_dataset_grounds_stat_cards(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        generator.set_breach_dataset(
            [
                {"organization": "Covenant Health", "date": "2026-04-10", "incident_type": "Ransomware",
                 "records_exposed": 1_000_000, "source": "HHS"},
                {"organization": "LabCorp", "date": "2026-05-02", "incident_type": "Hacking",
                 "records_exposed": 500_000, "source": "HIBP"},
                {"organization": "OutOfPeriod", "date": "2026-01-01", "incident_type": "Ransomware",
                 "records_exposed": 9_000_000, "source": "VCDB"},  # Q1 -> excluded
            ]
        )
        analysis = {
            "breach_landscape": {
                "stat_cards": [
                    {"label": "Total Incidents", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                    {"label": "Est. Total Impact", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                    {"label": "Ransomware", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                    {"label": "Records Exposed", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                ],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        cards = {c["label"]: c["value"] for c in analysis["breach_landscape"]["stat_cards"]}
        assert cards["Total Incidents"] == "2"  # only Q2 incidents
        assert cards["Ransomware"] == "1"
        assert cards["Records Exposed"] == "1.5M"  # 1.0M + 0.5M
        assert cards["Est. Total Impact"] != "N/A"  # records x per-record cost

    def test_breach_dataset_grounds_incidents_by_type_examples(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        generator.set_breach_dataset(
            [
                {"organization": "Covenant Health", "date": "2026-04-10", "incident_type": "Ransomware",
                 "records_exposed": 1_200_000, "source": "HHS"},
                {"organization": "LabCorp", "date": "2026-05-02", "incident_type": "Hacking",
                 "records_exposed": 500_000, "source": "HIBP"},
            ]
        )
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"label": "Total Incidents", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"}],
                "incidents_by_type": [{"type": "Placeholder", "current_count": 99, "notable_example": "x"}],
            }
        }
        generator.generate(analysis)
        text = _get_document_text(generator.doc)
        assert "Covenant Health" in text  # real named example from the dataset
        types = [row["type"] for row in analysis["breach_landscape"]["incidents_by_type"]]
        assert "Placeholder" not in types  # AI table replaced by grounded incidents

    def test_breach_dataset_lag_keeps_live_count(self, generator):
        # Live/AI count is 20; dataset has only 2 in-period incidents -> keep 20, enrich $/records.
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        generator.set_breach_dataset(
            [
                {"organization": "A", "date": "2026-04-10", "incident_type": "Ransomware",
                 "records_exposed": 1_000_000, "source": "HHS"},
                {"organization": "B", "date": "2026-05-10", "incident_type": "Hacking",
                 "records_exposed": 500_000, "source": "HIBP"},
            ]
        )
        analysis = {
            "breach_landscape": {
                "stat_cards": [
                    {"label": "Total Incidents", "value": "20", "prior_value": "N/A", "change_pct": "N/A"},
                    {"label": "Est. Total Impact", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                    {"label": "Records Exposed", "value": "N/A", "prior_value": "N/A", "change_pct": "N/A"},
                ],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        cards = {c["label"]: c["value"] for c in analysis["breach_landscape"]["stat_cards"]}
        assert cards["Total Incidents"] == "20"  # not shrunk to 2
        # Records Exposed left N/A in lag mode (an incomplete dataset would undercount);
        # impact is rescaled to the live count (20 x ~$4.88M default per-breach).
        assert cards["Est. Total Impact"] == "$98M"

    def test_breach_dataset_absent_leaves_values(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))  # no dataset supplied
        analysis = {
            "breach_landscape": {
                "stat_cards": [{"label": "Total Incidents", "value": "20", "prior_value": "N/A", "change_pct": "N/A"}],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        assert analysis["breach_landscape"]["stat_cards"][0]["value"] == "20"

    def test_prior_quarter_stats_stay_na_when_no_baseline(self, generator):
        from src.core.reporting_period import make_period

        generator.set_reporting_period(make_period(2026, "Q2"))
        analysis = {
            "breach_landscape": {
                "stat_cards": [
                    {"value": "20", "label": "Total Incidents", "prior_value": "N/A", "change_pct": "N/A"}
                ],
                "incidents_by_type": [],
            }
        }
        generator.generate(analysis)
        card = analysis["breach_landscape"]["stat_cards"][0]
        # No Q1 history -> honest N/A, no fabricated prior or percent.
        assert card["prior_value"] == "N/A"
        assert card["change_pct"] == "N/A"


class TestReportTypesList:
    """Tests for list_report_types functionality."""

    def test_list_includes_weekly(self):
        """list_report_types should include 'weekly'."""
        types = list_report_types()
        assert "weekly" in types

    def test_list_includes_quarterly(self):
        """list_report_types should include 'quarterly'."""
        types = list_report_types()
        assert "quarterly" in types

    def test_both_types_registered(self):
        """Both weekly and quarterly should be registered."""
        assert len(REPORT_REGISTRY) >= 2
        assert "weekly" in REPORT_REGISTRY
        assert "quarterly" in REPORT_REGISTRY


# =============================================================================
# Inline citation subscript rendering (brackets kept, document-wide)
# =============================================================================


class TestCitationSubscripts:
    def test_split_markers_keeps_brackets(self):
        from src.reports.base import _split_citation_markers

        parts = _split_citation_markers("system [3][4]. Done.")
        assert parts == [("system ", False), ("[3][4]", True), (". Done.", False)]

    def test_no_citations_returns_single_segment(self):
        from src.reports.base import _split_citation_markers

        assert _split_citation_markers("no citations here") == [("no citations here", False)]

    def test_paragraph_citations_become_subscript_with_brackets(self):
        from docx import Document

        from src.reports.base import _subscript_citations_in_paragraph

        doc = Document()
        para = doc.add_paragraph("breach at Stadler Rail [1] this week [2].")
        _subscript_citations_in_paragraph(para)

        superscript_runs = [r.text for r in para.runs if r.font.superscript]
        # Brackets are KEPT and the markers are superscripted.
        assert superscript_runs == ["[1]", "[2]"]
        assert "Stadler Rail" in "".join(r.text for r in para.runs)

    def test_document_wide_pass_covers_table_cells(self):
        from docx import Document

        from src.reports.weekly_report import WeeklyReportGenerator

        gen = WeeklyReportGenerator.__new__(WeeklyReportGenerator)
        gen.doc = Document()
        gen.doc.add_paragraph("Summary cites [1].")
        table = gen.doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].add_run("[3] CrowdStrike")

        gen._subscript_all_citations()

        # Body paragraph citation superscripted...
        body = gen.doc.paragraphs[0]
        assert [r.text for r in body.runs if r.font.superscript] == ["[1]"]
        # ...and the table-cell citation too.
        cell_para = table.rows[0].cells[0].paragraphs[0]
        assert [r.text for r in cell_para.runs if r.font.superscript] == ["[3]"]
        assert "CrowdStrike" in "".join(r.text for r in cell_para.runs)


class TestOTUnifiedTable:
    """The OT section is one table combining asset-detected vulns and owned advisories."""

    def _gen(self):
        from docx import Document

        gen = WeeklyReportGenerator.__new__(WeeklyReportGenerator)
        gen.doc = Document()
        return gen

    def test_combines_assets_and_advisories_in_one_table(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-2020-1467"], "cvss": 10.0, "affected_ot_devices_count": 13,
                 "is_known_exploited": True, "in_cisa_kev": True,
                 "description": "An elevation of privilege vulnerability exists when Windows "
                                "improperly handles hard links."},
            ],
            "ot_advisories": [
                {"advisory_id": "ICSA-26-190-02", "title": "Schneider PowerChute",
                 "products_affected": "PowerChute Serial Shutdown <=1.4",
                 "cves": ["CVE-2026-2399", "CVE-2026-2404"], "severity": "medium", "cvss": 6.1,
                 "affected_assets": 2, "claroty_status": "ok"},
            ],
        })
        tables = gen.doc.tables
        assert len(tables) == 1                    # ONE table, no lens split
        assert len(tables[0].columns) == 6
        hdr = [c.text for c in tables[0].rows[0].cells]
        assert hdr == ["Vulnerability", "Product", "Description", "Devices", "CVSS", "Exploited"]
        text = "\n".join(c.text for row in tables[0].rows for c in row.cells)
        assert "CVE-2020-1467" in text and "ICSA-26-190-02" in text  # both sources present
        assert "Windows" in text                                     # product parsed from description
        assert "PowerChute Serial Shutdown <=1.4" in text            # advisory product
        assert "CVE-2026-2399" in text                               # advisory CVE list shown

    def test_exploited_sorts_first(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-A"], "cvss": 10.0, "affected_ot_devices_count": 1,
                 "is_known_exploited": False, "description": "not exploited but CVSS 10"},
                {"cve_ids": ["CVE-B"], "cvss": 8.0, "affected_ot_devices_count": 1,
                 "is_known_exploited": True, "in_cisa_kev": True, "description": "exploited"},
            ],
        })
        first = gen.doc.tables[0].rows[1]
        assert "CVE-B" in first.cells[0].text          # exploited outranks higher CVSS
        assert "In the wild (KEV)" in first.cells[5].text  # Exploited column

    def test_ranked_by_device_count_within_exploited_group(self):
        # Same exploited status -> the one on more of your devices ranks higher.
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-FEW"], "cvss": 10.0, "affected_ot_devices_count": 3,
                 "is_known_exploited": True, "in_cisa_kev": True, "description": "few devices"},
                {"cve_ids": ["CVE-MANY"], "cvss": 7.5, "affected_ot_devices_count": 40,
                 "is_known_exploited": True, "in_cisa_kev": True, "description": "many devices"},
            ],
        })
        rows = gen.doc.tables[0].rows
        assert "CVE-MANY" in rows[1].cells[0].text   # 40 devices outranks 3, despite lower CVSS
        assert "CVE-FEW" in rows[2].cells[0].text

    def test_product_parsed_from_description(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-1"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "description": "Unspecified vulnerability in Oracle Java SE 7u67 and 8u20."},
                {"cve_ids": ["CVE-2"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "description": "Sandbox escape fixed in Firefox 148 and Thunderbird 148."},
            ],
        })
        text = "\n".join(c.text for row in gen.doc.tables[0].rows for c in row.cells)
        assert "Oracle Java SE" in text
        assert "Firefox" in text

    def test_product_prefers_kev_and_falls_back_to_dash(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-1"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "description": "obscure issue", "kev_affected_product": "Acme PLC"},
                {"cve_ids": ["CVE-2"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "description": "totally unrecognizable phrasing here"},
            ],
        })
        by_cve = {r.cells[0].text.split()[0]: r for r in gen.doc.tables[0].rows[1:]}
        assert by_cve["CVE-1"].cells[1].text.strip() == "Acme PLC"   # KEV product wins
        assert by_cve["CVE-2"].cells[1].text.strip() == "—"          # nothing recognizable

    def test_exploited_column_shows_kev_and_ransomware(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-1"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "is_known_exploited": True, "known_ransomware": True, "description": "d"},
            ],
        })
        exploited_cell = gen.doc.tables[0].rows[1].cells[5]
        assert "In the wild (KEV)" in exploited_cell.text
        assert "Ransomware" in exploited_cell.text

    def test_exploited_column_grades_exploitdb_and_epss(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                # public exploit exists (ExploitDB) but not KEV -> "Exploit public"
                {"cve_ids": ["CVE-PUB"], "cvss": 9.0, "affected_ot_devices_count": 3,
                 "is_known_exploited": False, "exploits_count": 2, "epss": 0.1, "description": "d"},
                # high EPSS, no KEV, no public exploit -> "Likely (EPSS xx%)"
                {"cve_ids": ["CVE-EPSS"], "cvss": 9.0, "affected_ot_devices_count": 2,
                 "is_known_exploited": False, "exploits_count": 0, "epss": 0.74, "description": "d"},
                # nothing -> dash
                {"cve_ids": ["CVE-NONE"], "cvss": 9.0, "affected_ot_devices_count": 1,
                 "is_known_exploited": False, "exploits_count": 0, "epss": 0.02, "description": "d"},
            ],
        })
        by_cve = {r.cells[0].text.split()[0]: r for r in gen.doc.tables[0].rows[1:]}
        assert "Exploit public" in by_cve["CVE-PUB"].cells[5].text
        assert "Likely (EPSS 74%)" in by_cve["CVE-EPSS"].cells[5].text
        assert by_cve["CVE-NONE"].cells[5].text.strip() == "—"

    def test_only_owned_advisories_included(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_advisories": [
                {"advisory_id": "ICSA-OWNED", "title": "owned", "cves": ["CVE-1"], "cvss": 7.5,
                 "affected_assets": 2, "claroty_status": "ok"},
                {"advisory_id": "ICSA-NOT-OWNED", "title": "not owned", "cves": ["CVE-2"], "cvss": 9.8,
                 "affected_assets": 0, "claroty_status": "ok"},
            ],
        })
        text = "\n".join(c.text for row in gen.doc.tables[0].rows for c in row.cells)
        assert "ICSA-OWNED" in text
        assert "ICSA-NOT-OWNED" not in text

    def test_groups_same_product_cves_into_one_row(self):
        gen = self._gen()
        firefox = [
            {"cve_ids": [f"CVE-2026-27{i}"], "cvss": 10.0, "affected_ot_devices_count": 1,
             "epss": 0.03, "description": f"Sandbox escape {i} fixed in Firefox 148."}
            for i in range(5)
        ]
        gen._add_ot_advisories({
            "ot_environment_exposure": firefox + [
                {"cve_ids": ["CVE-2020-1467"], "cvss": 10.0, "affected_ot_devices_count": 13,
                 "is_known_exploited": True, "in_cisa_kev": True,
                 "description": "Windows elevation of privilege via hard links."},
            ],
        })
        text = "\n".join(c.text for row in gen.doc.tables[0].rows for c in row.cells)
        assert "Firefox — 5 CVEs" in text            # 5 Firefox rows collapsed into one
        assert "CVE-2020-1467" in text               # single-product row stays individual
        # header + one Firefox group row + one Windows row = 3
        assert len(gen.doc.tables[0].rows) == 3

    def test_group_aggregates_strongest_exploitation(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [
                {"cve_ids": ["CVE-1"], "cvss": 9.0, "affected_ot_devices_count": 2,
                 "exploits_count": 0, "epss": 0.1, "description": "Firefox issue one."},
                {"cve_ids": ["CVE-2"], "cvss": 9.0, "affected_ot_devices_count": 2,
                 "is_known_exploited": True, "in_cisa_kev": True, "description": "Firefox issue two."},
            ],
        })
        # Group of two Firefox CVEs; one is KEV -> the group reads "In the wild (KEV)".
        row = gen.doc.tables[0].rows[1]
        assert "Firefox — 2 CVEs" in row.cells[0].text
        assert "In the wild (KEV)" in row.cells[5].text

    def test_advisory_devices_uses_env_assets(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_advisories": [
                {"advisory_id": "ICSA-1", "title": "t", "cves": ["CVE-1"], "cvss": 8.0,
                 "affected_assets": 7, "claroty_status": "ok"},
            ],
        })
        assert gen.doc.tables[0].rows[1].cells[3].text.strip() == "7"  # Env. Assets -> Devices

    def test_empty_shows_note_no_table(self):
        gen = self._gen()
        gen._add_ot_advisories({"ot_environment_exposure": [], "ot_advisories": []})
        assert gen.doc.tables == []
        body = "\n".join(p.text for p in gen.doc.paragraphs)
        assert "No High/Critical OT vulnerabilities or matching advisories" in body

    def test_claroty_error_note_when_nothing_rendered(self):
        gen = self._gen()
        gen._add_ot_advisories({
            "ot_environment_exposure": [],
            "ot_advisories": [
                {"advisory_id": "ICSA-1", "title": "t", "cves": ["CVE-1"], "cvss": 9.8,
                 "affected_assets": 0, "claroty_status": "error"},
            ],
        })
        assert gen.doc.tables == []
        body = "\n".join(p.text for p in gen.doc.paragraphs)
        assert "did not complete this run" in body


class TestPeerIncidentsFromIntel471:
    """Peer Incidents render directly from the Intel471 breach-alert feed."""

    def _gen(self):
        from docx import Document

        gen = WeeklyReportGenerator.__new__(WeeklyReportGenerator)
        gen.doc = Document()
        return gen

    def test_extract_maps_dedups_and_caps(self):
        gen = self._gen()
        breaches = [
            {"organization": "Shalina Healthcare", "incident_type": "Ransomware", "date": "2026-08-15T00:00:00"},
            {"organization": "Colla Health Inc.", "incident_type": "Data Leak", "date": "2026-08-14T10:00:00"},
            {"organization": "Shalina Healthcare", "incident_type": "Ransomware", "date": "2026-08-15"},  # dup org
            {"organization": "", "incident_type": "Breach", "date": "2026-08-13"},  # no org -> skipped
        ]
        out = gen._extract_intel471_breaches({"intel471_breaches": breaches})
        orgs = [i["organization"] for i in out]
        assert orgs == ["Shalina Healthcare", "Colla Health Inc."]  # deduped, most-recent first
        assert all(i["source"] == "Intel471" for i in out)
        assert all(len(i["date"]) == 10 for i in out)  # YYYY-MM-DD

    def test_section_renders_intel471_breach_orgs(self):
        gen = self._gen()
        gen._add_industry_incidents({
            "industry_incidents": [],
            "osint_sources_used": [],
            "intel471_breaches": [
                {"organization": "UAB Biotecha", "incident_type": "Data Leak", "date": "2026-08-12"},
                {"organization": "Subra Pharmacies", "incident_type": "Ransomware", "date": "2026-08-12"},
            ],
        })
        text = "\n".join(
            c.text for t in gen.doc.tables for row in t.rows for c in row.cells
        )
        assert "UAB Biotecha" in text
        assert "Subra Pharmacies" in text

    def test_ai_and_intel471_incidents_deduped_by_org(self):
        gen = self._gen()
        gen._add_industry_incidents({
            # AI already reported Molina; Intel471 also has it -> should appear once.
            "industry_incidents": [
                {"organization": "Molina Healthcare", "incident_type": "Unauthorized Access",
                 "date": "2026-08-12", "source": "Intel471"},
            ],
            "osint_sources_used": [],
            "intel471_breaches": [
                {"organization": "Molina Healthcare", "incident_type": "Breach", "date": "2026-08-12"},
                {"organization": "Colla Health Inc.", "incident_type": "Data Leak", "date": "2026-08-14"},
            ],
        })
        org_cells = [t.rows[r].cells[0].text for t in gen.doc.tables for r in range(1, len(t.rows))]
        assert org_cells.count("Molina Healthcare") == 1  # not double-listed
        assert "Colla Health Inc." in org_cells

    def test_ai_intel471_prose_mined_incidents_dropped(self):
        # An AI incident attributed to Intel471 (e.g. Amgen scraped from an INTSUM narrative)
        # must NOT render — Intel471 peers come only from the structured breach-alert feed.
        gen = self._gen()
        gen._add_industry_incidents({
            "industry_incidents": [
                {"organization": "Amgen", "incident_type": "Breach", "date": "2026-08-14",
                 "source": "Intel471"},  # prose-mined -> dropped
                {"organization": "Scottish Government", "incident_type": "Breach",
                 "date": "2026-08-13", "source": "Dark Reading", "osint_citation_number": 1},
            ],
            "osint_sources_used": [
                {"title": "Scottish Govt Data Breach", "url": "https://x", "citation_number": 1},
            ],
            "intel471_breaches": [],
        })
        org_cells = [t.rows[r].cells[0].text for t in gen.doc.tables for r in range(1, len(t.rows))]
        assert "Amgen" not in org_cells                 # Intel471-attributed AI incident dropped
        assert "Scottish Government" in org_cells        # OSINT-sourced AI incident kept

    def test_watchlist_lenses_survive_cap_and_lead(self):
        # 25 sector peers + 1 competitor + 1 supply-chain. The two watchlist hits must survive
        # the 20-row cap and appear first (competitor before supply chain).
        gen = self._gen()
        breaches = [
            {"organization": f"Sector Clinic {i}", "incident_type": "Ransomware",
             "date": f"2026-08-{10 + (i % 6):02d}", "relevance": "sector"}
            for i in range(25)
        ]
        breaches.append({"organization": "Qiagen NV", "incident_type": "Data Leak",
                         "date": "2026-08-10", "relevance": "competitor"})
        breaches.append({"organization": "Okta Inc.", "incident_type": "Unauthorized Access",
                         "date": "2026-08-10", "relevance": "supply-chain"})
        gen._add_industry_incidents({
            "industry_incidents": [], "osint_sources_used": [], "intel471_breaches": breaches,
        })
        org_col = [t.rows[r].cells[0].text for t in gen.doc.tables for r in range(1, len(t.rows))]
        assert len(org_col) == 20                         # capped
        assert "Qiagen NV" in org_col and "Okta Inc." in org_col   # watchlist hits survived
        assert org_col[0] == "Qiagen NV"                  # competitor leads
        assert org_col[1] == "Okta Inc."                  # supply chain second

    def test_relevance_column_labels_each_lens(self):
        gen = self._gen()
        gen._add_industry_incidents({
            # OSINT-sourced AI incident with no explicit lens -> classified from the org name
            "industry_incidents": [
                {"organization": "Pacific Biosciences", "incident_type": "Breach",
                 "date": "2026-08-12", "source": "Dark Reading", "osint_citation_number": 1},
            ],
            "osint_sources_used": [
                {"title": "PacBio Breach", "url": "https://x", "citation_number": 1},
            ],
            "intel471_breaches": [
                {"organization": "Okta", "incident_type": "Unauthorized Access",
                 "date": "2026-08-14", "relevance": "supply-chain"},
                {"organization": "Colla Health Inc.", "incident_type": "Data Leak",
                 "date": "2026-08-13", "relevance": "sector"},
            ],
        })
        # Map organization (col 0) -> relevance label (col 3)
        rel = {}
        for t in gen.doc.tables:
            for r in range(1, len(t.rows)):
                rel[t.rows[r].cells[0].text] = t.rows[r].cells[3].text
        assert rel["Pacific Biosciences"] == "Competitor"   # classified from name, no explicit lens
        assert rel["Okta"] == "Supply Chain"
        assert rel["Colla Health Inc."] == "Sector Peer"
